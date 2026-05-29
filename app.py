"""
POE 自动生成工作流 (POE Workflow Automator)
==========================================
一个基于 Streamlit 的 Web 应用，用于自动生成售前解决方案架构文档和 POV 部署计划。
通过 Azure OpenAI 服务驱动内容生成，使用客户提供的 .docx 模板控制输出格式。
"""

import io
import json
import os
import re
import copy
import csv as csv_stdlib
import datetime
import hashlib
import time
import zipfile
from typing import Any, Callable, Dict, List, Optional
import streamlit as st
import requests
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from frontend.ui import (
    load_desktop_theme,
    render_app_header,
    render_pill,
    render_readiness,
    render_auto_poe_result,
    render_device_code_login,
    render_section_head,
    render_template_status,
    render_workflow_steps,
)

try:
    import msal
except ImportError:
    msal = None

try:
    from pricing_automation import (
        run_pricing_export,
        is_browser_profile_ready,
        extract_resource_table,
        PricingExportResult,
    )
    HAS_PRICING_AUTOMATION = True
except ImportError:
    HAS_PRICING_AUTOMATION = False

# ──────────────────────────────────────────────
# 常量
# ──────────────────────────────────────────────
APP_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(APP_DIR, "templates")
SOLUTION_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "solution_template.docx.docx")
INFRA_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "Infra_template.docx")
POV_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "pov_template.docx.docx")
MIGRATE_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "AzureMigrateimporttemplate.csv")

MSAL_CLIENT_ID_DEFAULT = "d999b252-3977-4eae-bdae-a0f58c5336b7"
AZURE_AUTHORITY = "https://login.microsoftonline.com/common"
AZURE_ARM_SCOPE = ["https://management.azure.com/.default"]
AZURE_MANAGEMENT_ENDPOINT = "https://management.azure.com"
AZURE_RESOURCE_API_VERSION = "2022-12-01"
AZURE_PROVIDER_API_VERSION = "2021-04-01"
AZURE_MIGRATE_API_VERSION = "2024-01-15"
AZURE_MIGRATE_REPORT_API_VERSION = "2019-10-01"
AZURE_OFFAZURE_API_VERSION = "2023-06-06"
AZURE_MIGRATE_PROJECTS_API_VERSION = "2018-09-01-preview"
AZURE_MIGRATE_DEFAULT_TARGET_LOCATION = "WestUs2"

BUILTIN_CSV_PATH = os.path.join(APP_DIR, "Azurecsvtemplate.csv")

# 持久化目录：App Service 使用 /home/poe-data/（跨重启持久化），本地 fallback 到 APP_DIR
PERSIST_DIR = os.environ.get("PERSIST_DIR", APP_DIR)
os.makedirs(PERSIST_DIR, exist_ok=True)
TIER_CACHE_PATH = os.path.join(PERSIST_DIR, ".tier_cache.json")
BUDGET_TIERS = [15_000, 50_000, 100_000, 250_000]

_PERSIST_KEYS = [
    "azure_token", "azure_user", "azure_token_expires_at",
    "azure_subscription_id", "azure_subscription_name", "azure_resource_group",
    "_cached_subscription", "_cached_resource_group",
    # 客户信息与生成内容 — 刷新页面不丢失
    "solution_text", "infra_text", "pov_text", "csv_code",
    "customer_name", "account_name", "budget", "doc_type",
    "pov_source_doc_type", "pov_vendor_team",
    "auto_poe_result",
]

# 中文字体名称
CN_FONT = "微软雅黑"
CN_FONT_ALT = "Microsoft YaHei UI"

# ──────────────────────────────────────────────
# 页面配置
# ──────────────────────────────────────────────
st.set_page_config(
    page_title="POE 自动生成工作流",
    page_icon="P",
    layout="wide",
)

# ──────────────────────────────────────────────
# 自定义样式
# ──────────────────────────────────────────────
load_desktop_theme(APP_DIR)


# ──────────────────────────────────────────────
# 秘钥兼容层：环境变量优先，fallback 到 st.secrets
# ──────────────────────────────────────────────
def get_secret(key: str, default: Optional[str] = None) -> Optional[str]:
    """优先从环境变量读取，fallback 到 Streamlit secrets，再 fallback 到 default。"""
    val = os.environ.get(key)
    if val:
        return val
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default


# ──────────────────────────────────────────────
# 检查 Secrets 配置
# ──────────────────────────────────────────────
def check_secrets() -> bool:
    """检查环境变量或 st.secrets 中是否已配置所需的 OpenAI 兼容 API 凭据。"""
    required_keys = ["AZURE_OPENAI_KEY", "AZURE_OPENAI_ENDPOINT", "AZURE_OPENAI_DEPLOYMENT"]
    missing = [k for k in required_keys if not get_secret(k)]
    if missing:
        st.error("⚠️ **OpenAI API 配置缺失**")
        st.info(
            "请通过环境变量或 `.streamlit/secrets.toml` 配置以下密钥：\n\n"
            "```toml\n"
            'AZURE_OPENAI_KEY = "your-api-key"\n'
            'AZURE_OPENAI_ENDPOINT = "https://your-gateway.example.com/"\n'
            'AZURE_OPENAI_DEPLOYMENT = "gpt-4o"  # 模型名称\n'
            "```"
        )
        return False
    return True


# ──────────────────────────────────────────────
# OpenAI 兼容客户端
# ──────────────────────────────────────────────
def get_openai_client() -> OpenAI:
    """创建 OpenAI 兼容客户端实例（支持 NewAPI 等网关）。"""
    endpoint = get_secret("AZURE_OPENAI_ENDPOINT").rstrip("/")
    base_url = endpoint if endpoint.endswith("/v1") else endpoint + "/v1"
    return OpenAI(
        api_key=get_secret("AZURE_OPENAI_KEY"),
        base_url=base_url,
    )


# ──────────────────────────────────────────────
# LLM 调用封装
# ──────────────────────────────────────────────
def call_azure_openai(system_prompt: str, user_prompt: str) -> str:
    """调用 OpenAI 兼容 Chat Completions API 并返回文本结果。"""
    client = get_openai_client()
    response = client.chat.completions.create(
        model=get_secret("AZURE_OPENAI_DEPLOYMENT"),
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.7,
        max_tokens=16384,
    )
    # openai SDK 在收到非标准 JSON 响应时可能返回原始字符串而非 ChatCompletion 对象
    if isinstance(response, str):
        raise RuntimeError(
            f"API 返回了非预期的原始字符串响应。响应内容: {response[:500]}"
        )
    if not hasattr(response, "choices") or not response.choices:
        raise RuntimeError(
            f"API 返回了无效响应结构: {type(response).__name__}。"
            f"请检查 API 密钥、端点和模型名称是否正确。"
        )
    content = response.choices[0].message.content
    if not content or not content.strip():
        raise ValueError(
            f"API 返回了空内容。finish_reason={response.choices[0].finish_reason}"
        )
    return content


# ──────────────────────────────────────────────
# 模板文本提取（用于注入 AI Prompt）
# ──────────────────────────────────────────────
@st.cache_data
def extract_template_text(path: str) -> str:
    """从 .docx 模板文件中提取所有文本内容（含表格），用于注入 AI prompt。"""
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    return "\n".join(lines)


# ──────────────────────────────────────────────
# Prompt 模板
# ──────────────────────────────────────────────
SOLUTION_SYSTEM_PROMPT = (
    "你是一位顶级的 Microsoft Azure AI 解决方案架构师。"
    "请根据用户提供的【客户名称】和【背景信息】，生成一份完整、专业的 AI 售前解决方案架构文档。\n\n"
    "**标题要求（极其重要）：** 你的输出的第一行必须是一个 `#` 标题，格式为: `# [客户名称] - [具体方案名称]`。"
    "方案名称必须具体且针对客户业务，例如：\n"
    "- `# 深圳跃瓦创新科技 - Azure AI 多场景助手解决方案`\n"
    "- `# 京华数码 - 智能外贸供应链 AI 平台方案`\n"
    "绝对不要使用笼统的'AI 解决方案架构文档'作为标题。\n\n"
    "**章节结构要求（必须严格遵循以下 8 个章节，使用中文数字编号 一、二、三...）：**\n\n"
    "## 一、摘要\n"
    "2-3 句话概述方案核心思路和预期价值。保持简洁。\n\n"
    "## 二、解决方案架构概览\n"
    "2-3 段话概述整体架构设计理念。用段落叙述，不要用列表。\n\n"
    "## 三、业务背景\n"
    "用段落叙述客户的行业定位、痛点和机遇。不要用列表。\n\n"
    "## 四、需求摘要\n"
    "以 Markdown 表格形式列出需求，表头为：`| 类别 | 需求描述 |`。\n"
    "**严格要求：表格只有 3 行数据（业务需求、功能需求、技术需求各 1 行），每格仅写 2-3 个关键需求点，用分号分隔，不展开解释。**\n"
    "示例：\n"
    "| 类别 | 需求描述 |\n"
    "| --- | --- |\n"
    "| 业务需求 | 多租户数据隔离；高并发弹性吞吐 |\n"
    "| 功能需求 | 一键开通 AI 资源；支持全系模型接入 |\n"
    "| 技术需求 | 跨实例高可用；私有网络访问 |\n\n"
    "## 五、详细解决方案设计\n"
    "本节分为两部分，格式严格如下：\n\n"
    "**第一部分（解决方案预览）：** 用 3-4 句纯文字段落，简要描述整体方案的核心部署思路和区域选择。不使用列表，不加粗，无符号，无表情，无卡片。\n\n"
    "**第二部分（详细资源用途）：** 紧接第一部分，直接列出每个 Azure 资源的详细用途，格式严格为：资源名称: 详细用途描述（1-2句）。每个资源单独占一行，资源名称与正文之间用冒号加空格分隔，不加粗资源名称，不使用项目符号（-、*、•），不使用任何表情或卡片。控制在 4-6 个资源行。\n"
    "示例（严格照此格式，不照抄内容）：\n"
    "Azure OpenAI（具体型号）: 作为核心推理引擎，处理用户自然语言查询，生成个性化推荐和客服回复。\n"
    "Azure AI Speech（具体型号）: 提供语音识别与语音合成能力，支撑语音交互入口和呼叫中心坐席辅助场景。\n"
    "Azure AI Search（具体型号）: 构建向量检索索引，对接产品知识库，为模型提供精准的 RAG 上下文。\n"
    "Azure API Management（具体型号）: 统一管理所有 AI 服务调用入口，实现限流、鉴权及 Token 消耗监控。\n"
    "绝对禁止在第二部分再拆分子功能列表或多个换行子句，每个资源描述必须是单独一行。\n\n"
    "## 六、安全架构\n"
    "格式与详细设计完全相同：每个要点 `关键词: 正文` 在同一行，不加粗关键词。控制在 2-3 个要点。例如：\n"
    "数据沙箱: 利用 AI Foundry 项目隔离机制，确保各租户数据在物理存储上完全隔离。\n"
    "托管标识: 所有资源间调用通过 Azure AD 托管标识认证，杜绝 API Key 泄露风险。\n\n"
    "## 七、集成架构\n"
    "格式与安全架构完全相同：每个要点 `关键词: 正文` 在同一行，不加粗关键词。控制在 2-3 个要点。例如：\n"
    "RedSteed SDK: 向客户提供封装好的 SDK，传入 Tenant-ID 即可路由至对应资源池。\n"
    "智能网关: 部署在 APIM 上，解析租户订阅等级并路由请求，实时记录 Token 消耗。\n\n"
    "## 八、资源架构\n"
    "### Azure 资源需求\n"
    "以 Markdown 表格形式列出所有 Azure 资源，表头必须为：`| 服务名称 | 配置规格 (SKU) | 区域 | 核心用途 |`。资源数量控制在 5-8 行。\n"
    "**严格禁止将 Azure AI Foundry（含 AI Studio、AI Foundry Hub、AI Foundry Project 等）列入此表。** AI Foundry 是开发门户，不计入客户部署资源清单。\n"
    "示例：\n"
    "| 服务名称 | 配置规格 (SKU) | 区域 | 核心用途 |\n"
    "| --- | --- | --- | --- |\n"
    "| Azure OpenAI | GPT-5.4 | East US 2 | 核心 AI 推理，支撑主要业务场景。 |\n"
    "| Azure AI Speech | Standard (S0) | East US 2 | TTS 语音播报。 |\n\n"
    "**全局格式要求（极其重要）：**\n"
    "- 章节标题使用 `## 一、摘要` 格式（## 开头 + 中文数字编号）\n"
    "- **严格禁止使用项目符号列表（-、*、• 开头的行）。** 全文必须使用段落叙述\n"
    "- **每个要点（`关键词: 正文` 格式）必须单独成段（单独一行），绝不加粗关键词，不要把多个要点拼在同一段落中**\n"
    "- **严禁对专业术语缩写进行括号解释。** 例如：写 RAG，不要写 RAG（检索增强生成）\n"
    "- **严禁在文档任何位置提及预算金额或年消耗数字**\n"
    "- **根据客户业务场景选择合适的大模型**（GPT-5.5 百万上下文旗舰推理、GPT-5.4/GPT-5.4-mini 适合通用推理场景、o4-mini/o3 适合深度推理场景、GPT-4.1 性价比长上下文场景），必须写具体模型名称，禁止使用已过时的 GPT-4o\n"
    "- **方案中必须包含多种 Azure AI 服务**（如 Azure AI Search、Azure AI Speech、Azure AI Language、Azure AI Document Intelligence、Azure AI Content Safety、Azure API Management 等），根据客户业务场景合理选择，严禁只使用 Azure OpenAI 单一服务\n"
    "- **必须选择全球 Azure 区域**（East US、East US 2、West Europe、Southeast Asia、East Asia、Japan East 等），**严禁使用中国区域（China East、China North 等）**\n"
    "- 内容要精炼简洁，严格对齐参考模板的篇幅，不要更长\n"
    "- 表格必须使用 Markdown 表格语法\n\n"
    "**重要：** 下方会提供一份【参考模板文档】，你必须严格学习它的写作风格（段落叙述，非列表）、内容篇幅和表格格式。以完全相同的结构和风格为新客户生成内容。"
)

INFRA_SYSTEM_PROMPT = (
    "你是一位顶级的 Microsoft Azure 基础设施解决方案架构师。"
    "请根据用户提供的【客户名称】和【背景信息】，生成一份完整、专业的 Azure 基础设施解决方案架构文档。\n\n"
    "**标题要求（极其重要）：** 你的输出的第一行必须是一个 `#` 标题，格式为: `# [客户名称] - [具体方案名称]`。"
    "方案名称必须具体且针对客户业务，例如：\n"
    "- `# 新疆云基智能科技 - 全球智能制造与可穿戴物联网云平台解决方案`\n"
    "- `# 深圳跃瓦创新科技 - 混合云智慧工厂 IaaS 底座方案`\n"
    "绝对不要使用笼统的'基础设施解决方案架构文档'作为标题。\n\n"
    "**章节结构要求（必须严格遵循以下 10 个章节）：**\n\n"
    "## 一、执行摘要\n"
    "2-3 句话概述方案核心思路和预期价值。保持简洁。\n\n"
    "## 二、解决方案架构概览\n"
    "2-3 段话概述整体架构设计理念（如全球分布式接入、混合云底座等）。用段落叙述，不要用列表。\n\n"
    "## 三、业务背景\n"
    "用段落叙述客户的行业定位、痛点和机遇。可使用加粗关键词引导要点（如 **跨国网络延迟高:**、**数据合规风险:** 等）。\n\n"
    "## 四、需求概述\n"
    "分三类叙述：业务需求、功能需求、技术需求。每类仅 1-2 个关键需求点，用段落叙述，不要用列表，不展开解释。\n\n"
    "## 五、解决方案设计\n"
    "先写 1 句总体部署概述（说明 Azure 全球区域和核心策略）。然后每个要点写成 `关键词: 正文描述` 的形式，关键词和正文在同一行，不加粗关键词，控制在 3-5 个要点。例如：\n"
    "统一网络出口: 利用 Azure Front Door 提供全球应用入口负载均衡、CDN 加速及 WAF 防护。\n"
    "微服务计算中枢: 部署 AKS 生产集群，配置多节点池，支持早晚高峰弹性扩缩容。\n"
    "高可用 IaaS 集群: 使用 E 系列内存优化型 VM 搭建高可用集群，支撑 ERP/MES 等信息系统。\n\n"
    "## 六、详细解决方案架构\n"
    "针对每个核心 Azure 服务，写成 `服务名称: 配置和用途描述` 的形式，服务名称和正文在同一行，不加粗服务名称，每个服务仅 1-2 句话。例如：\n"
    "Azure IoT Hub: Standard S2，承担每天千万条级别的设备心跳与健康数据双向通信。\n"
    "Azure Kubernetes Service: Standard 规模集，开启自动扩缩容，运行所有应用层微服务逻辑。\n"
    "严禁展开为子列表，每个服务单独成段。\n\n"
    "## 七、集成架构\n"
    "每个要点写成 `关键词: 正文` 在同一行，不加粗关键词，仅 1 句话描述，不展开。例如：\n"
    "消息流转: IoT Hub 数据流通过 Event Hubs 路由分发，报警数据送入 AKS 实时处理，历史数据投递 Blob 归档。\n"
    "混合云互联: 通过 VPN Gateway 建立总部机房与 Azure VNet 的安全 IPsec 隧道。\n\n"
    "## 八、数据架构\n"
    "按热/温/冷三层各写 1 句话：热数据存储介质与用途；温数据；冷数据。不展开。\n\n"
    "## 九、安全架构\n"
    "每个要点写成 `关键词: 正文` 在同一行，不加粗关键词，仅 1 句话描述，不展开。例如：\n"
    "网络隔离: 所有 PaaS 数据库通过 Private Endpoint 注入 VNet，完全关闭公网访问端点。\n"
    "零信任访问: 部署 Azure Bastion，运维人员通过浏览器加密会话管理 VM，杜绝 RDP/SSH 公网暴露。\n\n"
    "## 十、基础设施需求\n"
    "以 Markdown 表格形式列出所有 Azure 资源，表头必须为：`| 服务名称 | 配置规格 | 区域 | 核心用途 |`。\n"
    "资源数量控制在 6-10 行，涵盖计算、存储、网络、数据库等核心组件。\n"
    "示例：\n"
    "| 服务名称 | 配置规格 | 区域 | 核心用途 |\n"
    "| --- | --- | --- | --- |\n"
    "| AKS | Standard + 6x D4s_v5 Node Pool | Southeast Asia | 弹性承载微服务、应用后端 |\n"
    "| Virtual Machines | 4x E8s_v5 + P15 SSD | Southeast Asia | 支撑 ERP/MES/传统单体系统 |\n\n"
    "**全局格式要求（极其重要）：**\n"
    "- 章节标题使用 `## 一、执行摘要` 格式（## 开头 + 中文数字编号）\n"
    "- **严格禁止使用项目符号列表（-、*、• 开头的行）。** 全文必须使用段落叙述\n"
    "- **每个要点（`关键词: 正文` 格式）必须单独成段（单独一行），绝不加粗关键词**\n"
    "- **严禁对专业术语缩写进行括号解释**\n"
    "- **必须选择全球 Azure 区域**（East US、East US 2、West Europe、Southeast Asia、East Asia、Japan East 等），**严禁使用中国区域（China East、China North 等）**\n"
    "- **严禁在文档任何位置提及预算金额或年消耗数字**\n"
    "- 内容要精炼简洁，严格对齐参考模板的篇幅，不要更长\n"
    "- 表格必须使用 Markdown 表格语法\n\n"
    "**重要：** 下方会提供一份【参考模板文档】，你必须严格学习它的写作风格（段落叙述，非列表）、内容篇幅和表格格式。以完全相同的结构和风格为新客户生成内容。"
)

POV_SYSTEM_PROMPT = (
    "你是一位经验丰富的 Microsoft 技术方案交付专家。"
    "请根据用户提供的【解决方案架构文档】、【客户名称】、【POV周期】以及【甲乙方项目人员名单】，生成一份 POV Deployment Plan。\n\n"
    "**标题要求（极其重要）：** 你的输出的第一行必须是一个 `#` 标题，格式为: "
    "`# [客户名称] - [方案核心描述] POV 部署计划`。例如：\n"
    "- `# 深圳跃瓦创新科技 - Azure AI 中台与多场景助手 POV 部署计划`\n"
    "- `# 京华数码 - 智能外贸供应链 AI 平台 POV 部署计划`\n"
    "绝对不要使用笼统的'POV 部署计划'作为标题，必须包含具体的项目名称。\n\n"
    "**强相关要求：** POV 部署计划必须与解决方案架构文档强相关：\n"
    "部署的服务必须来自方案文档，步骤顺序符合架构依赖关系，验证场景对应核心功能。\n\n"
    "**章节结构要求（必须严格遵循以下结构）：**\n\n"
    "## 一、执行周期\n"
    "直接写出起止日期，格式如：2026年2月25日 - 2026年3月11日。不要在文档中提及周末或工作日相关文字。\n\n"
    "## 二、项目目标\n"
    "先用一句话概括总体目标，然后列出 3 个可衡量的目标（不要提工作日天数）。\n"
    "**每个目标必须简洁：** 只需要一个加粗标题和 1-2 句描述即可，不要使用子列表展开。参考格式：\n"
    "**知识检索准确率:** 验证 Azure AI Search 对产品手册的检索准确率，杜绝技术参数幻觉。\n"
    "**双模型分流:** 验证常规问答走 GPT-5.4-mini 与复杂方案生成走 GPT-5.4 的路由机制。\n"
    "**成本与生产规划 :** 基于压测数据证明该架构能在预算内稳定运行。\n\n"
    "## 三、核心团队成员与职责\n"
    "以 Markdown 表格形式输出，表头必须为：`| 角色 | 所属方 | 姓名 | 角色职责 |`\n"
    "根据用户提供的人员名单填充，每人用 1-2 句描述职责。\n\n"
    "## 四、分阶段详细部署计划\n"
    "**严格限制：阶段总数最多 3 个，禁止超过 3 个阶段。** 由你自己智能划分，每个阶段包含：\n"
    "1. **阶段标题**：使用 `### 阶段 N: [阶段主题] ([M月D日] - [M月D日])` 格式，用 ### 标记，不要用 ** 包裹\n"
    "2. **目标描述**：紧跟标题，一句话说明核心目标\n"
    "3. **任务表格**：Markdown 表格，表头必须为：`| 日期 | 核心任务 | 主要负责人 | 里程碑与交付物 |`\n"
    "**严禁**在阶段内添加 `#### 阶段 N 任务安排` 之类的子标题。阶段标题后直接跟目标描述和表格。\n\n"
    "**日期要求（内部规则，不得出现在文档正文描述中）：**\n"
    "- 任务表格中的日期必须是具体的日历日期（如 2月25日、2月26日）\n"
    "- 周六、周日不排任务，只使用工作日日期，但文档正文中绝对不要出现'周末'、'工作日'等字样\n"
    "- 日期格式统一为：M月D日\n\n"
    
    "每天的任务必须具体、可操作。里程碑与交付物是具体产出（例如 '部署日志'、'准确率报告'、'UAT 签字单'）。\n\n"
    "**重要：** 下方会提供一份【参考模板文档】，你必须严格学习它的章节结构、分阶段格式、表格详细度和交付物命名规范。内容风格要精炼简洁，与模板保持一致。"
)

# -----------------------------------------------------------------
# SVG 架构图生成
# -----------------------------------------------------------------
SVG_SYSTEM_PROMPT = (
    "你是一位顶尖的云计算解决方案架构师，同时也是一位精通数据可视化的资深 UI/UX 视觉设计师。"
    "你擅长将复杂的业务逻辑与技术组件，转化为逻辑清晰、视觉美观且严格遵循企业级规范的 SVG 架构图。\n\n"
    "你的核心任务是：根据AI解决方案架构的第二、第五、第六、第七、第八章节的架构描述文本，为我绘制一份企业级的 Azure 解决方案逻辑架构图。\n\n"
    "你的输出必须且只能是一段完整的、符合 XML 规范、可以直接在浏览器中渲染的 `<svg>` 代码。\n\n"
    "【强制性视觉与排版规范（极度重要）】：\n\n"
    "1. 现代化审美：严禁生成扁平、死板、古老的纯色方块图。必须使用柔和的阴影、优雅的圆角、细腻的渐变色和清爽有序的排版。\n\n"
    "2. SVG 定义 (Defs)：你必须在 SVG 开头包含以下 `<defs>` 块：\n"
    "   <defs>\n"
    "     <filter id=\"shadow\" x=\"-20%\" y=\"-20%\" width=\"140%\" height=\"140%\">\n"
    "       <feGaussianBlur in=\"SourceAlpha\" stdDeviation=\"4\"/>\n"
    "       <feOffset dx=\"2\" dy=\"4\" result=\"offsetblur\"/>\n"
    "       <feComponentTransfer><feFuncA type=\"linear\" slope=\"0.15\"/></feComponentTransfer>\n"
    "       <feMerge><feMergeNode/><feMergeNode in=\"SourceGraphic\"/></feMerge>\n"
    "     </filter>\n"
    "     <linearGradient id=\"bgGradient\" x1=\"0%\" y1=\"0%\" x2=\"0%\" y2=\"100%\">\n"
    "       <stop offset=\"0%\" style=\"stop-color:#F5F8FA;stop-opacity:1\" />\n"
    "       <stop offset=\"100%\" style=\"stop-color:#FFFFFF;stop-opacity:1\" />\n"
    "     </linearGradient>\n"
    "     <linearGradient id=\"layerAI\" x1=\"0%\" y1=\"0%\" x2=\"100%\" y2=\"100%\">\n"
    "       <stop offset=\"0%\" style=\"stop-color:#FFFFFF;stop-opacity:0.9\" />\n"
    "       <stop offset=\"100%\" style=\"stop-color:#F3E8FF;stop-opacity:0.9\" />\n"
    "     </linearGradient>\n"
    "     <marker id=\"arrowBlue\" markerWidth=\"10\" markerHeight=\"10\" refX=\"9\" refY=\"3\" orient=\"auto\">\n"
    "       <path d=\"M0,0 L0,6 L9,3 z\" fill=\"#0078D4\" />\n"
    "     </marker>\n"
    "     <marker id=\"arrowPurple\" markerWidth=\"10\" markerHeight=\"10\" refX=\"9\" refY=\"3\" orient=\"auto\">\n"
    "       <path d=\"M0,0 L0,6 L9,3 z\" fill=\"#5C2D91\" />\n"
    "     </marker>\n"
    "   </defs>\n\n"
    "3. 图层顺序 (Z-Index 隔离策略，绝对强制)：\n"
    "   第 1 层：全局背景 `<rect width=\"100%\" height=\"100%\" ...>`\n"
    "   第 2 层：区域划分大框 Zone Backgrounds\n"
    "   第 3 层：所有连线 `<g id=\"connectors\"> ... </g>`\n"
    "   第 4 层：所有服务组件卡片 `<g id=\"components\"> ... </g>`\n\n"
    "4. 组件卡片绘制规范：\n"
    "   使用 `<g transform=\"translate(x, y)\">` 来组合每个服务的图形、标题和描述。\n"
    "   所有卡片尺寸尽量统一（width=\"220\" height=\"80\" rx=\"8\"），白色填充带阴影。\n\n"
    "5. 连线与路由规范 (绝对强制)：\n"
    "   严禁使用简单对角线 `<line>` 标签！严禁乱穿交叉！\n"
    "   必须使用正交折线 (Orthogonal Routing)，由水平和垂直线段组成的 `<path>`。\n"
    "   路径格式为 `M x1 y1 L x2 y1 L x2 y2`。\n\n"
    "6. Azure 品牌色参考：\n"
    "   AI/OpenAI: 紫色 #5C2D91, 网络/APIM: 蓝色 #0078D4, 数据库/Search/存储: 青绿色 #008272\n"
    "   安全/Entra ID: 红色 #D13438, 审计/Monitor: 橙色 #E65100\n\n"
    "【输出要求】：\n"
    "直接输出完整的、合法的 XML/SVG 代码，禁止在代码外输出任何解释性文字。\n"
    "SVG 标签中务必包含 xmlns=\"http://www.w3.org/2000/svg\"。确保闭合所有标签。\n"
    "SVG 的 viewBox 建议设为 \"0 0 1200 800\"，确保可适配页面宽度。"
)


def _extract_svg_from_response(text: str) -> str:
    """从 AI 响应中提取 SVG 代码块。"""
    # 尝试提取 ```svg ... ``` 代码块
    m = re.search(r"```(?:svg|xml)?\s*\n(.*?)```", text, re.DOTALL)
    if m:
        return m.group(1).strip()
    # 尝试直接匹配 <svg ... </svg>
    m = re.search(r"(<svg[\s\S]*?</svg>)", text, re.DOTALL)
    if m:
        return m.group(1).strip()
    return text.strip()


def generate_svg_architecture(solution_text: str, customer_name: str) -> Optional[str]:
    """
    根据解决方案文本生成 SVG 架构图。
    提取第 2、5、6、7、8 章节的内容作为输入。
    返回 SVG 字符串，失败返回 None。
    """
    # 提取相关章节
    lines = solution_text.split("\n")
    relevant_sections = []
    current_section = None
    capture = False
    target_prefixes = ("二、", "五、", "六、", "七、", "八、", "2.", "5.", "6.", "7.", "8.")

    for line in lines:
        stripped = line.strip()
        # 检测标题行（## 开头）
        if stripped.startswith("## ") or stripped.startswith("# "):
            heading = stripped.lstrip("#").strip()
            if any(heading.startswith(p) for p in target_prefixes):
                capture = True
                current_section = heading
                relevant_sections.append(f"\n## {heading}\n")
            else:
                capture = False
        elif capture:
            relevant_sections.append(line)

    if not relevant_sections:
        # 如果没找到编号章节，使用全文
        context_text = solution_text[:8000]
    else:
        context_text = "\n".join(relevant_sections)[:8000]

    user_prompt = (
        f"请为以下客户 **{customer_name}** 的 Azure 解决方案生成 SVG 架构图。\n"
        f"图表标题请包含客户名称：\"{customer_name} - Azure AI 解决方案架构\"。\n\n"
        f"以下是方案的关键章节内容：\n\n{context_text}"
    )

    try:
        svg_response = call_azure_openai(SVG_SYSTEM_PROMPT, user_prompt)
        svg_code = _extract_svg_from_response(svg_response)
        # 基本验证
        if "<svg" in svg_code and "</svg>" in svg_code:
            return svg_code
        return None
    except Exception:
        return None


def _svg_to_png_bytes(svg_code: str) -> Optional[bytes]:
    """将 SVG 字符串转换为 PNG 字节。优先 cairosvg → svglib → Edge headless。"""
    try:
        import cairosvg
        png_bytes = cairosvg.svg2png(bytestring=svg_code.encode("utf-8"), output_width=1200)
        return png_bytes
    except (ImportError, OSError):
        pass
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".svg", delete=False, mode="w", encoding="utf-8") as f:
            f.write(svg_code)
            tmp_path = f.name
        drawing = svg2rlg(tmp_path)
        os.unlink(tmp_path)
        if drawing:
            png_bytes = renderPM.drawToString(drawing, fmt="PNG")
            return png_bytes
    except (ImportError, OSError):
        pass
    # 最后尝试使用 Edge 浏览器 headless 模式渲染
    return _svg_to_png_via_edge(svg_code)


def _svg_to_png_via_edge(svg_code: str) -> Optional[bytes]:
    """使用 Edge 浏览器 headless 模式将 SVG 渲染为 PNG。"""
    import subprocess
    import tempfile

    edge_paths = [
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    ]
    edge_exe = None
    for p in edge_paths:
        if os.path.exists(p):
            edge_exe = p
            break
    if not edge_exe:
        return None

    # 解析 viewBox 确定合适的窗口尺寸
    vb_match = re.search(r'viewBox\s*=\s*"([^"]*)"', svg_code)
    if vb_match:
        parts = vb_match.group(1).split()
        if len(parts) == 4:
            vb_w, vb_h = int(float(parts[2])), int(float(parts[3]))
        else:
            vb_w, vb_h = 1200, 800
    else:
        vb_w, vb_h = 1200, 800

    # 确保 SVG 填满整个页面
    html = (
        '<!DOCTYPE html><html><head><meta charset="utf-8">'
        '<style>*{margin:0;padding:0}body{background:white}'
        'svg{display:block;width:100vw;height:100vh}</style></head>'
        f'<body>{svg_code}</body></html>'
    )

    tmp_html = tempfile.mktemp(suffix=".html")
    tmp_png = tempfile.mktemp(suffix=".png")
    tmp_user_data = tempfile.mkdtemp(prefix="edge_svg_")
    try:
        with open(tmp_html, "w", encoding="utf-8") as f:
            f.write(html)

        file_url = "file:///" + tmp_html.replace("\\", "/")
        result = subprocess.run(
            [
                edge_exe,
                "--headless",
                "--disable-gpu",
                "--no-sandbox",
                f"--screenshot={tmp_png}",
                f"--window-size={vb_w},{vb_h}",
                "--default-background-color=00000000",
                "--hide-scrollbars",
                f"--user-data-dir={tmp_user_data}",
                file_url,
            ],
            capture_output=True,
            timeout=30,
        )
        if os.path.exists(tmp_png) and os.path.getsize(tmp_png) > 0:
            with open(tmp_png, "rb") as f:
                return f.read()
    except (subprocess.TimeoutExpired, OSError):
        pass
    finally:
        for p in (tmp_html, tmp_png):
            try:
                os.unlink(p)
            except OSError:
                pass
        try:
            import shutil
            shutil.rmtree(tmp_user_data, ignore_errors=True)
        except Exception:
            pass
    return None


def _add_svg_image_to_doc(doc, svg_code: str, width_cm: float = 16) -> bool:
    """
    将 SVG 直接插入到 Word 文档。
    优先转为 PNG 插入（兼容性最好），若无法转换则通过 docx XML 直接嵌入 SVG。
    返回 True 如果成功插入。
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    # 先尝试 PNG 转换
    png_bytes = _svg_to_png_bytes(svg_code)
    if png_bytes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(png_bytes), width=Cm(width_cm))
        return True

    # PNG 转换失败，直接嵌入 SVG 为 docx image part（Word 365+ 支持）
    try:
        from docx.opc.part import Part as OpcPart
        from docx.opc.packuri import PackURI

        svg_bytes = svg_code.encode("utf-8")
        # 提取 viewBox 尺寸来计算比例
        vb_match = re.search(r'viewBox\s*=\s*"([^"]*)"', svg_code)
        if vb_match:
            parts = vb_match.group(1).split()
            if len(parts) == 4:
                vb_w, vb_h = float(parts[2]), float(parts[3])
            else:
                vb_w, vb_h = 1200, 800
        else:
            vb_w, vb_h = 1200, 800

        width_emu = int(width_cm * 360000)  # cm to EMU
        height_emu = int(width_emu * vb_h / max(vb_w, 1))

        # 添加 SVG 作为 image part
        part = doc.part
        svg_part = OpcPart(
            PackURI("/word/media/architecture.svg"),
            "image/svg+xml",
            svg_bytes,
            part.package,
        )
        r_id = part.relate_to(svg_part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")

        # 创建内联图片 XML（不要用 .set() 设置 xmlns 属性，OxmlElement 通过前缀自动处理命名空间）
        inline = OxmlElement("wp:inline")
        inline.set("distT", "0")
        inline.set("distB", "0")
        inline.set("distL", "0")
        inline.set("distR", "0")

        extent = OxmlElement("wp:extent")
        extent.set("cx", str(width_emu))
        extent.set("cy", str(height_emu))
        inline.append(extent)

        docPr = OxmlElement("wp:docPr")
        docPr.set("id", "1")
        docPr.set("name", "Architecture Diagram")
        inline.append(docPr)

        graphic = OxmlElement("a:graphic")

        graphicData = OxmlElement("a:graphicData")
        graphicData.set("uri", "http://schemas.openxmlformats.org/drawingml/2006/picture")

        pic = OxmlElement("pic:pic")

        nvPicPr = OxmlElement("pic:nvPicPr")
        cNvPr = OxmlElement("pic:cNvPr")
        cNvPr.set("id", "0")
        cNvPr.set("name", "architecture.svg")
        nvPicPr.append(cNvPr)
        nvPicPr.append(OxmlElement("pic:cNvPicPr"))
        pic.append(nvPicPr)

        blipFill = OxmlElement("pic:blipFill")
        blip = OxmlElement("a:blip")
        blip.set(_qn("r:embed"), r_id)
        blipFill.append(blip)
        stretch = OxmlElement("a:stretch")
        stretch.append(OxmlElement("a:fillRect"))
        blipFill.append(stretch)
        pic.append(blipFill)

        spPr = OxmlElement("pic:spPr")
        xfrm = OxmlElement("a:xfrm")
        off = OxmlElement("a:off")
        off.set("x", "0")
        off.set("y", "0")
        xfrm.append(off)
        ext = OxmlElement("a:ext")
        ext.set("cx", str(width_emu))
        ext.set("cy", str(height_emu))
        xfrm.append(ext)
        spPr.append(xfrm)
        prstGeom = OxmlElement("a:prstGeom")
        prstGeom.set("prst", "rect")
        spPr.append(prstGeom)
        pic.append(spPr)

        graphicData.append(pic)
        graphic.append(graphicData)
        inline.append(graphic)

        # 将 inline 放到段落 run 的 drawing 元素中
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        drawing = OxmlElement("w:drawing")
        drawing.append(inline)
        run._element.append(drawing)
        return True
    except Exception:
        return False


CSV_SYSTEM_PROMPT = (
    "你是一位 Azure 迁移专家。用户会提供一份 Azure 价格估算表（包含资源名称、SKU、估算金额等）和一份 Azure Migrate 导入 CSV 模板。\n\n"
    "你的任务是：\n"
    "1. 分析价格估算表中的资源列表和金额\n"
    "2. 倒推客户在本地环境可能使用什么配置的 VM\n"
    "3. 按照 CSV 模板格式填充数据\n\n"
    "**CSV 填充规则：**\n"
    "- **必填列**：*Server name, *Cores, *Memory (In MB), *OS name\n"
    "- **Server name 格式**：用描述性命名，格式为 `服务类型-区域-规模-序号`，例如：\n"
    "  TTS-EastAsia-37M-01, LLM-4o-EUS2-01, Search-S1-EastAsia-01, CosmosDB-Serverless-01, Storage-Blob-2TB-01\n"
    "- **Cores/Memory**：根据价格表中 Azure VM/服务规格倒推，常规值：4/8192, 8/16384, 8/32768\n"
    "- **OS name**：大多数写 Linux，数据库类可写 Linux\n"
    "- **磁盘字段也必须填写**：\n"
    "  - Number of disks: 根据服务类型填 1 或 2\n"
    "  - Disk 1 size (In GB): 64/128/256/512/2048 等\n"
    "  - Disk 1 read throughput (MB per second): 根据时期情况填写\n"
    "  - Disk 1 write throughput (MB per second):  根据时期情况填写\n"
    "  - Disk 1 read ops (operations per second): 根据时期情况填写\n"
    "  - Disk 1 write ops (operations per second): 根据时期情况填写\n"
    "  - 如果有第二块盘（数据库、搜索等服务），填写 Disk 2 的相同字段\n"
    "- 其他列留空\n"
    "- 预估 VM 总数量和配置要合理，使得迁移后的 Azure 费用大约等于用户提供的年消耗预算\n\n"
    "**输出格式：**\n"
    "只输出纯粹的 CSV 内容（包含表头行），不要输出 Markdown 代码块标记、解释性文字或其他内容。"
)

# ──────────────────────────────────────────────
# Word 文档生成 —— 通用工具函数
# ──────────────────────────────────────────────
def _set_run_font(run, font_name=CN_FONT, size_pt=None, bold=None, color_rgb=None):
    """为 run 设置字体（含中文 eastAsia 字体）。"""
    run.font.name = font_name
    # python-docx 需要同时设置 eastAsia 字体才能在 Word 中正确显示中文
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb


def _add_styled_paragraph(doc, text, font_name=CN_FONT, size_pt=9, bold=False,
                          color_rgb=None, alignment=None, indent=True):
    """添加一个带完整样式的段落。indent=True 时添加首行缩进。"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    # 首行缩进（约 1 个 Tab = 0.74cm）
    if indent and alignment is None:
        p.paragraph_format.first_line_indent = Cm(0.74)
    # 处理 **加粗** 和普通文字的混合
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        is_bold = bold or (i % 2 == 1)
        _set_run_font(run, font_name=font_name, size_pt=size_pt, bold=is_bold,
                       color_rgb=color_rgb)
    return p


def _add_styled_heading(doc, text, level=1):
    """添加一个使用中文字体的标题。"""
    heading = doc.add_heading("", level=level)
    run = heading.add_run(text)
    size_map = {1: 18, 2: 14, 3: 12}
    _set_run_font(run, font_name=CN_FONT, size_pt=size_map.get(level, 12), bold=True)
    return heading


def _parse_markdown_table(lines: List[str]) -> Optional[List[List[str]]]:
    """
    尝试从 Markdown 行列表中解析表格。
    返回二维数组 (包含表头)，如果不是表格则返回 None。
    """
    if len(lines) < 2:
        return None
    # 检查是否是 Markdown 表格（至少有 | 分隔符和分隔行 ---）
    if "|" not in lines[0]:
        return None

    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # 跳过分隔行 |---|---|
        if re.match(r"^\|[\s\-:|]+\|$", stripped):
            continue
        # 解析单元格
        cells = [c.strip() for c in stripped.split("|")]
        # 去掉首尾空元素（因为 | 在开头和结尾会产生空字符串）
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows if len(rows) >= 2 else None


def _add_word_table(doc, table_data: list[list[str]]):
    """将二维数组写入 Word 表格，应用专业样式。"""
    if not table_data:
        return

    num_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = "Table Grid"

    for ri, row_data in enumerate(table_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""  # 清空默认段落文本
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            is_header = (ri == 0)
            _set_run_font(
                run,
                font_name=CN_FONT,
                size_pt=9,
                bold=is_header,
            )
            # 表头行背景色
            if is_header:
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(
                    qn("w:shd"),
                    {qn("w:fill"): "156082", qn("w:val"): "clear"},
                )
                shading.append(shading_elem)
                run.font.color.rgb = RGBColor(255, 255, 255)


def _markdown_to_docx(doc, markdown_text: str, body_size=9):
    """
    将 AI 返回的 Markdown 文本解析并写入 Word 文档。
    支持: 标题 (#/##/###)、列表 (-/*)、Markdown 表格、加粗 (**)、普通段落。
    """
    # 预处理：将 <br> 变体转换为换行符
    markdown_text = re.sub(r"<br\s*/?>", "\n", markdown_text)
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # ── 跳过 --- 分隔线 ──
        if stripped == '---' or stripped == '***' or stripped == '___':
            i += 1
            continue

        # ── 标题 ──
        if stripped.startswith("#### "):
            _add_styled_heading(doc, stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith("### "):
            _add_styled_heading(doc, stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            _add_styled_heading(doc, stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            _add_styled_heading(doc, stripped[2:], level=1)
            i += 1
            continue

        # ── 独立的 **加粗行**（如阶段标题），转为三级标题 ──
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            title_text = stripped[2:-2]
            _add_styled_heading(doc, title_text, level=3)
            i += 1
            continue

        # ── Markdown 表格 ──
        if "|" in stripped and not stripped.startswith("-"):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            table_data = _parse_markdown_table(table_lines)
            if table_data:
                _add_word_table(doc, table_data)
                doc.add_paragraph()  # 表格后空行
            else:
                # 不是表格，作为普通文本处理
                for tl in table_lines:
                    _add_styled_paragraph(doc, tl.strip(), size_pt=body_size)
            continue

        # ── 无序列表 ──
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            _add_styled_paragraph(doc, f"•  {text}", size_pt=body_size)
            i += 1
            continue

        # ── 有序列表 ──
        if stripped[0].isdigit() and ". " in stripped[:5]:
            _add_styled_paragraph(doc, stripped, size_pt=body_size)
            i += 1
            continue

        # ── 普通段落 ──
        _add_styled_paragraph(doc, stripped, size_pt=body_size)
        i += 1


# ──────────────────────────────────────────────
# Word 文档生成 —— 基于模板
# ──────────────────────────────────────────────
def _load_template(template_path: str) -> Document:
    """
    加载 .docx 模板文件作为基础文档。
    如果模板不存在，则返回一个空白 Document。
    """
    if os.path.exists(template_path):
        doc = Document(template_path)
        # 清空模板中的所有正文段落（保留样式定义、页面设置、页眉页脚）
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        # 清空表格
        for t in doc.tables:
            t._element.getparent().remove(t._element)
        return doc
    else:
        return Document()


def _extract_title(content: str, fallback: str = "") -> str:
    """从 AI 生成的 Markdown 内容中提取第一个 # 标题作为文档标题。"""
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            return stripped[2:].strip()
    return fallback


def _strip_first_heading(content: str) -> str:
    """去掉 Markdown 内容中的第一个 # 标题行（因为封面已经显示了标题）。"""
    lines = content.split("\n")
    result = []
    found = False
    for line in lines:
        stripped = line.strip()
        if not found and stripped.startswith("# ") and not stripped.startswith("## "):
            found = True
            continue  # 跳过第一个 # 标题
        result.append(line)
    return "\n".join(result)


def _add_page_break(doc):
    """在文档中添加分页符。"""
    from docx.oxml.ns import qn as _qn
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(_qn("w:br"), {_qn("w:type"): "page"})
    run._element.append(br)


def _add_toc(doc):
    """插入 Word 目录域（用户打开文档后按 Ctrl+A → F9 即可更新）。"""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    # 目录标题
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目录")
    _set_run_font(run, font_name=CN_FONT, size_pt=16, bold=True)

    doc.add_paragraph()  # 空行

    # 插入 TOC 域代码
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(_qn("w:fldCharType"), "begin")
    run._element.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(_qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(_qn("w:fldCharType"), "separate")
    run._element.append(fldChar_separate)

    # 占位文本（打开 Word 后会自动替换）
    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "（请右键点击此处 → 更新域，生成目录）"
    placeholder.append(placeholder_text)
    run._element.append(placeholder)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(_qn("w:fldCharType"), "end")
    run._element.append(fldChar_end)


def create_solution_docx(content: str, customer_name: str, svg_code: Optional[str] = None) -> bytes:
    """
    基于 solution 模板生成解决方案架构 Word 文档。
    布局: 封面标题（独占一页） → 目录（独占一页） → 正文
    如果提供 svg_code，则在第二章节结束后插入架构图。
    """
    doc = _load_template(SOLUTION_TEMPLATE_PATH)
    title = _extract_title(content, f"{customer_name} - AI 解决方案架构文档")
    body_content = _strip_first_heading(content)

    # ---- 第 1 页：封面标题 ----
    for _ in range(8):
        doc.add_paragraph()

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    # 与模板一致: 18pt #4874CB
    _set_run_font(run, font_name=CN_FONT_ALT, size_pt=18,
                   bold=True, color_rgb=RGBColor(0x48, 0x74, 0xCB))

    # 封面分页
    _add_page_break(doc)

    # ---- 第 2 页：目录 ----
    _add_toc(doc)

    # 目录分页
    _add_page_break(doc)

    # ---- 第 3 页起：正文内容（已去掉第一个 # 标题） ----
    if svg_code:
        # 在第二章节后插入架构图
        # 查找第三个 ## 标题（即第三章开头），在其前面插入图片
        lines = body_content.split("\n")
        h2_count = 0
        split_idx = len(lines)
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith("## ") or (stripped.startswith("# ") and not stripped.startswith("## ")):
                h2_count += 1
                if h2_count == 3:  # 第三章开头
                    split_idx = idx
                    break

        # 渲染第二章及之前的内容
        part1 = "\n".join(lines[:split_idx])
        _markdown_to_docx(doc, part1, body_size=9)

        # 插入架构图
        doc.add_paragraph()  # 空行
        _add_svg_image_to_doc(doc, svg_code, width_cm=16)

        # 图注
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"图：{customer_name} Azure AI 解决方案架构图")
        _set_run_font(cap_run, font_name=CN_FONT, size_pt=8, bold=False)

        doc.add_paragraph()  # 空行

        # 渲染剩余内容
        part2 = "\n".join(lines[split_idx:])
        if part2.strip():
            _markdown_to_docx(doc, part2, body_size=9)
    else:
        _markdown_to_docx(doc, body_content, body_size=9)

    # 导出
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_pov_docx(content: str, customer_name: str) -> bytes:
    """
    基于 POV 模板生成 POV 部署计划 Word 文档。
    布局: 封面标题（独占一页） → 正文
    """
    doc = _load_template(POV_TEMPLATE_PATH)
    title = _extract_title(content, f"{customer_name} - POV 部署计划")
    body_content = _strip_first_heading(content)

    # ---- 第 1 页：封面标题 ----
    for _ in range(8):
        doc.add_paragraph()

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    # 与模板一致: 22pt #156082
    _set_run_font(run, font_name=CN_FONT_ALT, size_pt=22,
                   bold=True, color_rgb=RGBColor(0x15, 0x60, 0x82))

    # 封面分页
    _add_page_break(doc)

    # ---- 第 2 页起：正文内容（已去掉第一个 # 标题） ----
    _markdown_to_docx(doc, body_content, body_size=9)

    # 导出
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_infra_docx(content: str, customer_name: str) -> bytes:
    """
    基于 Infra 模板生成基础设施解决方案 Word 文档。
    布局: 封面标题（独占一页） → 目录（独占一页） → 正文
    """
    doc = _load_template(INFRA_TEMPLATE_PATH)
    title = _extract_title(content, f"{customer_name} - 基础设施解决方案架构文档")
    body_content = _strip_first_heading(content)

    # ---- 第 1 页：封面标题 ----
    for _ in range(8):
        doc.add_paragraph()

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    # 与 AI 解决方案一致: 18pt #4874CB
    _set_run_font(run, font_name=CN_FONT_ALT, size_pt=18,
                   bold=True, color_rgb=RGBColor(0x48, 0x74, 0xCB))

    # 封面分页
    _add_page_break(doc)

    # ---- 第 2 页：目录 ----
    _add_toc(doc)

    # 目录分页
    _add_page_break(doc)

    # ---- 第 3 页起：正文内容（已去掉第一个 # 标题） ----
    _markdown_to_docx(doc, body_content, body_size=9)

    # 导出
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ──────────────────────────────────────────────
# 辅助：日期前缀文件名
# ──────────────────────────────────────────────
def _date_prefix():
    """返回当前日期前缀，如 0225"""
    return datetime.date.today().strftime("%m%d")


# ──────────────────────────────────────────────
# 全自动 POE：MSAL 登录与 Azure ARM API
# ──────────────────────────────────────────────
def _get_msal_client_id() -> str:
    """返回 MSAL Public Client ID。Client ID 不是密钥，可使用默认值。"""
    return get_secret("MSAL_CLIENT_ID", MSAL_CLIENT_ID_DEFAULT)


def _is_azure_token_valid() -> bool:
    expires_at = st.session_state.get("azure_token_expires_at", 0)
    return bool(st.session_state.get("azure_token")) and time.time() < expires_at


def msal_device_code_login() -> None:
    """通过 MSAL Device Code Flow 登录 Azure，并将 access token 放入 session state。"""
    if msal is None:
        raise RuntimeError("缺少 msal 依赖，请确认 requirements.txt 已包含 msal。")

    app = msal.PublicClientApplication(
        _get_msal_client_id(),
        authority=AZURE_AUTHORITY,
    )
    flow = app.initiate_device_flow(scopes=AZURE_ARM_SCOPE)
    if "user_code" not in flow:
        raise RuntimeError(f"无法启动 Microsoft 登录流程：{flow}")

    user_code = flow["user_code"]
    verify_url = flow.get("verification_uri", "https://microsoft.com/devicelogin")
    render_device_code_login(user_code, verify_url)

    result = app.acquire_token_by_device_flow(flow)
    if "access_token" not in result:
        err = result.get("error_description") or result.get("error") or "Microsoft 登录失败。"
        if "7000218" in str(err):
            err += "\n\n请在 Azure Portal → 应用注册 → 身份验证 → 高级设置中，将「允许公共客户端流」设为「是」。"
        raise RuntimeError(err)

    account = result.get("id_token_claims", {}) or result.get("account", {}) or {}
    username = account.get("preferred_username") or account.get("email") or account.get("username") or "Azure 用户"
    expires_in = int(result.get("expires_in", 3600))
    st.session_state["azure_token"] = result["access_token"]
    st.session_state["azure_user"] = username
    st.session_state["azure_token_expires_at"] = time.time() + max(expires_in - 300, 300)
    persist_session_state()


def clear_azure_login() -> None:
    for key in [
        "azure_token",
        "azure_user",
        "azure_token_expires_at",
        "azure_subscription_id",
        "azure_subscription_name",
        "azure_resource_group",
        "_cached_subscription",
        "_cached_resource_group",
    ]:
        st.session_state.pop(key, None)
    clear_session_persist()
def _format_arm_error(response: requests.Response) -> str:
    try:
        payload = response.json()
        error = payload.get("error", payload)
        message = error.get("message") or str(error)
    except Exception:
        message = response.text
    return f"Azure API {response.status_code}: {message[:1200]}"


def _retry_after_seconds(response: requests.Response, default: int = 10) -> int:
    retry_after = response.headers.get("Retry-After")
    if not retry_after:
        return default
    try:
        return max(1, min(int(retry_after), 60))
    except ValueError:
        return default


def _poll_azure_lro(
    operation_url: str,
    token: str,
    initial_delay: int = 10,
    timeout_seconds: int = 900,
) -> Dict[str, Any]:
    """轮询 ARM long-running operation，直到 Succeeded/Failed/Canceled。"""
    if operation_url.startswith("/"):
        operation_url = f"{AZURE_MANAGEMENT_ENDPOINT}{operation_url}"

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    deadline = time.time() + timeout_seconds
    delay = max(1, min(initial_delay, 60))

    while time.time() < deadline:
        time.sleep(delay)
        response = requests.get(operation_url, headers=headers, timeout=90)
        if response.status_code >= 400:
            raise RuntimeError(_format_arm_error(response))

        payload: Dict[str, Any] = {}
        if response.content:
            try:
                payload = response.json()
            except ValueError:
                payload = {}

        status = str(
            payload.get("status")
            or payload.get("properties", {}).get("provisioningState")
            or ""
        ).strip()
        status_lower = status.lower()
        if status_lower in {"succeeded", "completed"}:
            return payload
        if status_lower in {"failed", "canceled", "cancelled"}:
            raise RuntimeError(f"Azure 长操作失败：{payload or status}")
        if response.status_code in {200, 204} and not status:
            return payload

        delay = _retry_after_seconds(response, default=10)

    raise TimeoutError("等待 Azure 长操作完成超时。")


def azure_arm_request(
    method: str,
    path_or_url: str,
    token: str,
    body: Optional[Dict[str, Any]] = None,
    timeout: int = 90,
    poll_lro: bool = True,
    lro_timeout: int = 900,
    max_retries: int = 3,
) -> Dict[str, Any]:
    """调用 Azure ARM REST API。path_or_url 可传完整 URL 或 ARM 相对路径。"""
    url = path_or_url if path_or_url.startswith("http") else f"{AZURE_MANAGEMENT_ENDPOINT}{path_or_url}"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    last_exc: Optional[Exception] = None
    for attempt in range(max_retries):
        try:
            response = requests.request(method, url, headers=headers, json=body, timeout=timeout)
            break
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as exc:
            last_exc = exc
            if attempt < max_retries - 1:
                time.sleep(5 * (attempt + 1))
            else:
                raise RuntimeError(
                    f"HTTPSConnectionPool(host='management.azure.com', port=443): "
                    f"Max retries exceeded with url: {path_or_url.split('?')[0]} — {exc}"
                ) from exc
    if response.status_code >= 400:
        raise RuntimeError(_format_arm_error(response))
    payload: Dict[str, Any] = {}
    if response.content:
        try:
            payload = response.json()
        except ValueError:
            payload = {}

    lro_url = response.headers.get("Azure-AsyncOperation") or response.headers.get("Location")
    if poll_lro and response.status_code in {201, 202} and lro_url:
        final_payload = _poll_azure_lro(
            lro_url,
            token,
            initial_delay=_retry_after_seconds(response, default=10),
            timeout_seconds=lro_timeout,
        )
        return payload or final_payload

    return payload


def azure_arm_list(path: str, token: str) -> List[Dict[str, Any]]:
    """读取 ARM 列表接口并处理 nextLink 分页。"""
    items: List[Dict[str, Any]] = []
    next_url = f"{AZURE_MANAGEMENT_ENDPOINT}{path}"
    while next_url:
        payload = azure_arm_request("GET", next_url, token)
        items.extend(payload.get("value", []))
        next_url = payload.get("nextLink")
    return items


def list_azure_subscriptions(token: str) -> List[Dict[str, Any]]:
    subscriptions = azure_arm_list(f"/subscriptions?api-version={AZURE_RESOURCE_API_VERSION}", token)
    return sorted(subscriptions, key=lambda item: item.get("displayName", ""))


def list_azure_resource_groups(subscription_id: str, token: str) -> List[Dict[str, Any]]:
    groups = azure_arm_list(
        f"/subscriptions/{subscription_id}/resourceGroups?api-version={AZURE_RESOURCE_API_VERSION}",
        token,
    )
    return sorted(groups, key=lambda item: item.get("name", ""))


def _subscription_label(subscription: Dict[str, Any]) -> str:
    display_name = subscription.get("displayName") or subscription.get("subscriptionId")
    state = subscription.get("state", "Unknown")
    return f"{display_name} ({state})"


def _resource_group_label(resource_group: Dict[str, Any]) -> str:
    name = resource_group.get("name", "")
    location = resource_group.get("location", "")
    return f"{name} ({location})" if location else name


# ──────────────────────────────────────────────
# 全自动 POE：文档、Migrate、打包
# ──────────────────────────────────────────────
def _safe_azure_name(value: str, fallback: str, suffix: str = "", max_len: int = 60) -> str:
    base = re.sub(r"[^A-Za-z0-9-]+", "-", value or "").strip("-")
    base = re.sub(r"-+", "-", base)
    if not base:
        base = fallback
    reserve = len(suffix)
    return f"{base[: max_len - reserve].strip('-')}{suffix}".strip("-")


AZURE_MIGRATE_PROJECT_LOCATIONS = {
    "centralus", "westeurope", "uksouth", "ukwest", "northeurope", "westus2",
    "southeastasia", "eastasia", "centralindia", "southindia", "canadacentral",
    "australiasoutheast", "japanwest", "japaneast", "brazilsouth", "koreacentral",
    "koreasouth", "francecentral", "switzerlandnorth", "australiaeast", "uaenorth",
    "southafricanorth", "germanywestcentral", "norwayeast", "jioindiawest",
    "swedencentral", "qatarcentral", "polandcentral", "italynorth", "israelcentral",
    "spaincentral", "mexicocentral", "newzealandnorth", "indonesiacentral",
    "malaysiawest", "chilecentral", "austriaeast", "belgiumcentral", "denmarkeast",
}

AZURE_MIGRATE_LOCATION_ALIASES = {
    "westus": "westus2",
    "west us": "westus2",
    "east us 2": "eastus2",
    "eastus 2": "eastus2",
}


def _normalize_azure_location(location: str) -> str:
    return re.sub(r"\s+", "", (location or "").strip().lower())


def resolve_migrate_project_location(subscription_id: str, resource_group: str, token: str) -> str:
    resource_group_payload = azure_arm_request(
        "GET",
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}?api-version={AZURE_RESOURCE_API_VERSION}",
        token,
    )
    normalized = _normalize_azure_location(resource_group_payload.get("location", ""))
    normalized = AZURE_MIGRATE_LOCATION_ALIASES.get(normalized, normalized)
    if normalized in AZURE_MIGRATE_PROJECT_LOCATIONS:
        return normalized
    return "westus2"


def _workday_info(start_date: datetime.date, end_date: datetime.date) -> tuple[list[str], list[str]]:
    workdays = []
    weekends = []
    current = start_date
    while current <= end_date:
        target = f"{current.month}月{current.day}日"
        if current.weekday() < 5:
            workdays.append(target)
        else:
            weekends.append(target)
        current += datetime.timedelta(days=1)
    return workdays, weekends


def has_meaningful_pov_team(vendor_team: Optional[str]) -> bool:
    text = str(vendor_team or "").strip()
    if not text:
        return False
    placeholders = {"技术负责人", "Azure架构师", "Azure 架构师", "项目经理", "负责人"}
    for raw_line in text.splitlines():
        line = raw_line.strip().strip("：:")
        if not line:
            continue
        if ":" in raw_line or "：" in raw_line:
            _, value = re.split(r"[:：]", raw_line, maxsplit=1)
            if value.strip():
                return True
            if line in placeholders:
                continue
        elif line not in placeholders:
            return True
    return False


def build_pov_prompt(
    solution_text: str,
    customer_name: str,
    pov_start: datetime.date,
    pov_end: datetime.date,
    vendor_team: str,
    pov_ref: str,
) -> str:
    workdays, weekends = _workday_info(pov_start, pov_end)
    pov_prompt = (
        f"以下是已生成的解决方案架构文档，请据此生成 POV 部署计划：\n\n"
        f"{solution_text}\n\n"
        f"## 补充信息\n- **客户名称**：{customer_name}\n"
        f"- **POV 周期**：{pov_start.strftime('%Y/%m/%d')} - {pov_end.strftime('%Y/%m/%d')}\n\n"
        f"## 可用工作日清单（共 {len(workdays)} 天，必须且只能使用这些日期）\n"
        f"{'、'.join(workdays)}\n\n"
        f"## 禁用日期（周末，严禁安排任何任务）\n"
        f"{'、'.join(weekends) if weekends else '无'}\n\n"
        f"## 乙方项目人员\n{vendor_team.strip()}\n\n"
        f"请根据客户背景信息自动生成合理的甲方人员（2-3人，包含项目负责人和技术对接人，一定要中文名！）。"
    )
    if pov_ref:
        pov_prompt += (
            "\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
            f"{pov_ref}"
        )
    return pov_prompt


def generate_solution_artifact(
    current_doc_type: str,
    customer_name: str,
    account_name: str,
    customer_bg: str,
    solution_ref: str,
    infra_ref: str,
) -> Dict[str, Any]:
    system_prompt = SOLUTION_SYSTEM_PROMPT if current_doc_type == "AI" else INFRA_SYSTEM_PROMPT
    ref_text = solution_ref if current_doc_type == "AI" else infra_ref
    user_ctx = (
        f"## 客户信息\n- **客户名称**：{customer_name}\n\n"
        f"## 客户背景\n{customer_bg}"
    )
    if ref_text:
        user_ctx += (
            "\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
            f"{ref_text}"
        )

    content = call_azure_openai(system_prompt, user_ctx)
    if current_doc_type == "AI":
        # 生成 SVG 架构图
        svg_code = generate_svg_architecture(content, customer_name)
        docx_bytes = create_solution_docx(content=content, customer_name=customer_name, svg_code=svg_code)
        file_name = f"{account_name}-Solution Architecture.docx"
    else:
        docx_bytes = create_infra_docx(content=content, customer_name=customer_name)
        file_name = f"{account_name}-Infra Solution Architecture.docx"

    result = {"content": content, "bytes": docx_bytes, "file_name": file_name}
    if current_doc_type == "AI" and svg_code:
        result["svg_code"] = svg_code
    return result


def generate_pov_artifact(
    solution_text: str,
    customer_name: str,
    account_name: str,
    pov_ref: str,
    pov_start: datetime.date,
    pov_end: datetime.date,
    vendor_team: str,
) -> Dict[str, Any]:
    pov_prompt = build_pov_prompt(solution_text, customer_name, pov_start, pov_end, vendor_team, pov_ref)
    content = call_azure_openai(POV_SYSTEM_PROMPT, pov_prompt)
    docx_bytes = create_pov_docx(content=content, customer_name=customer_name)
    return {
        "content": content,
        "bytes": docx_bytes,
        "file_name": f"{account_name}-PostAssessment POVdeployment.docx",
    }


def register_azure_provider(subscription_id: str, namespace: str, token: str) -> None:
    azure_arm_request(
        "POST",
        f"/subscriptions/{subscription_id}/providers/{namespace}/register?api-version={AZURE_PROVIDER_API_VERSION}",
        token,
    )


def _extract_sas_url(payload: Any) -> Optional[str]:
    if isinstance(payload, str):
        if payload.startswith("http") and "sig=" in payload:
            return payload
        return None
    if isinstance(payload, dict):
        for value in payload.values():
            found = _extract_sas_url(value)
            if found:
                return found
    if isinstance(payload, list):
        for value in payload:
            found = _extract_sas_url(value)
            if found:
                return found
    return None


def _arm_path_with_api_version(path_or_id: str, api_version: str) -> str:
    separator = "&" if "?" in path_or_id else "?"
    return f"{path_or_id}{separator}api-version={api_version}"


def _resource_id(subscription_id: str, resource_group: str, provider_path: str) -> str:
    return f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}/{provider_path}"


def _migrate_project_path(subscription_id: str, resource_group: str, project_name: str) -> str:
    return (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/migrateProjects/{project_name}"
        f"?api-version={AZURE_MIGRATE_PROJECTS_API_VERSION}"
    )


def _migrate_solution_path(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    solution_name: str,
) -> str:
    return (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/migrateProjects/{project_name}"
        f"/solutions/{solution_name}?api-version={AZURE_MIGRATE_PROJECTS_API_VERSION}"
    )


def _migrate_solution_id(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    solution_name: str,
) -> str:
    return _resource_id(
        subscription_id,
        resource_group,
        f"providers/Microsoft.Migrate/migrateProjects/{project_name}/solutions/{solution_name}",
    )


def _servers_solution_details(extended_details: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    details = {
        "dependencyEnabledMachines": "0",
        "machinesHavingSqlServers": "0",
        "machinesHavingWebServers": "0",
        "serversOnLinux": "0",
        "serversOnWindows": "0",
        "serversOnOther": "0",
    }
    if extended_details:
        details.update(extended_details)
    return {
        "assessmentCount": 0,
        "groupCount": 0,
        "extendedDetails": details,
    }


def register_migrate_tool(subscription_id: str, resource_group: str, project_name: str, tool: str, token: str) -> None:
    azure_arm_request(
        "POST",
        (
            f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
            f"/providers/Microsoft.Migrate/migrateProjects/{project_name}"
            f"/registerTool?api-version={AZURE_MIGRATE_PROJECTS_API_VERSION}"
        ),
        token,
        {"tool": tool},
    )


def put_migrate_solution(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    solution_name: str,
    properties: Dict[str, Any],
    token: str,
) -> Dict[str, Any]:
    return azure_arm_request(
        "PUT",
        _migrate_solution_path(subscription_id, resource_group, project_name, solution_name),
        token,
        {"properties": properties},
    )


def ensure_migrate_solution(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    solution_name: str,
    properties: Dict[str, Any],
    token: str,
    progress: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    path = _migrate_solution_path(subscription_id, resource_group, project_name, solution_name)
    existing = _try_get_existing_resource(path, token)
    if existing and existing.get("properties", {}).get("details"):
        if progress:
            progress(f"  ✅ Solution 已存在且 details 完整，复用: {solution_name}")
        return existing
    result = put_migrate_solution(subscription_id, resource_group, project_name, solution_name, properties, token)
    if progress:
        progress(f"  ✅ Solution 已补齐: {solution_name}")
    return result


def ensure_portal_menu_solutions(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    master_site_id: str,
    token: str,
    progress: Callable[[str], None],
) -> None:
    """补齐 Azure Portal 项目菜单 blade 会枚举的默认 solution，避免前端读取 undefined.details。"""
    default_solutions = [
        (
            "Servers-Discovery-ServerDiscovery",
            {
                "tool": "ServerDiscovery",
                "purpose": "Discovery",
                "goal": "Servers",
                "status": "Inactive",
                "details": _servers_solution_details({"masterSiteId": master_site_id}),
            },
        ),
        (
            "Servers-Migration-ServerMigration",
            {
                "tool": "ServerMigration",
                "purpose": "Migration",
                "goal": "Servers",
                "status": "Active",
                "details": _servers_solution_details(),
            },
        ),
        (
            "Servers-Migration-ServerMigration_DataReplication",
            {
                "tool": "ServerMigration_DataReplication",
                "purpose": "Migration",
                "goal": "Servers",
                "status": "Inactive",
                "details": _servers_solution_details(),
            },
        ),
    ]
    for solution_name, properties in default_solutions:
        ensure_migrate_solution(
            subscription_id,
            resource_group,
            project_name,
            solution_name,
            properties,
            token,
            progress,
        )


def refresh_migrate_project_summary(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    token: str,
) -> None:
    azure_arm_request(
        "POST",
        (
            f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
            f"/providers/Microsoft.Migrate/migrateProjects/{project_name}"
            f"/refreshSummary?api-version={AZURE_MIGRATE_PROJECTS_API_VERSION}"
        ),
        token,
        {"goal": "Servers"},
        poll_lro=False,
    )


def _server_summary_count(migrate_project: Dict[str, Any]) -> int:
    servers = migrate_project.get("properties", {}).get("summary", {}).get("servers", {})
    direct_count = int(servers.get("discoveredCount") or 0)
    extended = servers.get("extendedSummary") or {}
    microsoft_count = int(extended.get("microsoftMachinesCount") or 0)
    return max(direct_count, microsoft_count)


def wait_for_portal_inventory_summary(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    expected_machine_count: int,
    token: str,
    progress: Callable[[str], None],
    timeout_seconds: int = 300,
) -> int:
    """等待 migrateProject summary 反映 Import CSV 库存，这是 Portal 全部库存 blade 使用的项目汇总层。"""
    project_path = _migrate_project_path(subscription_id, resource_group, project_name)
    deadline = time.time() + timeout_seconds
    attempt = 0
    last_count = 0

    while time.time() < deadline:
        attempt += 1
        try:
            refresh_migrate_project_summary(subscription_id, resource_group, project_name, token)
        except Exception:
            pass

        project = azure_arm_request("GET", project_path, token, poll_lro=False)
        last_count = _server_summary_count(project)
        if last_count >= expected_machine_count:
            return last_count

        progress(
            f"等待 Azure Portal 全部库存汇总刷新，当前 {last_count}/{expected_machine_count} 台"
            f"（第 {attempt} 次检查）"
        )
        time.sleep(20)

    raise TimeoutError(
        "Azure Portal 全部库存汇总未刷新到预期数量。"
        f"当前 {last_count}/{expected_machine_count} 台；"
        "请检查 ServerDiscovery_Import solution 与 Import Site 关联。"
    )


def _format_import_job_error(job: Dict[str, Any]) -> str:
    props = job.get("properties", {}) if isinstance(job, dict) else {}
    summary = props.get("errorSummary") if isinstance(props, dict) else {}
    parts = []
    if isinstance(summary, dict):
        error_count = summary.get("errorCount")
        warning_count = summary.get("warningCount")
        if error_count is not None:
            parts.append(f"errors={error_count}")
        if warning_count is not None:
            parts.append(f"warnings={warning_count}")
        errors = summary.get("errors")
        if isinstance(errors, list) and errors:
            preview = "; ".join(str(item) for item in errors[:5])
            parts.append(f"details={preview}")
    result = props.get("jobResult") or job.get("status")
    if result:
        parts.insert(0, f"jobResult={result}")
    return "；".join(parts) if parts else str(job)[:800]


def _dedupe_machines_by_discovery_arm_id(machines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    deduped: List[Dict[str, Any]] = []
    seen = set()
    for machine in machines:
        props = machine.get("properties", {})
        key = str(props.get("discoveryMachineArmId") or machine.get("id") or machine.get("name") or "").lower()
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(machine)
    return deduped


def wait_for_import_site_import(
    subscription_id: str,
    resource_group: str,
    site_name: str,
    token: str,
    progress: Callable[[str], None],
    job_arm_id: Optional[str] = None,
    timeout_seconds: int = 900,
) -> List[Dict[str, Any]]:
    """等待 OffAzure import site 完成 CSV 解析，并返回导入到 import site 的机器。"""
    machines_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.OffAzure/importSites/{site_name}/machines"
        f"?api-version={AZURE_OFFAZURE_API_VERSION}"
    )
    jobs_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.OffAzure/importSites/{site_name}/importJobs"
        f"?api-version={AZURE_OFFAZURE_API_VERSION}"
    )
    job_path = (
        _arm_path_with_api_version(job_arm_id, AZURE_OFFAZURE_API_VERSION)
        if job_arm_id and (job_arm_id.startswith("/") or job_arm_id.startswith("http"))
        else None
    )
    deadline = time.time() + timeout_seconds
    attempt = 0
    last_job: Dict[str, Any] = {}

    while time.time() < deadline:
        attempt += 1
        try:
            machines = azure_arm_list(machines_path, token)
            if machines:
                return machines
        except Exception:
            pass

        job: Dict[str, Any] = {}
        if job_path:
            try:
                job = azure_arm_request("GET", job_path, token, poll_lro=False)
            except Exception:
                job = {}
        if not job:
            try:
                jobs = azure_arm_list(jobs_path, token)
                if jobs:
                    job = jobs[-1]
            except Exception:
                job = {}
        if job:
            last_job = job
            props = job.get("properties", {})
            result = str(props.get("jobResult") or job.get("status") or "Unknown").strip()
            imported_count = props.get("numberOfMachinesImported")
            if result in {"Completed", "CompletedWithWarnings"}:
                machines = azure_arm_list(machines_path, token)
                if machines:
                    return machines
            if result in {"Failed", "CompletedWithErrors"}:
                raise RuntimeError(f"Azure Migrate CSV 导入失败：{_format_import_job_error(job)}")
            suffix = f"，已导入 {imported_count} 台" if imported_count is not None else ""
            progress(f"服务器清单导入任务状态：{result}{suffix}（第 {attempt} 次检查）")
        else:
            progress(f"等待 Azure Migrate 创建服务器清单导入任务...（第 {attempt} 次检查）")
        time.sleep(15)

    detail = f"最后一次任务状态：{_format_import_job_error(last_job)}" if last_job else "未查询到导入任务。"
    raise TimeoutError(f"等待 Azure Migrate 导入服务器清单超时（15 分钟）。{detail}")


def wait_for_project_machines(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    collector_name: str,
    token: str,
    progress: Callable[[str], None],
    site_name: Optional[str] = None,
    timeout_seconds: int = 600,
) -> List[Dict[str, Any]]:
    """等待 Azure Migrate 导入完成，通过直接查询 machines 列表来判断。"""
    machines_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/machines?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    project_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    deadline = time.time() + timeout_seconds
    attempt = 0

    def _filter_current_import(machines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        if not site_name:
            return _dedupe_machines_by_discovery_arm_id(machines)
        marker = f"/importsites/{site_name}".lower()
        return _dedupe_machines_by_discovery_arm_id([
            machine for machine in machines
            if marker in str(machine.get("properties", {}).get("discoveryMachineArmId", "")).lower()
        ])

    while time.time() < deadline:
        attempt += 1
        # 先尝试直接列出 machines
        try:
            all_machines = azure_arm_list(machines_path, token)
            machines = _filter_current_import(all_machines)
            if machines:
                return machines
        except Exception:
            pass
        # 也检查 project 的 numberOfMachines
        try:
            project = azure_arm_request("GET", project_path, token)
            machine_count = project.get("properties", {}).get("numberOfMachines", 0) or 0
            if machine_count > 0:
                machines = _filter_current_import(azure_arm_list(machines_path, token))
                if machines:
                    return machines
        except Exception:
            pass
        progress(f"服务器清单仍在导入中，继续等待 Azure Migrate 发现结果...（第 {attempt} 次检查）")
        time.sleep(15)
    raise TimeoutError("等待 Azure Migrate 导入服务器清单超时（10 分钟），请在 Azure Portal 检查导入状态。")


def list_imported_machines(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    collector_name: str,
    token: str,
) -> List[Dict[str, Any]]:
    candidate_paths = [
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}/providers/Microsoft.Migrate/assessmentProjects/{project_name}/machines?api-version={AZURE_MIGRATE_API_VERSION}",
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}/providers/Microsoft.Migrate/assessmentProjects/{project_name}/importcollectors/{collector_name}/machines?api-version={AZURE_MIGRATE_API_VERSION}",
    ]
    last_error = None
    for path in candidate_paths:
        try:
            machines = azure_arm_list(path, token)
            if machines:
                return machines
        except Exception as exc:
            last_error = exc
    if last_error:
        raise RuntimeError(f"无法读取 Azure Migrate 已导入服务器列表：{last_error}")
    return []


def wait_for_assessment_complete(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    group_name: str,
    assessment_name: str,
    token: str,
    progress: Callable[[str], None],
    timeout_seconds: int = 600,
) -> Dict[str, Any]:
    path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/groups/{group_name}/assessments/{assessment_name}"
        f"?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        assessment = azure_arm_request("GET", path, token)
        props = assessment.get("properties", {})
        status = props.get("status", "Unknown")
        stage = props.get("stage", "Unknown")
        provisioning_state = props.get("provisioningState", "Unknown")
        if status == "Completed":
            return assessment
        if status in {"Invalid", "OutOfSync", "OutDated", "Deleted"}:
            raise RuntimeError(f"Azure Migrate 评估状态异常：{status}")
        if provisioning_state in {"Failed", "Canceled"}:
            raise RuntimeError(f"Azure Migrate 评估资源状态异常：{provisioning_state}")
        progress(f"评估仍在计算中，当前状态：{status}，阶段：{stage}，资源状态：{provisioning_state}")
        time.sleep(15)
    raise TimeoutError("等待 Azure Migrate 评估完成超时，请稍后在 Azure Portal 检查评估结果。")


def parse_annual_budget_usd(raw_budget: Optional[str]) -> Optional[float]:
    """把 UI 里的年预估消耗解析成 USD 数字；无法可靠解析时返回 None。"""
    if raw_budget is None:
        return None
    text = str(raw_budget).strip()
    if not text:
        return None

    normalized = (
        text.lower()
        .replace(",", "")
        .replace("$", "")
        .replace("usd", "")
        .replace("美元", "")
        .replace("美金", "")
        .replace("预估", "")
        .replace("年消耗", "")
        .replace("年度", "")
        .replace("每年", "")
        .replace("+", "")
        .strip()
    )
    if not normalized or normalized in {"na", "n/a", "none", "null", "未填写", "无"}:
        return None

    multiplier = 1.0
    if "万" in normalized or re.search(r"\d\s*w\b", normalized):
        multiplier = 10_000.0
        normalized = normalized.replace("万", "").replace("w", "")
    elif "million" in normalized:
        multiplier = 1_000_000.0
        normalized = normalized.replace("million", "")
    elif re.search(r"\d\s*m\b", normalized):
        multiplier = 1_000_000.0
        normalized = re.sub(r"\bm\b", "", normalized)
    elif re.search(r"\d\s*k\b", normalized):
        multiplier = 1_000.0
        normalized = re.sub(r"\bk\b", "", normalized)

    match = re.search(r"(\d+(?:\.\d+)?)", normalized)
    if not match:
        return None
    return float(match.group(1)) * multiplier


# ──────────────────────────────────────────────
# 内置 CSV 模板 + 分规模机器选择学习
# ──────────────────────────────────────────────

def load_builtin_csv_template() -> str:
    with open(BUILTIN_CSV_PATH, "r", encoding="utf-8-sig") as f:
        return f.read()


def _get_template_machine_names() -> List[str]:
    csv_text = load_builtin_csv_template()
    names: List[str] = []
    for line in csv_text.strip().split("\n")[1:]:
        if not line.strip():
            continue
        name = line.split(",", 1)[0].strip()
        if name:
            names.append(name)
    return names


def _safe_csv_prefix(account_name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9\-]", "-", (account_name or "customer").strip()).strip("-") or "customer"


def prefix_csv_server_names(csv_text: str, prefix: str) -> str:
    """给 CSV 中的服务器名添加客户前缀，并将序号随机化以避免多人导入时序号冲突。"""
    import random as _rng
    lines = csv_text.strip().split("\n")
    if not lines:
        return csv_text

    # 使用 prefix 作为随机种子，确保同一客户每次生成的序号相同（tier cache 可匹配）
    seed = int(hashlib.md5(prefix.encode()).hexdigest(), 16) % (2**32)
    rng = _rng.Random(seed)

    # 收集所有数字后缀并生成随机映射
    num_pattern = re.compile(r"(\d+)")
    # 找出所有行中使用的数字，生成一个不重复的随机序号池
    all_numbers = set()
    for line in lines[1:]:
        if not line.strip():
            continue
        name = line.split(",", 1)[0].strip()
        for m in num_pattern.finditer(name):
            all_numbers.add(int(m.group()))

    # 生成随机映射：原始序号 → 随机序号（范围扩大避免冲突）
    max_num = max(all_numbers) if all_numbers else 50
    pool_start = rng.randint(100, 800)
    number_map: Dict[int, int] = {}
    used_numbers = set()
    for orig_num in sorted(all_numbers):
        new_num = pool_start + rng.randint(1, 5)
        while new_num in used_numbers:
            new_num += rng.randint(1, 3)
        number_map[orig_num] = new_num
        used_numbers.add(new_num)
        pool_start = new_num

    result = [lines[0]]
    for line in lines[1:]:
        if not line.strip():
            continue
        parts = line.split(",", 1)
        if len(parts) >= 2:
            original_name = parts[0].strip()
            # 替换名称中的数字为随机化后的数字
            def _replace_num(m):
                orig = int(m.group())
                return str(number_map.get(orig, orig))
            randomized_name = num_pattern.sub(_replace_num, original_name)
            result.append(f"{prefix}-{randomized_name},{parts[1]}")
        else:
            result.append(line)
    return "\n".join(result)


def _csv_template_hash() -> str:
    try:
        with open(BUILTIN_CSV_PATH, "rb") as f:
            return hashlib.md5(f.read()).hexdigest()[:12]
    except Exception:
        return "unknown"


def snap_budget_to_tier(annual_budget: float) -> int:
    if annual_budget is None or annual_budget <= 0:
        return BUDGET_TIERS[-1]
    for tier in BUDGET_TIERS:
        if annual_budget <= tier * 1.15:
            return tier
    return BUDGET_TIERS[-1]


def load_tier_cache() -> Dict[str, Any]:
    if not os.path.exists(TIER_CACHE_PATH):
        return {}
    try:
        with open(TIER_CACHE_PATH, "r", encoding="utf-8") as f:
            cache = json.load(f)
        if cache.get("template_hash") != _csv_template_hash():
            return {}
        created = cache.get("created_at", "")
        if created:
            created_dt = datetime.datetime.fromisoformat(created)
            if (datetime.datetime.now() - created_dt).days > 7:
                return {}
        return cache
    except Exception:
        return {}


def save_tier_cache(cache: Dict[str, Any]) -> None:
    cache["created_at"] = datetime.datetime.now().isoformat()
    cache["template_hash"] = _csv_template_hash()
    with open(TIER_CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump(cache, f, indent=2, ensure_ascii=False)


def _get_session_persist_path() -> str:
    """返回当前 Streamlit session 对应的持久化文件路径（per-session 隔离）。"""
    try:
        from streamlit.runtime.scriptrunner import get_script_run_ctx
        ctx = get_script_run_ctx()
        if ctx and ctx.session_id:
            sid = hashlib.md5(ctx.session_id.encode()).hexdigest()[:8]
        else:
            sid = "default"
    except Exception:
        sid = "default"
    return os.path.join(PERSIST_DIR, f".session_persist_{sid}.json")


def _cleanup_stale_sessions(max_age_hours: int = 24) -> None:
    """删除超过 max_age_hours 的旧 session 持久化文件。"""
    try:
        cutoff = time.time() - max_age_hours * 3600
        for fname in os.listdir(PERSIST_DIR):
            if fname.startswith(".session_persist_") and fname.endswith(".json"):
                fpath = os.path.join(PERSIST_DIR, fname)
                if os.path.getmtime(fpath) < cutoff:
                    os.remove(fpath)
    except Exception:
        pass


def persist_session_state() -> None:
    data: Dict[str, Any] = {}
    for k in _PERSIST_KEYS:
        if k in st.session_state:
            data[k] = st.session_state[k]
    if not data:
        return
    try:
        path = _get_session_persist_path()
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
    except Exception:
        pass


def restore_session_state() -> None:
    _cleanup_stale_sessions()
    path = _get_session_persist_path()
    if not os.path.exists(path):
        return
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return
    for k, v in data.items():
        if k not in st.session_state:
            st.session_state[k] = v


def clear_session_persist() -> None:
    try:
        os.remove(_get_session_persist_path())
    except OSError:
        pass


def _assessed_machine_monthly_cost(machine: Dict[str, Any]) -> float:
    props = machine.get("properties", {})
    cost = 0.0
    for key in ("monthlyComputeCostForRecommendedSize", "monthlyStorageCost", "monthlyBandwidthCost"):
        try:
            cost += float(props.get(key, 0) or 0)
        except (TypeError, ValueError):
            pass
    for comp in props.get("costComponents") or []:
        if str(comp.get("name", "")).lower() == "monthlysecuritycost":
            try:
                cost += float(comp.get("value", 0) or 0)
            except (TypeError, ValueError):
                pass
    return cost


def _strip_account_prefix(display_name: str, prefix: str) -> str:
    if prefix and display_name.lower().startswith(prefix.lower() + "-"):
        return display_name[len(prefix) + 1:]
    return display_name


def learn_tier_machine_selections(
    assessed_machines: List[Dict[str, Any]],
    account_prefix: str,
    progress: Callable[[str], None],
) -> Dict[str, Any]:
    machine_costs: List[Dict[str, Any]] = []
    for m in assessed_machines:
        display_name = m.get("properties", {}).get("displayName", "")
        monthly_cost = _assessed_machine_monthly_cost(m)
        original_name = _strip_account_prefix(display_name, account_prefix)
        machine_costs.append({
            "template_name": original_name,
            "monthly_cost": monthly_cost,
        })

    total_monthly = sum(mc["monthly_cost"] for mc in machine_costs)
    total_annual = total_monthly * 12
    progress(
        f"全量评估学习基准：{len(machine_costs)} 台服务器，"
        f"年化 {_format_usd(total_annual)}"
    )

    machine_costs.sort(key=lambda x: x["monthly_cost"])

    cache: Dict[str, Any] = {
        "total_monthly": total_monthly,
        "total_annual": total_annual,
        "machine_count": len(machine_costs),
        "tiers": {},
    }

    for tier in BUDGET_TIERS:
        target_monthly = tier / 12
        target_max_monthly = tier * 1.2 / 12

        if total_annual <= tier * 1.2:
            selected = list(machine_costs)
        else:
            selected: List[Dict[str, Any]] = []
            running = 0.0
            for mc in machine_costs:
                if running >= target_monthly:
                    break
                if running + mc["monthly_cost"] <= target_max_monthly:
                    selected.append(mc)
                    running += mc["monthly_cost"]

            if sum(s["monthly_cost"] for s in selected) < target_monthly:
                remaining = [mc for mc in machine_costs if mc not in selected]
                running = sum(s["monthly_cost"] for s in selected)
                for mc in remaining:
                    if running + mc["monthly_cost"] > target_max_monthly:
                        break
                    selected.append(mc)
                    running += mc["monthly_cost"]
                    if running >= target_monthly:
                        break

        sel_monthly = sum(s["monthly_cost"] for s in selected)
        sel_names = [s["template_name"] for s in selected]

        cache["tiers"][str(tier)] = {
            "machine_names": sel_names,
            "machine_count": len(sel_names),
            "expected_monthly": round(sel_monthly, 2),
            "expected_annual": round(sel_monthly * 12, 2),
        }
        progress(
            f"  规模 {_format_usd(float(tier))}：选择 {len(sel_names)}/{len(machine_costs)} 台，"
            f"预期年化 {_format_usd(sel_monthly * 12)}"
        )

    return cache


def get_machine_ids_for_tier(
    tier: int,
    machines: List[Dict[str, Any]],
    account_prefix: str,
    cache: Dict[str, Any],
) -> List[str]:
    tier_data = cache.get("tiers", {}).get(str(tier))
    if not tier_data:
        return [m.get("id") for m in machines if m.get("id")]

    selected_template_names = {n.lower() for n in tier_data["machine_names"]}

    # 先尝试按名称匹配
    selected_ids: List[str] = []
    for m in machines:
        display_name = m.get("properties", {}).get("displayName", "")
        original_name = _strip_account_prefix(display_name, account_prefix)
        if original_name.lower() in selected_template_names:
            mid = m.get("id")
            if mid:
                selected_ids.append(mid)

    # 如果名称匹配失败（例如缓存来自旧命名方案），按缓存的机器数量选取
    expected_count = tier_data.get("machine_count", len(selected_template_names))
    if len(selected_ids) < expected_count:
        all_ids = [m.get("id") for m in machines if m.get("id")]
        selected_ids = all_ids[:expected_count]

    return selected_ids


def _assessment_cost_component(assessment: Dict[str, Any], component_name: str) -> float:
    props = assessment.get("properties", {})
    for component in props.get("costComponents") or []:
        if str(component.get("name") or "").lower() == component_name.lower():
            try:
                return float(component.get("value") or 0)
            except (TypeError, ValueError):
                return 0.0
    return 0.0


def assessment_monthly_total_cost(assessment: Dict[str, Any]) -> float:
    props = assessment.get("properties", {})

    def _num(name: str) -> float:
        try:
            return float(props.get(name) or 0)
        except (TypeError, ValueError):
            return 0.0

    return (
        _num("monthlyComputeCost")
        + _num("monthlyStorageCost")
        + _num("monthlyBandwidthCost")
        + _assessment_cost_component(assessment, "MonthlySecurityCost")
    )


def _format_usd(value: Optional[float]) -> str:
    if value is None:
        return "未填写"
    return f"${value:,.2f}"


def _assessment_settings_snapshot(settings: Dict[str, Any]) -> Dict[str, Any]:
    keys = [
        "azureLocation",
        "sizingCriterion",
        "reservedInstance",
        "azureHybridUseBenefit",
        "linuxAzureHybridUseBenefit",
        "azureSecurityOfferingType",
        "scalingFactor",
        "discountPercentage",
    ]
    return {key: settings.get(key) for key in keys}


# Azure Migrate 评估地区名称映射（解决方案文档中的区域名 → API azureLocation 值）
_REGION_NAME_TO_LOCATION = {
    "east us": "EastUs", "east us 2": "EastUs2", "west us": "WestUs",
    "west us 2": "WestUs2", "west us 3": "WestUs3", "central us": "CentralUs",
    "north central us": "NorthCentralUs", "south central us": "SouthCentralUs",
    "west europe": "WestEurope", "north europe": "NorthEurope",
    "southeast asia": "SoutheastAsia", "east asia": "EastAsia",
    "japan east": "JapanEast", "japan west": "JapanWest",
    "australia east": "AustraliaEast", "australia southeast": "AustraliaSoutheast",
    "uk south": "UKSouth", "uk west": "UKWest",
    "canada central": "CanadaCentral", "canada east": "CanadaEast",
    "korea central": "KoreaCentral", "korea south": "KoreaSouth",
    "france central": "FranceCentral", "germany west central": "GermanyWestCentral",
    "switzerland north": "SwitzerlandNorth", "norway east": "NorwayEast",
    "brazil south": "BrazilSouth", "south africa north": "SouthAfricaNorth",
    "uae north": "UAENorth", "india central": "CentralIndia",
    "india south": "SouthIndia", "india west": "WestIndia",
    "sweden central": "SwedenCentral", "qatar central": "QatarCentral",
}


def _extract_dominant_region(solution_text: str) -> Optional[str]:
    """
    从解决方案文本中提取出现频率最高的 Azure 区域，返回对应的 azureLocation API 值。
    如果无法识别任何区域，返回 None。
    """
    if not solution_text:
        return None
    text_lower = solution_text.lower()
    region_counts: Dict[str, int] = {}
    # 按名称长度降序匹配，避免 "east us" 匹配到 "east us 2" 的情况
    sorted_regions = sorted(_REGION_NAME_TO_LOCATION.keys(), key=len, reverse=True)
    for region_name in sorted_regions:
        count = text_lower.count(region_name)
        if count > 0:
            location_val = _REGION_NAME_TO_LOCATION[region_name]
            region_counts[location_val] = region_counts.get(location_val, 0) + count
    if not region_counts:
        return None
    # 返回出现次数最多的区域
    return max(region_counts, key=region_counts.get)


def _build_assessment_body(target_location: Optional[str] = None) -> Dict[str, Any]:
    return {
        "properties": {
            "groupType": "Import",
            "assessmentType": "MachineAssessment",
            "azureLocation": target_location or AZURE_MIGRATE_DEFAULT_TARGET_LOCATION,
            "azureOfferCode": "MSAZR0003P",
            "azurePricingTier": "Standard",
            "azureStorageRedundancy": "LocallyRedundant",
            "scalingFactor": 1.3,
            "percentile": "Percentile95",
            "timeRange": "Day",
            "currency": "USD",
            "azureHybridUseBenefit": "Yes",
            "linuxAzureHybridUseBenefit": "Yes",
            "azureSecurityOfferingType": "MDC",
            "discountPercentage": 0,
            "sizingCriterion": "PerformanceBased",
            "azureDiskTypes": ["Premium", "StandardSSD", "Standard"],
            "azureVmFamilies": [
                "Dv2_series", "Dv3_series", "DSv2_series", "Dsv3_series", "Ev3_series",
                "Esv3_series", "F_series", "Fs_series", "Fsv2_series", "M_series", "D_series",
                "DS_series", "H_series", "Lsv2_series",
            ],
            "vmUptime": {"daysPerMonth": 31, "hoursPerDay": 24},
            "reservedInstance": "RI3Year",
            "stage": "InProgress",
        },
    }


def tune_assessment_to_budget(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    group_name: str,
    assessment_name: str,
    assessment_path: str,
    assessment_body: Dict[str, Any],
    assessment: Dict[str, Any],
    annual_budget: Optional[float],
    token: str,
    progress: Callable[[str], None],
) -> tuple[Dict[str, Any], List[Dict[str, Any]], bool]:
    history: List[Dict[str, Any]] = []
    target_min = annual_budget if annual_budget and annual_budget > 0 else None
    # 目标上限为当前档次的天花板（不超过下一个 tier），而非固定 120%
    if annual_budget and annual_budget > 0:
        current_tier = snap_budget_to_tier(annual_budget)
        tier_idx = BUDGET_TIERS.index(current_tier) if current_tier in BUDGET_TIERS else -1
        if tier_idx < len(BUDGET_TIERS) - 1:
            target_max = float(BUDGET_TIERS[tier_idx + 1])
        else:
            target_max = annual_budget * 1.5  # 最大档次无上限约束，放宽到 150%
        target_mid = (target_min + target_max) / 2
    else:
        target_max = None
        target_mid = None

    def _record(round_name: str, action: str, current: Dict[str, Any], met: bool) -> None:
        monthly_total = assessment_monthly_total_cost(current)
        annual_total = monthly_total * 12
        history.append({
            "round": round_name,
            "action": action,
            "monthly_total": monthly_total,
            "annual_total": annual_total,
            "target_annual": annual_budget,
            "target_min": target_min,
            "target_max": target_max,
            "met_target": met,
            "settings": _assessment_settings_snapshot(assessment_body["properties"]),
        })

    def _in_target_range(current: Dict[str, Any]) -> bool:
        if target_min is None or target_max is None:
            return True
        annual_total = assessment_monthly_total_cost(current) * 12
        return target_min <= annual_total <= target_max

    def _next_patch(annual_total: float) -> tuple[str, Dict[str, Any]]:
        settings = assessment_body["properties"]
        if target_max is not None and target_mid is not None and annual_total > target_max:
            current_discount = float(settings.get("discountPercentage") or 0)
            current_discount_factor = max(1 - current_discount / 100, 0.01)
            undiscounted_annual = annual_total / current_discount_factor
            required_discount = 100 * (1 - (target_mid / max(undiscounted_annual, 1)))
            next_discount = min(max(current_discount, required_discount), 99.0)
            if next_discount > current_discount + 0.1:
                return (
                    f"设置折扣为 {next_discount:.2f}%，控制年化估算不超过预估值 20%",
                    {"discountPercentage": round(next_discount, 2)},
                )
            current_factor = float(settings.get("scalingFactor") or 1.3)
            if current_factor > 1.0:
                next_factor = max(1.0, current_factor * 0.8)
                return (
                    f"降低舒适因子到 {next_factor:.2f}，控制年化估算不超过预估值 20%",
                    {"scalingFactor": round(next_factor, 2)},
                )
            if settings.get("azureSecurityOfferingType") != "NO":
                return ("关闭安全成本估算，控制年化估算不超过预估值 20%", {"azureSecurityOfferingType": "NO"})
            return ("已达到自动降价边界", {})

        if target_mid is None:
            return ("未设置预算，不调整", {})

        current_discount = float(settings.get("discountPercentage") or 0)
        if current_discount > 0:
            return ("取消折扣以提高年化估算", {"discountPercentage": 0})
        if settings.get("azureHybridUseBenefit") != "No" or settings.get("linuxAzureHybridUseBenefit") != "No":
            return (
                "关闭 Azure Hybrid Benefit，把 OS 许可成本计入估算",
                {"azureHybridUseBenefit": "No", "linuxAzureHybridUseBenefit": "No"},
            )
        current_factor = float(settings.get("scalingFactor") or 1.3)
        factor_ratio = min(max(target_mid / max(annual_total, 1), 1.05), 1.35)
        next_factor = min(current_factor * factor_ratio, 5.0)
        if next_factor > current_factor + 0.01:
            return (f"提高舒适因子到 {next_factor:.2f}", {"scalingFactor": round(next_factor, 2)})
        if settings.get("reservedInstance") != "None":
            return ("切换为按量计费以提高年化估算", {"reservedInstance": "None"})
        return ("已达到自动提价边界", {})

    if annual_budget is None or annual_budget <= 0:
        monthly_total = assessment_monthly_total_cost(assessment)
        progress(
            "未填写可解析的预估年消耗，跳过价格校准。"
            f"当前 Azure Migrate 年化估算：{_format_usd(monthly_total * 12)}"
        )
        _record("initial", "未填写可解析预算，未调整评估设置", assessment, True)
        return assessment, history, True

    monthly_total = assessment_monthly_total_cost(assessment)
    annual_total = monthly_total * 12
    progress(
        "Azure Migrate 当前年化估算："
        f"{_format_usd(annual_total)}；目标区间：{_format_usd(target_min)} - {_format_usd(target_max)}"
    )
    if _in_target_range(assessment):
        _record("initial", "初始 Portal 默认评估已在目标区间内", assessment, True)
        return assessment, history, True
    direction = "高于" if target_max is not None and annual_total > target_max else "低于"
    _record("initial", f"初始 Portal 默认评估{direction}目标区间", assessment, False)

    for round_index in range(1, 4):
        round_name = f"round-{round_index}"
        action, patch = _next_patch(annual_total)
        if not patch:
            progress(f"评估年化估算仍不在目标区间内，{action}。")
            break
        progress(f"评估年化估算不在目标区间，开始自动调整（{round_name}）：{action}")
        assessment_body["properties"].update(patch)
        assessment_body["properties"]["stage"] = "InProgress"
        azure_arm_request("PUT", assessment_path, token, assessment_body)
        assessment = wait_for_assessment_complete(
            subscription_id, resource_group, project_name, group_name, assessment_name, token, progress
        )
        monthly_total = assessment_monthly_total_cost(assessment)
        annual_total = monthly_total * 12
        met = _in_target_range(assessment)
        progress(
            f"{round_name} 重新计算完成：月估算 {_format_usd(monthly_total)}，"
            f"年化 {_format_usd(annual_total)}，目标区间 {_format_usd(target_min)} - {_format_usd(target_max)}"
        )
        _record(round_name, action, assessment, met)
        if met:
            return assessment, history, True

    progress(
        f"已自动调整 3 轮，但 Azure Migrate 年化估算仍未落入用户预估年消耗的档次区间"
        f"（{_format_usd(target_min)} ~ {_format_usd(target_max)}）；"
        "请到 Azure Portal 的评估设置中手动调整后重新导出。"
    )
    return assessment, history, False


def download_assessment_report(
    subscription_id: str,
    resource_group: str,
    project_name: str,
    group_name: str,
    assessment_name: str,
    token: str,
) -> bytes:
    download_url_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/groups/{group_name}/assessments/{assessment_name}/downloadUrl"
        f"?api-version={AZURE_MIGRATE_REPORT_API_VERSION}"
    )
    payload = azure_arm_request("POST", download_url_path, token, poll_lro=False)
    report_url = payload.get("assessmentReportUrl")
    if not report_url:
        raise RuntimeError(f"Azure Migrate 未返回评估报告下载地址：{payload}")

    response = requests.get(report_url, timeout=180)
    if response.status_code >= 400:
        raise RuntimeError(f"下载 Azure Migrate 评估报告失败：{response.status_code} {response.text[:800]}")
    if not response.content:
        raise RuntimeError("Azure Migrate 评估报告为空。")
    return response.content


def _try_get_existing_resource(path: str, token: str) -> Optional[Dict[str, Any]]:
    """尝试 GET 某资源，如果存在返回 dict，不存在返回 None。"""
    try:
        result = azure_arm_request("GET", path, token)
        if result and result.get("id"):
            return result
    except Exception:
        pass
    return None


def wait_for_group_machine_membership(
    group_path: str,
    token: str,
    expected_machine_count: int,
    progress: Callable[[str], None],
    timeout_seconds: int = 600,
) -> Dict[str, Any]:
    """等待 updateMachines 完成，并返回最新评估组信息。"""
    deadline = time.time() + timeout_seconds
    attempt = 0
    last_group: Dict[str, Any] = {}

    while time.time() < deadline:
        attempt += 1
        last_group = azure_arm_request("GET", group_path, token, poll_lro=False)
        props = last_group.get("properties", {})
        machine_count = int(props.get("machineCount") or 0)
        provisioning_state = props.get("provisioningState", "Unknown")
        supported_types = sorted(
            str(item).strip()
            for item in (props.get("supportedAssessmentTypes") or [])
            if str(item).strip()
        )
        supported_label = ", ".join(supported_types) if supported_types else "未返回"

        if machine_count >= expected_machine_count and provisioning_state not in {"Failed", "Canceled"}:
            progress(f"  ℹ️ 评估组已包含 {machine_count} 台服务器；支持类型: {supported_label}")
            return last_group
        if provisioning_state in {"Failed", "Canceled"}:
            raise RuntimeError(f"评估组更新失败，资源状态：{provisioning_state}")

        progress(
            f"评估组仍在关联服务器，当前 {machine_count}/{expected_machine_count} 台；"
            f"支持类型: {supported_label}（第 {attempt} 次检查）"
        )
        time.sleep(10)

    props = last_group.get("properties", {})
    raise TimeoutError(
        "等待评估组关联服务器超时。"
        f"当前 machineCount={props.get('machineCount')}，"
        f"supportedAssessmentTypes={props.get('supportedAssessmentTypes')}"
    )


def run_azure_migrate_assessment(
    token: str,
    subscription_id: str,
    resource_group: str,
    account_name: str,
    assessment_name: str,
    annual_budget_text: Optional[str],
    progress: Callable[[str], None],
    target_location: Optional[str] = None,
) -> Dict[str, Any]:
    # ── 加载内置 CSV 模板，给服务器名加上客户前缀 ──
    progress("加载内置服务器清单模板...")
    csv_text_raw = load_builtin_csv_template()
    safe_prefix = _safe_csv_prefix(account_name)
    csv_text = prefix_csv_server_names(csv_text_raw, safe_prefix)
    csv_bytes = csv_text.encode("utf-8-sig")
    progress(f"  ✅ 已为所有服务器名添加前缀: {safe_prefix}-")

    annual_budget = parse_annual_budget_usd(annual_budget_text)
    tier = snap_budget_to_tier(annual_budget) if annual_budget and annual_budget > 0 else BUDGET_TIERS[-1]
    progress(f"客户预估年消耗: {_format_usd(annual_budget)}，匹配规模档位: {_format_usd(float(tier))}")

    safe_base = _safe_azure_name(account_name, f"poe-{_date_prefix()}", max_len=36).lower()
    run_suffix = str(int(time.time()))
    short_run_suffix = run_suffix[-6:]
    project_name = _safe_azure_name(safe_base, "poe", "project", 55)
    site_name = _safe_azure_name(safe_base, "poe", "site", 24)
    master_site_name = _safe_azure_name(safe_base, "poe", "masterSite", 55)
    collector_name = _safe_azure_name(safe_base, "poe", "collector", 55)
    group_name = _safe_azure_name(safe_base, "poe", f"group-{run_suffix}", 55)
    assessment_resource_name = _safe_azure_name(assessment_name, "poe-assessment", max_len=55)
    project_location = resolve_migrate_project_location(subscription_id, resource_group, token)
    annual_budget = parse_annual_budget_usd(annual_budget_text)
    migrate_project_id = _resource_id(
        subscription_id,
        resource_group,
        f"providers/Microsoft.Migrate/migrateProjects/{project_name}",
    )
    assessment_project_id = _resource_id(
        subscription_id,
        resource_group,
        f"providers/Microsoft.Migrate/assessmentProjects/{project_name}",
    )
    master_site_id = _resource_id(
        subscription_id,
        resource_group,
        f"providers/Microsoft.OffAzure/masterSites/{master_site_name}",
    )
    import_site_id = _resource_id(
        subscription_id,
        resource_group,
        f"providers/Microsoft.OffAzure/importSites/{site_name}",
    )

    progress("注册 Microsoft.Migrate 与 Microsoft.OffAzure 资源提供程序...")
    register_azure_provider(subscription_id, "Microsoft.Migrate", token)
    register_azure_provider(subscription_id, "Microsoft.OffAzure", token)

    # ── Step 1: 创建 migrateProject（Portal 可见） ──
    progress(f"创建 Azure Migrate 项目（区域：{project_location}）...")
    migrate_project_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/migrateProjects/{project_name}"
        f"?api-version={AZURE_MIGRATE_PROJECTS_API_VERSION}"
    )
    existing_mp = _try_get_existing_resource(migrate_project_path, token)
    if existing_mp:
        mp_id = existing_mp.get("id", project_name)
        progress(f"成功复用 Azure Migrate 项目：{project_name}")
    else:
        migrate_project_body = {
            "properties": {},
            "location": project_location,
            "tags": {"Migrate Project": project_name, "createdBy": "POE Workflow"},
            "identity": {"type": "SystemAssigned"},
        }
        try:
            mp_result = azure_arm_request("PUT", migrate_project_path, token, migrate_project_body)
        except Exception:
            migrate_project_body.pop("identity", None)
            mp_result = azure_arm_request("PUT", migrate_project_path, token, migrate_project_body)
        mp_id = mp_result.get("id", project_name)
        progress(f"成功创建 Azure Migrate 项目：{project_name}")

    # ── Step 2: 注册 Portal 同款 Discovery Import 与 Assessment 工具 ──
    progress("注册 ServerDiscovery_Import 与 ServerAssessment 工具...")
    for tool in ("ServerDiscovery_Import", "ServerAssessment"):
        try:
            register_migrate_tool(subscription_id, resource_group, project_name, tool, token)
            progress(f"  ✅ 工具已注册: {tool}")
        except Exception:
            progress(f"  ℹ️ 工具可能已注册: {tool}")

    # ── Step 3: 创建 ServerAssessment Solution ──
    assessment_solution_name = "Servers-Assessment-ServerAssessment"
    assessment_solution_path = _migrate_solution_path(
        subscription_id, resource_group, project_name, assessment_solution_name
    )
    existing_sol = _try_get_existing_resource(assessment_solution_path, token)
    if existing_sol:
        assessment_solution_id = existing_sol.get("id", "")
        progress(f"成功复用评估 Solution：{assessment_solution_name}")
    else:
        sol_result = put_migrate_solution(
            subscription_id,
            resource_group,
            project_name,
            assessment_solution_name,
            {
                "tool": "ServerAssessment",
                "purpose": "Assessment",
                "goal": "Servers",
                "status": "Active",
                "details": _servers_solution_details({
                    "projectId": assessment_project_id,
                    "avsAssessment": "0",
                    "azureSqlAssessment": "0",
                    "azureVmAssessment": "0",
                    "azureWebAppAssessment": "0",
                    "businessCaseCount": "0",
                }),
            },
            token,
        )
        assessment_solution_id = sol_result.get("id", "")
        progress(f"成功创建评估 Solution：{assessment_solution_name}")

    # ── Step 4: 创建 assessmentProject 并关联 Solution ──
    progress("创建 Assessment Project...")
    ap_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    existing_ap = _try_get_existing_resource(ap_path, token)
    ap_body = {
        "kind": "Migrate",
        "properties": {
            "projectStatus": "Active",
            "assessmentSolutionId": assessment_solution_id,
            "publicNetworkAccess": "Enabled",
        },
        "location": project_location,
        "tags": {"createdBy": "POE Workflow"},
    }
    if existing_ap and str(existing_ap.get("kind") or "").lower() == "migrate":
        ap_id = existing_ap.get("id", project_name)
        progress(f"成功复用 Assessment Project：{project_name}")
    else:
        ap_result = azure_arm_request("PUT", ap_path, token, ap_body)
        ap_id = ap_result.get("id", project_name)
        progress(f"成功创建 Assessment Project：{project_name}")

    # Portal 的评估 blade 通过 Assessment Solution 的 projectId 找到 assessmentProject。
    assessment_solution = put_migrate_solution(
        subscription_id,
        resource_group,
        project_name,
        assessment_solution_name,
        {
            "tool": "ServerAssessment",
            "purpose": "Assessment",
            "goal": "Servers",
            "status": "Active",
            "details": _servers_solution_details({
                "projectId": assessment_project_id,
                "avsAssessment": "0",
                "azureSqlAssessment": "0",
                "azureVmAssessment": "0",
                "azureWebAppAssessment": "0",
                "businessCaseCount": "0",
            }),
        },
        token,
    )
    assessment_solution_id = assessment_solution.get("id", assessment_solution_id)
    progress("  ✅ Assessment Solution 已关联 assessmentProject")

    # ── Step 5: 创建 Portal Discovery Import 链路（Master Site + Discovery Solution + Import Site） ──
    progress("创建 Portal 可识别的 Discovery Import 链路...")
    discovery_solution_name = "Servers-Discovery-ServerDiscovery_Import"
    discovery_solution_id = _migrate_solution_id(
        subscription_id, resource_group, project_name, discovery_solution_name
    )
    master_site_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.OffAzure/masterSites/{master_site_name}"
        f"?api-version={AZURE_OFFAZURE_API_VERSION}"
    )
    existing_master_site = _try_get_existing_resource(master_site_path, token)
    existing_sites = []
    if existing_master_site:
        existing_sites = existing_master_site.get("properties", {}).get("sites") or []
    master_site_result = azure_arm_request("PUT", master_site_path, token, {
        "kind": "Migrate",
        "location": project_location,
        "tags": {"Migrate Project": project_name, "createdBy": "POE Workflow"},
        "properties": {
            "allowMultipleSites": True,
            "publicNetworkAccess": "Enabled",
            "sites": existing_sites,
        },
    })
    master_site_id = master_site_result.get("id", master_site_id)
    progress(f"成功创建 Master Site：{master_site_name}")
    ensure_portal_menu_solutions(
        subscription_id,
        resource_group,
        project_name,
        master_site_id,
        token,
        progress,
    )

    put_migrate_solution(
        subscription_id,
        resource_group,
        project_name,
        discovery_solution_name,
        {
            "tool": "ServerDiscovery_Import",
            "purpose": "Discovery",
            "goal": "Servers",
            "status": "Inactive",
            "details": _servers_solution_details({
                "importSiteId": import_site_id,
            }),
        },
        token,
    )
    progress("  ✅ Discovery Import Solution 已关联 importSite")

    # ── Step 6: 创建 Import Site ──
    progress("创建 Import Site...")
    site_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.OffAzure/importSites/{site_name}"
        f"?api-version={AZURE_OFFAZURE_API_VERSION}"
    )
    existing_site = _try_get_existing_resource(site_path, token)
    if existing_site:
        site_id = existing_site.get("id", site_name)
        progress(f"成功复用 Import Site：{site_name}")
    else:
        site_result = azure_arm_request("PUT", site_path, token, {
            "location": project_location,
            "properties": {
                "masterSiteId": master_site_id,
                "discoverySolutionId": discovery_solution_id,
            },
        })
        site_id = site_result.get("id", site_name)
        progress(f"成功创建 Import Site：{site_name}")

    normalized_sites = {str(site).lower(): site for site in existing_sites}
    normalized_sites.setdefault(import_site_id.lower(), import_site_id)
    azure_arm_request("PUT", master_site_path, token, {
        "kind": "Migrate",
        "location": project_location,
        "tags": {"Migrate Project": project_name, "createdBy": "POE Workflow"},
        "properties": {
            "allowMultipleSites": True,
            "publicNetworkAccess": "Enabled",
            "sites": list(normalized_sites.values()),
        },
    })
    progress("  ✅ Master Site 已关联 Import Site")

    # ── Step 7: 创建 Import Collector，关联 Import Site 到 Assessment Project ──
    progress("创建 Import Collector 关联 Import Site 到 Assessment Project...")
    collector_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/importcollectors/{collector_name}?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    discovery_site_id = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.OffAzure/importSites/{site_name}"
    )
    coll_result = azure_arm_request("PUT", collector_path, token, {
        "properties": {"discoverySiteId": discovery_site_id}
    })
    progress(f"成功创建 Import Collector：{collector_name}")

    # ── Step 8: 检查当前 Import Site 是否已有库存 —— 有则跳过 CSV 导入 ──
    site_machines_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.OffAzure/importSites/{site_name}"
        f"/machines?api-version={AZURE_OFFAZURE_API_VERSION}"
    )
    try:
        existing_site_machines = azure_arm_list(site_machines_path, token)
    except Exception:
        existing_site_machines = []

    if existing_site_machines:
        progress(f"  ℹ️ Import Site 已有 {len(existing_site_machines)} 台服务器库存，跳过 CSV 重新导入")
        imported_site_machines = existing_site_machines
        portal_inventory_count = len(existing_site_machines)
    else:
        # ── 获取 SAS URL 并上传 CSV ──
        progress("获取 CSV 上传地址并上传服务器清单...")
        import_uri_path = (
            f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
            f"/providers/Microsoft.OffAzure/importSites/{site_name}/importUri"
            f"?api-version={AZURE_OFFAZURE_API_VERSION}"
        )
        import_uri_payload = azure_arm_request("POST", import_uri_path, token, {})
        sas_url = _extract_sas_url(import_uri_payload)
        if not sas_url:
            raise RuntimeError(f"Azure 未返回可用的 CSV 上传 SAS URL：{import_uri_payload}")
        import_job_arm_id = import_uri_payload.get("jobArmId") if isinstance(import_uri_payload, dict) else None

        upload_response = requests.put(
            sas_url,
            data=csv_bytes,
            headers={"x-ms-blob-type": "BlockBlob", "Content-Type": "text/csv"},
            timeout=180,
        )
        if upload_response.status_code >= 400:
            raise RuntimeError(f"上传 CSV 到 Azure Migrate 失败：{upload_response.status_code} {upload_response.text[:800]}")
        progress(f"  ✅ CSV 已上传（{len(csv_bytes)} 字节）")

        # ── Step 9: 触发 Import Job（回传 importUri 返回的 SasUriResponse） ──
        progress("触发 Import Job 导入服务器清单...")
        import_trigger_body = dict(import_uri_payload) if isinstance(import_uri_payload, dict) else {}
        import_trigger_body["uri"] = sas_url
        if import_job_arm_id:
            import_trigger_body["jobArmId"] = import_job_arm_id
        job_result = azure_arm_request("POST", import_uri_path, token, import_trigger_body)
        import_job_arm_id = (
            job_result.get("jobArmId")
            or import_job_arm_id
            or job_result.get("id")
            if isinstance(job_result, dict)
            else import_job_arm_id
        )
        progress("成功触发 Import Job")

        # ── Step 10: 等待 OffAzure Import Site 完成 CSV 解析 ──
        imported_site_machines = wait_for_import_site_import(
            subscription_id,
            resource_group,
            site_name,
            token,
            progress,
            job_arm_id=import_job_arm_id,
        )
        progress(f"  ✅ Import Site 已导入 {len(imported_site_machines)} 台服务器")

        # Import Collector 在导入完成后再 PUT 一次，触发 assessmentProject 拉取刚导入的机器。
        azure_arm_request("PUT", collector_path, token, {
            "properties": {"discoverySiteId": discovery_site_id}
        })
        progress("  ✅ Import Collector 已刷新同步")

        portal_inventory_count = wait_for_portal_inventory_summary(
            subscription_id,
            resource_group,
            project_name,
            len(imported_site_machines),
            token,
            progress,
        )
        progress(f"  ✅ Azure Portal 全部库存汇总已刷新: {portal_inventory_count} 台服务器")

    # ── Step 11: 等待 Assessment Project 可读取机器 ──
    machines = wait_for_project_machines(
        subscription_id,
        resource_group,
        project_name,
        collector_name,
        token,
        progress,
        site_name=site_name,
    )
    all_machine_ids = [machine.get("id") for machine in machines if machine.get("id")]
    if not all_machine_ids:
        raise RuntimeError("Azure Migrate 未返回可加入评估的服务器，请检查 CSV 导入结果。")
    progress(f"  ✅ 已发现 {len(all_machine_ids)} 台服务器")

    # ── Step 12: 分规模学习 — 确定当前规模应选哪些服务器 ──
    cache = load_tier_cache()
    tier_cached = bool(cache.get("tiers", {}).get(str(tier)))

    if not tier_cached:
        progress(f"规模 {_format_usd(float(tier))} 尚未学习，开始全量评估学习...")
        learning_group_name = _safe_azure_name(safe_base, "poe", f"learn-{run_suffix}", 55)
        learning_group_path = (
            f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
            f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
            f"/groups/{learning_group_name}?api-version={AZURE_MIGRATE_API_VERSION}"
        )
        azure_arm_request("PUT", learning_group_path, token, {
            "properties": {"groupType": "Import"},
            "eTag": "",
        })
        progress(f"  ✅ 已创建学习评估组: {learning_group_name}")

        learn_update_path = (
            f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
            f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
            f"/groups/{learning_group_name}/updateMachines?api-version={AZURE_MIGRATE_API_VERSION}"
        )
        azure_arm_request("POST", learn_update_path, token, {
            "eTag": "*",
            "properties": {"operationType": "Add", "machines": all_machine_ids},
        })
        wait_for_group_machine_membership(
            learning_group_path, token,
            expected_machine_count=len(all_machine_ids),
            progress=progress,
        )
        progress(f"  ✅ 学习评估组已关联 {len(all_machine_ids)} 台服务器")

        learning_assess_name = _safe_azure_name("learning", "poe-assess", max_len=55)
        learning_assess_path = (
            f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
            f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
            f"/groups/{learning_group_name}/assessments/{learning_assess_name}"
            f"?api-version={AZURE_MIGRATE_API_VERSION}"
        )
        learning_body = _build_assessment_body()
        azure_arm_request("PUT", learning_assess_path, token, learning_body)
        learning_assessment = wait_for_assessment_complete(
            subscription_id, resource_group, project_name,
            learning_group_name, learning_assess_name, token, progress,
        )
        progress("  ✅ 学习评估完成，正在分析各服务器单机成本...")

        learning_assessed = azure_arm_list(
            f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
            f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
            f"/groups/{learning_group_name}/assessments/{learning_assess_name}/assessedMachines"
            f"?api-version={AZURE_MIGRATE_API_VERSION}",
            token,
        )

        cache = learn_tier_machine_selections(learning_assessed, safe_prefix, progress)
        save_tier_cache(cache)
        progress("  ✅ 所有规模学习完成，结果已缓存到本地。")
    else:
        tier_info = cache["tiers"][str(tier)]
        progress(
            f"已命中学习缓存：规模 {_format_usd(float(tier))}，"
            f"选择 {tier_info['machine_count']} 台服务器，"
            f"预期年化 {_format_usd(tier_info['expected_annual'])}"
        )

    selected_ids = get_machine_ids_for_tier(tier, machines, safe_prefix, cache)
    if not selected_ids:
        selected_ids = all_machine_ids
    progress(f"当前规模选定 {len(selected_ids)}/{len(all_machine_ids)} 台服务器进入最终评估")

    # ── Step 13: 创建最终评估组并通过 updateMachines 加入选定服务器 ──
    progress("创建最终评估组...")
    group_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/groups/{group_name}?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    existing_group = _try_get_existing_resource(group_path, token)
    if existing_group:
        existing_group_type = str(existing_group.get("properties", {}).get("groupType", "")).strip().lower()
        if existing_group_type and existing_group_type != "import":
            # 同名组如果不是 Import 类型，groupType 无法修改，只能改名新建。
            group_name = _safe_azure_name(group_name, "poe-group", f"-imp-{int(time.time())}", 55)
            group_path = (
                f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
                f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
                f"/groups/{group_name}?api-version={AZURE_MIGRATE_API_VERSION}"
            )
            existing_group = None
            progress(f"  ⚠️ 发现同名评估组类型为 {existing_group_type}，改用新组名: {group_name}")

    if not existing_group:
        azure_arm_request("PUT", group_path, token, {
            "properties": {"groupType": "Import"},
            "eTag": "",
        })
        progress(f"  ✅ 已创建 Import 评估组: {group_name}")
    else:
        progress(f"  ✅ 复用评估组: {group_name}")

    update_machines_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/groups/{group_name}/updateMachines?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    azure_arm_request("POST", update_machines_path, token, {
        "eTag": "*",
        "properties": {
            "operationType": "Add",
            "machines": selected_ids,
        },
    })
    progress(f"  ✅ 已向评估组添加服务器: {len(selected_ids)} 台")

    group_payload = wait_for_group_machine_membership(
        group_path,
        token,
        expected_machine_count=len(selected_ids),
        progress=progress,
    )
    supported_types = {
        str(item).strip()
        for item in (group_payload.get("properties", {}).get("supportedAssessmentTypes") or [])
        if str(item).strip()
    }
    if supported_types and "MachineAssessment" not in supported_types:
        progress(
            "  ℹ️ 评估组尚未显式返回 MachineAssessment；"
            "继续按服务器评估类型创建 Azure VM 评估。"
        )

    # ── Step 14: 创建最终评估 ──
    progress("创建 Azure Migrate 最终评估...")
    assessment_path = (
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/groups/{group_name}/assessments/{assessment_resource_name}"
        f"?api-version={AZURE_MIGRATE_API_VERSION}"
    )
    assessment_body = _build_assessment_body(target_location=target_location)
    if target_location:
        progress(f"  ℹ️ 评估目标区域已设为：{target_location}（与解决方案架构一致）")
    if annual_budget is not None:
        progress(f"已解析用户预估年消耗：{_format_usd(annual_budget)}")
    else:
        progress("未解析到有效预估年消耗；评估会按 Portal 默认设置创建。")
    azure_arm_request("PUT", assessment_path, token, assessment_body)
    assessment = wait_for_assessment_complete(
        subscription_id, resource_group, project_name, group_name, assessment_resource_name, token, progress
    )
    assessment, tuning_history, budget_target_met = tune_assessment_to_budget(
        subscription_id=subscription_id,
        resource_group=resource_group,
        project_name=project_name,
        group_name=group_name,
        assessment_name=assessment_resource_name,
        assessment_path=assessment_path,
        assessment_body=assessment_body,
        assessment=assessment,
        annual_budget=annual_budget,
        token=token,
        progress=progress,
    )
    try:
        refresh_migrate_project_summary(subscription_id, resource_group, project_name, token)
    except Exception:
        pass

    progress("读取评估结果并下载 Azure Migrate Portal 同源 Excel 报告...")
    assessed_machines = azure_arm_list(
        f"/subscriptions/{subscription_id}/resourceGroups/{resource_group}"
        f"/providers/Microsoft.Migrate/assessmentProjects/{project_name}"
        f"/groups/{group_name}/assessments/{assessment_resource_name}/assessedMachines"
        f"?api-version={AZURE_MIGRATE_API_VERSION}",
        token,
    )
    excel_bytes = download_assessment_report(
        subscription_id, resource_group, project_name, group_name, assessment_resource_name, token
    )
    progress(f"  ✅ Azure Migrate 导出报告已下载（{len(excel_bytes)} 字节）")
    return {
        "project_name": project_name,
        "site_name": site_name,
        "collector_name": collector_name,
        "group_name": group_name,
        "assessment_name": assessment_resource_name,
        "portal_inventory_count": portal_inventory_count,
        "migrate_project_id": migrate_project_id,
        "assessment_project_id": assessment_project_id,
        "import_site_id": import_site_id,
        "assessment": assessment,
        "assessed_machines": assessed_machines,
        "excel_bytes": excel_bytes,
        "budget_target": annual_budget,
        "monthly_cost": assessment_monthly_total_cost(assessment),
        "annualized_cost": assessment_monthly_total_cost(assessment) * 12,
        "budget_target_met": budget_target_met,
        "tuning_history": tuning_history,
        "tier": tier,
        "selected_machine_count": len(selected_ids),
        "total_machine_count": len(all_machine_ids),
    }


def fix_assessment_excel_timestamps(
    excel_bytes: bytes,
    pov_start: datetime.date,
    pov_end: datetime.date,
) -> bytes:
    """
    修改 Assessment Excel 报告中的时间戳，使其位于用户输入的 POV 时间区间内。
    - Assessment_Summary sheet: "Created on (UTC)" 列 — 纯文本格式 M/D/YYYY H:MM:SS AM
    - Assessment_Properties sheet: "Performance history start time" = Created on (UTC)
      "Performance history end time" = start + 1 天，纯文本格式，不改时分秒
    """
    import random
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(excel_bytes))

    # 在 POV 区间内随机选一天作为评估创建日
    total_days = (pov_end - pov_start).days
    if total_days <= 0:
        total_days = 1
    random_day_offset = random.randint(1, max(total_days - 1, 1))
    created_date = pov_start + datetime.timedelta(days=random_day_offset)

    def _format_as_text(dt_val: datetime.datetime) -> str:
        """格式化为 M/D/YYYY H:MM:SS AM/PM 纯文本。"""
        hour = dt_val.hour
        ampm = "AM" if hour < 12 else "PM"
        hour_12 = hour % 12
        if hour_12 == 0:
            hour_12 = 12
        return f"{dt_val.month}/{dt_val.day}/{dt_val.year} {hour_12}:{dt_val.minute:02d}:{dt_val.second:02d} {ampm}"

    def _parse_time_from_value(orig):
        """从原始单元格值中提取时分秒。"""
        if isinstance(orig, datetime.datetime):
            return orig.hour, orig.minute, orig.second
        # 尝试从文本中解析时间部分
        text = str(orig).strip()
        import re
        m = re.search(r'(\d{1,2}):(\d{2}):(\d{2})\s*(AM|PM)?', text, re.IGNORECASE)
        if m:
            h, mi, s = int(m.group(1)), int(m.group(2)), int(m.group(3))
            ampm = (m.group(4) or "").upper()
            if ampm == "PM" and h != 12:
                h += 12
            elif ampm == "AM" and h == 12:
                h = 0
            return h, mi, s
        return 2, 35, 35  # 默认时间

    # 用于记录 Created on (UTC) 最终时间，供 Performance history start 使用
    created_datetime = None

    # ── 修改 Assessment_Summary sheet ──
    if "Assessment_Summary" in wb.sheetnames:
        ws = wb["Assessment_Summary"]
        header_row = 1
        created_col = None
        for col in range(1, ws.max_column + 1):
            cell_val = ws.cell(row=header_row, column=col).value
            if cell_val and "created on" in str(cell_val).lower():
                created_col = col
                break
        if created_col:
            for row in range(2, ws.max_row + 1):
                orig = ws.cell(row=row, column=created_col).value
                if orig is not None:
                    h, mi, s = _parse_time_from_value(orig)
                    created_datetime = datetime.datetime(
                        created_date.year, created_date.month, created_date.day, h, mi, s
                    )
                    ws.cell(row=row, column=created_col).value = _format_as_text(created_datetime)

    # 如果没从 Summary 里获取到，使用默认时间
    if created_datetime is None:
        created_datetime = datetime.datetime(
            created_date.year, created_date.month, created_date.day, 2, 35, 35
        )

    # ── 修改 Assessment_Properties sheet ──
    # Performance history start = Created on (UTC)
    # Performance history end = start + 1天
    perf_end_datetime = created_datetime + datetime.timedelta(days=1)

    if "Assessment_Properties" in wb.sheetnames:
        ws = wb["Assessment_Properties"]
        header_row = 1
        prop_col = None
        val_col = None
        for col in range(1, min(ws.max_column + 1, 20)):
            cell_val = ws.cell(row=header_row, column=col).value
            if cell_val:
                lower_val = str(cell_val).lower()
                if "property" in lower_val or "name" in lower_val or "key" in lower_val:
                    prop_col = col
                elif "value" in lower_val:
                    val_col = col

        if prop_col and val_col:
            for row in range(2, ws.max_row + 1):
                prop_name = str(ws.cell(row=row, column=prop_col).value or "").lower()
                if "performance history start" in prop_name:
                    ws.cell(row=row, column=val_col).value = _format_as_text(created_datetime)
                elif "performance history end" in prop_name:
                    ws.cell(row=row, column=val_col).value = _format_as_text(perf_end_datetime)
        else:
            start_col = None
            end_col = None
            for col in range(1, min(ws.max_column + 1, 50)):
                cell_val = ws.cell(row=header_row, column=col).value
                if cell_val:
                    lower_val = str(cell_val).lower()
                    if "performance history start" in lower_val:
                        start_col = col
                    elif "performance history end" in lower_val:
                        end_col = col
            if start_col:
                for row in range(2, ws.max_row + 1):
                    if ws.cell(row=row, column=start_col).value is not None:
                        ws.cell(row=row, column=start_col).value = _format_as_text(created_datetime)
            if end_col:
                for row in range(2, ws.max_row + 1):
                    if ws.cell(row=row, column=end_col).value is not None:
                        ws.cell(row=row, column=end_col).value = _format_as_text(perf_end_datetime)

    # 保存
    out_buf = io.BytesIO()
    wb.save(out_buf)
    return out_buf.getvalue()


def create_poe_zip(artifacts: List[Dict[str, Any]]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for artifact in artifacts:
            zip_file.writestr(artifact["file_name"], artifact["bytes"])
    return buffer.getvalue()


def get_existing_solution_text(current_doc_type: str) -> Optional[str]:
    key = "solution_text" if current_doc_type == "AI" else "infra_text"
    text = st.session_state.get(key)
    if isinstance(text, str) and text.strip():
        return text.strip()
    return None


def create_solution_artifact_from_text(
    current_doc_type: str,
    content: str,
    customer_name: str,
    account_name: str,
) -> Dict[str, Any]:
    if current_doc_type == "AI":
        docx_bytes = create_solution_docx(content=content, customer_name=customer_name)
        file_name = f"{account_name}-Solution Architecture.docx"
    else:
        docx_bytes = create_infra_docx(content=content, customer_name=customer_name)
        file_name = f"{account_name}-Infra Solution Architecture.docx"
    return {"content": content, "bytes": docx_bytes, "file_name": file_name}


def create_pov_artifact_from_text(
    content: str,
    customer_name: str,
    account_name: str,
) -> Dict[str, Any]:
    docx_bytes = create_pov_docx(content=content, customer_name=customer_name)
    return {
        "content": content,
        "bytes": docx_bytes,
        "file_name": f"{account_name}-PostAssessment POVdeployment.docx",
    }


def get_generated_migrate_csv_text() -> Optional[str]:
    csv_text = st.session_state.get("csv_code")
    if isinstance(csv_text, str) and csv_text.strip():
        return csv_text.strip()
    return None


def resolve_auto_inventory_csv(uploaded_inventory: Any) -> tuple[Optional[bytes], str]:
    if uploaded_inventory is not None:
        return uploaded_inventory.getvalue(), "手动上传 CSV"
    generated_csv = get_generated_migrate_csv_text()
    if generated_csv:
        return generated_csv.encode("utf-8-sig"), "Azure Migrate CSV 标签页生成"
    return None, "未提供"


def format_auto_poe_log(message: str) -> str:
    text = str(message or "").strip()
    if not text:
        return ""
    # 保留 Markdown 格式行（标题、颜色标注等）原样输出
    if text.startswith("**") or text.startswith("---") or ":green[" in text or ":red[" in text or ":orange[" in text:
        return text
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"^[^\w\u4e00-\u9fff]+", "", text).strip()
    text = text.removesuffix("...").strip()
    text = re.sub(r"https?://\S+", "[url]", text)
    text = re.sub(r"/subscriptions/[^\s；;，,]+", lambda match: match.group(0).rstrip("/").split("/")[-1], text, flags=re.I)
    if not text:
        return ""
    return text


def should_display_auto_poe_log(message: str) -> bool:
    text = str(message or "").strip()
    if not text:
        return False
    # Markdown 格式行（分模块标题、带颜色的结果行）始终显示
    if text.startswith("**") or text.startswith("---") or ":green[" in text or ":red[" in text or ":orange[" in text:
        return True
    interim_markers = [
        "正在",
        "开始",
        "等待",
        "仍在",
        "继续",
        "第 ",
        "注册 ",
        "读取评估结果",
    ]
    result_markers = [
        "成功",
        "完成",
        "已",
        "复用",
        "解析",
        "目标区间",
        "不在目标区间",
        "未填写",
        "未解析",
        "请到 Azure Portal",
        "学习",
        "规模",
        "选定",
        "命中",
        "缓存",
        "前缀",
    ]
    if any(marker in text for marker in interim_markers) and not any(marker in text for marker in result_markers):
        return False
    return any(marker in text for marker in result_markers)


def run_full_auto_poe(
    current_doc_type: str,
    customer_name: str,
    account_name: str,
    customer_bg: str,
    solution_ref: str,
    infra_ref: str,
    pov_ref: str,
    token: str,
    subscription_id: str,
    resource_group: str,
    assessment_name: str,
    annual_budget_text: Optional[str],
    pov_start: Optional[datetime.date],
    pov_end: Optional[datetime.date],
    vendor_team: str,
    existing_solution_text: Optional[str],
    existing_pov_text: Optional[str],
    progress: Callable[[str], None],
) -> Dict[str, Any]:
    progress("---")
    progress("**一、AI 解决方案架构文档**")
    if existing_solution_text and existing_solution_text.strip():
        solution_artifact = create_solution_artifact_from_text(
            current_doc_type,
            existing_solution_text.strip(),
            customer_name,
            account_name,
        )
        progress(f"  :green[✅ 已复用] `{solution_artifact['file_name']}`")
    else:
        solution_artifact = generate_solution_artifact(
            current_doc_type, customer_name, account_name, customer_bg, solution_ref, infra_ref
        )
        progress(f"  :green[✅ 已生成] `{solution_artifact['file_name']}`")

    target_key = "solution_text" if current_doc_type == "AI" else "infra_text"
    st.session_state[target_key] = solution_artifact["content"]
    st.session_state["customer_name"] = customer_name
    st.session_state["account_name"] = account_name

    progress("")
    progress("**二、POV 部署计划文档**")
    if existing_pov_text and existing_pov_text.strip():
        pov_artifact = create_pov_artifact_from_text(
            existing_pov_text.strip(),
            customer_name,
            account_name,
        )
        progress(f"  :green[✅ 已复用] `{pov_artifact['file_name']}`")
    else:
        if not pov_start or not pov_end:
            raise RuntimeError("缺少 POV 开始日期或结束日期，无法生成 POV 部署计划。")
        if not has_meaningful_pov_team(vendor_team):
            raise RuntimeError("缺少 POV 项目人员，无法生成 POV 部署计划。")
        pov_artifact = generate_pov_artifact(
            solution_artifact["content"],
            customer_name,
            account_name,
            pov_ref,
            pov_start,
            pov_end,
            vendor_team,
        )
        progress(f"  :green[✅ 已生成] `{pov_artifact['file_name']}`")
    st.session_state["pov_text"] = pov_artifact["content"]
    st.session_state["pov_source_doc_type"] = current_doc_type

    # ── Step 2.5: Pricing Calculator 自动导出 ──
    pricing_artifact = None
    pricing_result = None
    if HAS_PRICING_AUTOMATION:
        progress("")
        progress("**三、Azure 价格估算表**")
        progress("  正在自动化 Azure Pricing Calculator 导出...")
        try:
            annual_budget_val = parse_annual_budget_usd(annual_budget_text) or 0.0
            # 计算预算上限（下一个档位），防止超标
            budget_cap = 0.0
            if annual_budget_val > 0:
                for tier in BUDGET_TIERS:
                    if tier > annual_budget_val:
                        budget_cap = tier
                        break
                if budget_cap == 0:
                    budget_cap = annual_budget_val * 1.5  # 最高档位用 1.5 倍封顶
            pricing_result = run_pricing_export(
                solution_text=solution_artifact["content"],
                annual_budget=annual_budget_val,
                account_name=account_name,
                progress=progress,
                headed=False,
                budget_cap=budget_cap,
            )
            if pricing_result.xlsx_bytes:
                pricing_artifact = {
                    "file_name": f"{account_name}-Azure calculator.xlsx",
                    "bytes": pricing_result.xlsx_bytes,
                }
                progress(f"  :green[✅ 已导出] `{pricing_artifact['file_name']}`")
                if pricing_result.fallbacks:
                    fallback_msgs = [f"{fb.original_service}({fb.original_sku})→{fb.resolved_sku}" for fb in pricing_result.fallbacks if fb.resolved_service != "(跳过)"]
                    if fallback_msgs:
                        progress(f"  ⚠️ 资源降级: {'; '.join(fallback_msgs)}")
            else:
                progress(f"  ⚠️ 价格导出未成功: {pricing_result.error or '未知错误'}")
        except Exception as e:
            progress(f"  ⚠️ 价格导出异常（不影响其他文档）: {e}")

    # 从解决方案文档中提取主要区域，用于评估目标地区
    target_location = _extract_dominant_region(solution_artifact["content"])

    progress("")
    progress("**四、Azure Migrate 评估报告**")
    migrate_result = run_azure_migrate_assessment(
        token, subscription_id, resource_group, account_name, assessment_name, annual_budget_text, progress,
        target_location=target_location,
    )
    # 修改 Excel 中的时间戳，使其位于用户 POV 时间区间内
    excel_bytes = migrate_result["excel_bytes"]
    if pov_start and pov_end:
        try:
            excel_bytes = fix_assessment_excel_timestamps(excel_bytes, pov_start, pov_end)
            progress("  ✅ 已修正评估报告时间戳至 POV 区间内")
        except Exception:
            pass  # 时间戳修改失败不阻塞主流程
    assessment_artifact = {
        "file_name": f"{account_name}-Azure Migrate Assessment.xlsx",
        "bytes": excel_bytes,
    }
    progress(f"  :green[✅ 已生成] `{assessment_artifact['file_name']}`")

    # 打包所有交付物
    all_artifacts = [solution_artifact, pov_artifact, assessment_artifact]
    if pricing_artifact:
        all_artifacts.append(pricing_artifact)
    zip_bytes = create_poe_zip(all_artifacts)
    progress(f"成功生成 POE 套件：{account_name}-POE-Complete.zip（{len(all_artifacts)} 个文件）")
    return {
        "zip_bytes": zip_bytes,
        "zip_name": f"{account_name}-POE-Complete.zip",
        "solution": solution_artifact,
        "pov": pov_artifact,
        "assessment": assessment_artifact,
        "pricing": pricing_artifact,
        "pricing_result": pricing_result,
        "migrate": migrate_result,
    }


def render_full_auto_poe_area(
    current_doc_type: str,
    account_name: str,
    customer_name: str,
    budget: str,
    customer_bg: str,
    solution_ref: str,
    infra_ref: str,
    pov_ref: str,
) -> None:
    azure_logged_in = _is_azure_token_valid()
    selected_subscription = st.session_state.get("_cached_subscription")
    selected_resource_group = st.session_state.get("_cached_resource_group")
    resolved_customer_name = customer_name.strip() or str(st.session_state.get("customer_name") or "").strip()
    resolved_account_name = (
        account_name.strip()
        or str(st.session_state.get("account_name") or "").strip()
        or resolved_customer_name
    )
    resolved_budget_text = budget.strip() or str(st.session_state.get("budget") or "").strip()
    default_assessment_name = _safe_azure_name(resolved_account_name or resolved_customer_name, "poe", "assessment", 55)
    assessment_name = st.session_state.get("auto_assessment_name", default_assessment_name)
    budget_value = parse_annual_budget_usd(resolved_budget_text)
    pov_start = st.session_state.get("pov_start_date")
    pov_end = st.session_state.get("pov_end_date")
    vendor_team = st.session_state.get("pov_vendor_team", "")
    existing_solution_text = get_existing_solution_text(current_doc_type)
    existing_pov_text = st.session_state.get("pov_text")
    existing_pov_text = existing_pov_text.strip() if isinstance(existing_pov_text, str) else None
    pov_source_doc_type = st.session_state.get("pov_source_doc_type")
    if existing_pov_text and pov_source_doc_type and pov_source_doc_type != current_doc_type:
        existing_pov_text = None
    builtin_csv_ok = os.path.exists(BUILTIN_CSV_PATH)
    matched_tier = snap_budget_to_tier(budget_value) if budget_value and budget_value > 0 else None
    tier_cache = load_tier_cache()
    tier_learned = bool(tier_cache.get("tiers", {}).get(str(matched_tier))) if matched_tier else False
    has_existing_solution = bool(existing_solution_text)
    has_existing_pov = bool(existing_pov_text)
    pov_dates_filled = bool(pov_start and pov_end)
    pov_team_ready = has_meaningful_pov_team(vendor_team)
    solution_ready = has_existing_solution or (bool(resolved_customer_name) and bool(customer_bg.strip()))
    pov_ready = has_existing_pov or (pov_dates_filled and pov_team_ready)
    status_slot = st.empty()

    token = st.session_state.get("azure_token")

    # ── 就绪状态 + 步骤 ──
    with status_slot.container():
        render_workflow_steps([
            {"title": "登录 Azure", "state": "done" if azure_logged_in else "ready"},
            {"title": "选择目标", "state": "done" if selected_subscription and selected_resource_group else ("ready" if azure_logged_in else "blocked")},
            {"title": "内置模板", "state": "done" if builtin_csv_ok else "blocked"},
            {"title": "生成套件", "state": "ready" if False else "blocked"},
        ])
        readiness_items = [
            ("客户名称", bool(resolved_customer_name), "用于文档标题和输出文件名"),
            ("方案文档", solution_ready, "已生成则复用；未生成时会按客户背景生成"),
            ("预估年消耗", budget_value is not None and budget_value > 0, "用于校准迁移评估年化估算和规模匹配"),
            ("POV 文档/输入", pov_ready, "已生成则复用；未生成时需填写时间区间和项目人员"),
            ("Azure 登录", azure_logged_in, "用于创建 Azure Migrate 项目和评估"),
            ("订阅与资源组", bool(selected_subscription and selected_resource_group), "评估资源会创建到这里"),
            ("内置服务器模板", builtin_csv_ok, "内置 Azurecsvtemplate.csv 模板"),
            ("评估名称", bool(assessment_name.strip()), "用于 Azure Migrate 评估资源"),
        ]
        render_readiness(readiness_items)

    # ── Azure 登录 + 订阅/资源组 — 单行 ──
    az_c1, az_c2, az_c3 = st.columns([1, 1.5, 1.5])
    with az_c1:
        if azure_logged_in:
            st.success(f"已登录：{st.session_state.get('azure_user', 'Azure 用户')}", icon="✅")
            if st.button("退出", use_container_width=True, key="btn_azure_logout"):
                clear_azure_login()
                st.rerun()
        else:
            if st.button("登录 Azure", type="primary", use_container_width=True, key="btn_azure_login"):
                try:
                    msal_device_code_login()
                    st.rerun()
                except Exception as exc:
                    st.error(f"登录失败：{exc}")
    with az_c2:
        if azure_logged_in:
            try:
                subscriptions = list_azure_subscriptions(token)
                if subscriptions:
                    subscription_labels = [_subscription_label(sub) for sub in subscriptions]
                    default_index = next(
                        (
                            idx for idx, sub in enumerate(subscriptions)
                            if "Microsoft" in (sub.get("displayName") or "")
                            and ("合作伙伴" in (sub.get("displayName") or "") or "Partner" in (sub.get("displayName") or ""))
                        ),
                        0,
                    )
                    sub_label = st.selectbox("订阅", subscription_labels, index=default_index, key="auto_subscription_select")
                    selected_subscription = subscriptions[subscription_labels.index(sub_label)]
                    subscription_id = selected_subscription.get("subscriptionId")
                    st.session_state["azure_subscription_id"] = subscription_id
                    st.session_state["azure_subscription_name"] = selected_subscription.get("displayName", "")
                    st.session_state["_cached_subscription"] = selected_subscription
                    persist_session_state()
            except Exception as exc:
                st.error(f"读取订阅失败：{exc}")
        else:
            st.selectbox("订阅", ["登录后选择"], disabled=True, key="_sub_placeholder")
    with az_c3:
        if azure_logged_in and selected_subscription:
            try:
                groups = list_azure_resource_groups(subscription_id, token)
                if groups:
                    group_labels = [_resource_group_label(group) for group in groups]
                    rg_label = st.selectbox("资源组", group_labels, key="auto_resource_group_select")
                    selected_resource_group = groups[group_labels.index(rg_label)]
                    st.session_state["azure_resource_group"] = selected_resource_group.get("name")
                    st.session_state["_cached_resource_group"] = selected_resource_group
                    persist_session_state()
                else:
                    st.warning("当前订阅下没有资源组。")
            except Exception as exc:
                st.error(f"读取资源组失败：{exc}")
        else:
            st.selectbox("资源组", ["登录后选择"], disabled=True, key="_rg_placeholder")

    # ── 评估名称 + 模板状态 — 单行 ──
    name_c1, name_c2 = st.columns([1, 2])
    with name_c1:
        assessment_name = st.text_input("评估名称", value=default_assessment_name, key="auto_assessment_name")
    with name_c2:
        if builtin_csv_ok:
            template_names = _get_template_machine_names()
            if matched_tier:
                tier_label = _format_usd(float(matched_tier))
                if tier_learned:
                    learned_count = tier_cache["tiers"][str(matched_tier)]["machine_count"]
                    st.caption(f"内置模板 {len(template_names)} 台 | 规模 {tier_label} | 已学习 → {learned_count} 台")
                else:
                    st.caption(f"内置模板 {len(template_names)} 台 | 规模 {tier_label} | 首次运行将自动学习")
            else:
                st.caption(f"内置模板 {len(template_names)} 台 | 请填写预估年消耗匹配规模")
        else:
            st.error("Azurecsvtemplate.csv 未找到")

    # ── 浏览器 Profile 状态（价格计算器自动化）──
    if HAS_PRICING_AUTOMATION:
        browser_ready = is_browser_profile_ready()
        browser_col1, browser_col2 = st.columns([2, 1])
        with browser_col1:
            if browser_ready:
                st.caption("🌐 价格计算器浏览器 Profile 已就绪")
            else:
                st.caption("🌐 价格计算器浏览器 Profile 未初始化（首次需手动登录）")
        with browser_col2:
            if st.button("初始化浏览器" if not browser_ready else "重新登录", key="btn_init_browser", use_container_width=True):
                with st.spinner("正在打开浏览器，请在弹出窗口中登录 Azure..."):
                    try:
                        from pricing_automation import initialize_browser_profile
                        initialize_browser_profile()
                        st.success("✅ 浏览器 Profile 已保存")
                        st.rerun()
                    except Exception as e:
                        st.error(f"浏览器初始化失败: {e}")

    workflow_ready = all(item[1] for item in readiness_items)

    if st.button(
        "生成完整 POE 套件",
        type="primary",
        use_container_width=True,
        key="btn_full_auto_poe",
        disabled=not workflow_ready,
    ):
        cust = resolved_customer_name
        acct = resolved_account_name or cust
        if not cust:
            st.warning("请输入客户名称。")
            return
        if not solution_ready:
            st.warning("请先在「解决方案文档」标签页生成/导入方案文档，或填写客户背景以便自动生成。")
            return
        if budget_value is None or budget_value <= 0:
            st.warning("请输入可解析的预估年消耗，例如 500k、50万 或 500000。")
            return
        if not has_existing_pov and not pov_dates_filled:
            st.warning("请在上方公共输入区域填写 POV 开始日期和结束日期。")
            return
        if not has_existing_pov and pov_dates_filled and pov_end < pov_start:
            st.warning("POV 结束日期不能早于开始日期，请修正。")
            return
        if not has_existing_pov and not pov_team_ready:
            st.warning("请在上方公共输入区域填写乙方项目人员。")
            return
        if not _is_azure_token_valid():
            st.warning("请先登录 Microsoft Azure 账户。")
            return
        if not selected_subscription or not selected_resource_group:
            st.warning("请选择订阅和资源组。")
            return
        if not builtin_csv_ok:
            st.warning("内置服务器清单模板 Azurecsvtemplate.csv 未找到。")
            return
        if not assessment_name.strip():
            st.warning("请输入评估名称。")
            return

        try:
            st.session_state["auto_poe_running"] = True
            with st.status("正在生成完整 POE 套件", expanded=True) as status:
                stop_placeholder = st.empty()
                log_placeholder = st.empty()
                log_lines: List[str] = []
                stop_placeholder.button(
                    "停止生成", type="secondary", use_container_width=True, key="btn_stop_auto_poe",
                    on_click=lambda: st.session_state.update({"auto_poe_stop": True}),
                )

                def progress(message: str) -> None:
                    if st.session_state.get("auto_poe_stop"):
                        st.session_state.pop("auto_poe_stop", None)
                        st.session_state.pop("auto_poe_running", None)
                        raise RuntimeError("用户已手动停止生成。")
                    formatted = format_auto_poe_log(message)
                    if formatted and should_display_auto_poe_log(formatted):
                        log_lines.append(formatted)
                        # 使用 markdown 渲染，支持加粗、颜色等格式
                        log_placeholder.markdown("\n\n".join(log_lines[-120:]))

                result = run_full_auto_poe(
                    current_doc_type=current_doc_type,
                    customer_name=cust,
                    account_name=acct,
                    customer_bg=customer_bg.strip(),
                    solution_ref=solution_ref,
                    infra_ref=infra_ref,
                    pov_ref=pov_ref,
                    token=st.session_state["azure_token"],
                    subscription_id=selected_subscription["subscriptionId"],
                    resource_group=selected_resource_group["name"],
                    assessment_name=assessment_name.strip(),
                    annual_budget_text=resolved_budget_text,
                    pov_start=pov_start,
                    pov_end=pov_end,
                    vendor_team=vendor_team,
                    existing_solution_text=existing_solution_text,
                    existing_pov_text=existing_pov_text,
                    progress=progress,
                )
                stop_placeholder.empty()
                st.session_state.pop("auto_poe_running", None)
                st.session_state["auto_poe_zip_bytes"] = result["zip_bytes"]
                st.session_state["auto_poe_zip_name"] = result["zip_name"]
                if result["solution"].get("svg_code"):
                    st.session_state["svg_code"] = result["solution"]["svg_code"]
                st.session_state["auto_poe_result"] = {
                    "customer_name": cust,
                    "zip_name": result["zip_name"],
                    "solution_file_name": result["solution"]["file_name"],
                    "pov_file_name": result["pov"]["file_name"],
                    "assessment_file_name": result["assessment"]["file_name"],
                    "pricing_file_name": result["pricing"]["file_name"] if result.get("pricing") else None,
                    "pricing_fallbacks": [
                        f"{fb.original_service}({fb.original_sku})→{fb.resolved_sku}"
                        for fb in (result.get("pricing_result") or PricingExportResult()).fallbacks
                        if fb.resolved_service != "(跳过)"
                    ] if HAS_PRICING_AUTOMATION else [],
                    "project_name": result["migrate"]["project_name"],
                    "assessment_name": result["migrate"]["assessment_name"],
                    "machine_count": len(result["migrate"].get("assessed_machines", [])),
                    "portal_inventory_count": result["migrate"].get("portal_inventory_count", 0),
                    "annualized_cost": result["migrate"].get("annualized_cost"),
                    "budget_target": result["migrate"].get("budget_target"),
                    "budget_target_met": result["migrate"].get("budget_target_met", True),
                    "csv_source": "内置模板（Azurecsvtemplate.csv）",
                    "tier": result["migrate"].get("tier"),
                    "selected_machine_count": result["migrate"].get("selected_machine_count"),
                    "total_machine_count": result["migrate"].get("total_machine_count"),
                }
                status.update(label="完整 POE 套件生成完成", state="complete")
                st.session_state["auto_poe_finish_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        except Exception as exc:
            st.session_state.pop("auto_poe_running", None)
            st.session_state["auto_poe_error"] = str(exc)
            st.error(f"全自动生成失败：{exc}")

    if "auto_poe_zip_bytes" in st.session_state:
        result = st.session_state.get("auto_poe_result", {})
        annualized_cost = result.get("annualized_cost")
        budget_target = result.get("budget_target")
        generated_items = [
            ("解决方案架构文档", result.get("solution_file_name") or "-"),
            ("POV 文档", result.get("pov_file_name") or "-"),
            ("迁移评估文档", result.get("assessment_file_name") or "-"),
        ]
        if result.get("pricing_file_name"):
            generated_items.append(("价格估算表", result.get("pricing_file_name")))
        render_auto_poe_result(
            customer_name=result.get("customer_name") or customer_name.strip() or account_name.strip() or "该客户",
            generated_items=generated_items,
            migrate_items=[
                ("Azure Migrate 项目", result.get("project_name") or "-"),
                ("迁移评估名称", result.get("assessment_name") or "-"),
                ("CSV 来源", result.get("csv_source") or "-"),
                ("规模档位", _format_usd(float(result["tier"])) if result.get("tier") else "-"),
                ("选定服务器", f"{result.get('selected_machine_count', 0)}/{result.get('total_machine_count', 0)} 台"),
                ("Portal 库存", f"{result.get('portal_inventory_count', 0)} 台"),
                ("评估服务器数", f"{result.get('machine_count', 0)} 台"),
                ("年化估算", _format_usd(annualized_cost)),
                ("用户预估", _format_usd(budget_target)),
            ],
        )
        if budget_target and not result.get("budget_target_met", True):
            st.warning("自动调整 3 轮后仍未落入用户预估年消耗的档次区间，请在 Azure Portal 的评估设置中手动调整。")
        pricing_fallbacks = result.get("pricing_fallbacks", [])
        if pricing_fallbacks:
            st.warning(f"价格计算器资源降级: {'; '.join(pricing_fallbacks)}")
        finish_time = st.session_state.get("auto_poe_finish_time")
        if finish_time:
            st.info(f"✅ 任务完成时间：{finish_time}")
        st.download_button(
            label="下载全部 POE 文档 (.zip)",
            data=st.session_state["auto_poe_zip_bytes"],
            file_name=st.session_state["auto_poe_zip_name"],
            mime="application/zip",
            use_container_width=True,
            key="dl_auto_poe_zip",
        )


# ──────────────────────────────────────────────
# 主界面
# ──────────────────────────────────────────────
def main():
    restore_session_state()
    render_app_header()

    if not check_secrets():
        st.stop()

    # 侧边栏
    with st.sidebar:
        st.markdown("### 操作")
        if st.button("清除所有结果", use_container_width=True):
            for key in [
                "solution_text", "infra_text", "pov_text", "customer_name", "account_name", "csv_code",
                "budget", "doc_type", "pov_source_doc_type", "yearly_excel_bytes", "yearly_excel_name", "yearly_messages",
                "auto_poe_zip_bytes", "auto_poe_zip_name", "auto_poe_result", "auto_poe_error",
                "pov_vendor_team",
            ]:
                st.session_state.pop(key, None)
            clear_session_persist()
            st.rerun()

        st.markdown("---")
        st.markdown("### 模板状态")
        sol_ok = os.path.exists(SOLUTION_TEMPLATE_PATH)
        infra_ok = os.path.exists(INFRA_TEMPLATE_PATH)
        pov_ok = os.path.exists(POV_TEMPLATE_PATH)
        csv_ok = os.path.exists(MIGRATE_TEMPLATE_PATH)
        render_template_status([
            ("AI Solution", sol_ok),
            ("Infra", infra_ok),
            ("POV", pov_ok),
            ("CSV", csv_ok),
        ])

    solution_ref = extract_template_text(SOLUTION_TEMPLATE_PATH) if sol_ok else ""
    infra_ref = extract_template_text(INFRA_TEMPLATE_PATH) if infra_ok else ""
    pov_ref = extract_template_text(POV_TEMPLATE_PATH) if pov_ok else ""

    # ════════════════════════════════════════════════════
    # 公共输入区域
    # ════════════════════════════════════════════════════
    render_section_head(
        "客户信息",
        "这些输入会贯穿方案文档、POV 计划、CSV 推导和 Azure Migrate 评估。",
        render_pill("必填项优先", "accent"),
    )
    c0, c1, c2 = st.columns([1.5, 2, 1])
    with c0:
        account_name = st.text_input("账户名", placeholder="例如：Tetherflow", help="用于生成下载文件名的前缀")
    with c1:
        customer_name = st.text_input("客户名称", placeholder="例如：宇宙无敌科技有限公司")
    with c2:
        budget = st.text_input("预估年消耗 (USD)", placeholder="例如：500k+")

    customer_bg = st.text_area(
        "客户背景信息",
        placeholder="粘贴客户背景资料，包括行业、规模、现有 IT 环境、核心需求和已知约束。",
        height=125,
    )

    pov_dc1, pov_dc2, pov_dc3, pov_dc4 = st.columns([1, 1, 1, 1])
    with pov_dc1:
        pov_start_input = st.date_input("POV 开始日期", value=None, key="pov_start_date")
    with pov_dc2:
        pov_end_input = st.date_input("POV 结束日期", value=None, key="pov_end_date")
    with pov_dc3:
        pov_tech_lead = st.text_input("技术负责人", key="_pov_tech_lead", help="填写我方技术负责人姓名")
    with pov_dc4:
        pov_architect = st.text_input("架构师", key="_pov_architect", help="填写我方架构师姓名")
    _parts = []
    if (pov_tech_lead or "").strip():
        _parts.append(f"技术负责人: {pov_tech_lead.strip()}")
    if (pov_architect or "").strip():
        _parts.append(f"架构师: {pov_architect.strip()}")
    st.session_state["pov_vendor_team"] = ", ".join(_parts)

    st.divider()

    # ════════════════════════════════════════════════════
    # Tab 布局
    # ════════════════════════════════════════════════════
    tab_auto, tab_sol, tab_pov, tab_csv, tab_yearly = st.tabs([
        "全自动POE生成", "解决方案文档", "POV 部署计划", "Azure Migrate CSV", "年度价格表"
    ])

    dp = _date_prefix()  # 日期前缀

    # ─────────── Tab 1: 全自动 POE 生成 ───────────
    with tab_auto:
        auto_doc_type_label = st.radio(
            "生成文档类型",
            ["AI 解决方案", "Infra 基础设施"],
            horizontal=True,
            key="auto_doc_type_radio",
        )
        auto_doc_type = "AI" if auto_doc_type_label == "AI 解决方案" else "Infra"
        render_full_auto_poe_area(
            current_doc_type=auto_doc_type,
            account_name=account_name,
            customer_name=customer_name,
            budget=budget,
            customer_bg=customer_bg,
            solution_ref=solution_ref,
            infra_ref=infra_ref,
            pov_ref=pov_ref,
        )

    # ─────────── Tab 2: 解决方案文档 ───────────
    with tab_sol:
        # 文档类型切换
        doc_type = st.radio(
            "选择文档类型",
            ["AI 解决方案", "Infra 基础设施"],
            horizontal=True,
            key="doc_type_radio",
            index=0 if st.session_state.get("doc_type", "AI") == "AI" else 1,
        )
        current_doc_type = "AI" if doc_type == "AI 解决方案" else "Infra"
        st.session_state["doc_type"] = current_doc_type

        # 文档来源切换
        doc_source = st.radio(
            "文档来源",
            ["AI 生成", "手动导入"],
            horizontal=True,
            key="doc_source_radio",
        )

        left, right = st.columns([1, 1])
        with left:
            if doc_source == "手动导入":
                # ── 手动导入：两步流程 ──
                # Step 1：上传 / 粘贴，确认后暂存原文
                # Step 2：AI 按模板格式重新生成
                if "imported_doc_text" not in st.session_state:
                    # ── Step 1：上传或粘贴 ──
                    uploaded_doc = st.file_uploader(
                        "上传已有的 .docx 文档",
                        type=["docx"],
                        key="upload_existing_doc",
                        help="上传后将自动提取文档文本内容",
                    )
                    manual_text = st.text_area(
                        "或直接粘贴文本内容",
                        height=200,
                        key="manual_doc_text",
                        placeholder="将已有的解决方案文档内容粘贴到此处...",
                    )

                    if st.button("确认导入", type="primary", use_container_width=True, key="btn_import"):
                        imported_text = ""
                        if uploaded_doc is not None:
                            doc = Document(uploaded_doc)
                            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                            for table in doc.tables:
                                for row in table.rows:
                                    cells = [cell.text.strip() for cell in row.cells]
                                    paragraphs.append(" | ".join(cells))
                            imported_text = "\n\n".join(paragraphs)
                        elif manual_text.strip():
                            imported_text = manual_text.strip()
                        else:
                            st.warning("请上传文档或粘贴文本。")
                            st.stop()

                        st.session_state["imported_doc_text"] = imported_text
                        # 同步写入 solution_text / infra_text，使 POV 等后续步骤可立即识别到文档
                        target_key = "solution_text" if current_doc_type == "AI" else "infra_text"
                        st.session_state[target_key] = imported_text
                        st.session_state["customer_name"] = customer_name.strip() if customer_name.strip() else "未命名客户"
                        st.session_state["account_name"] = account_name.strip() if account_name.strip() else (customer_name.strip() or "未命名客户")
                        st.session_state["budget"] = budget
                        st.session_state.pop("pov_text", None)
                        persist_session_state()
                        st.rerun()

                else:
                    # ── Step 2：确认内容 + AI 重新生成 ──
                    imported_text = st.session_state["imported_doc_text"]
                    st.success(f"文档已导入（共 {len(imported_text)} 字符）")
                    st.text_area(
                        "导入内容预览",
                        value=imported_text[:600] + "\n\n..." if len(imported_text) > 600 else imported_text,
                        height=160,
                        disabled=True,
                        key="preview_imported",
                    )

                    c_reimport, c_regen = st.columns(2)
                    with c_reimport:
                        if st.button("重新上传", use_container_width=True, key="btn_reimport"):
                            st.session_state.pop("imported_doc_text", None)
                            st.rerun()
                    with c_regen:
                        if st.button("AI 重新生成", type="primary", use_container_width=True, key="btn_regen_import"):
                            cust = customer_name.strip() or st.session_state.get("customer_name", "未命名客户")
                            system_prompt = SOLUTION_SYSTEM_PROMPT if current_doc_type == "AI" else INFRA_SYSTEM_PROMPT
                            ref_text = solution_ref if current_doc_type == "AI" else infra_ref
                            user_ctx = (
                                f"## 客户信息\n- **客户名称**：{cust}\n\n"
                            )
                            if customer_bg.strip():
                                user_ctx += (
                                    f"## 客户背景信息\n{customer_bg.strip()}\n\n"
                                )
                            user_ctx += (
                                f"## 已有解决方案文档（请基于以上客户信息和以下已有文档，按照要求的章节格式重新整理生成，不要照抄原文）\n\n"
                                f"{imported_text}"
                            )
                            if ref_text:
                                user_ctx += (
                                    f"\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
                                    f"{ref_text}"
                                )
                            try:
                                with st.spinner("正在基于导入内容 AI 重新生成..."):
                                    result_text = call_azure_openai(system_prompt, user_ctx)
                                    target_key = "solution_text" if current_doc_type == "AI" else "infra_text"
                                    st.session_state[target_key] = result_text
                                    st.session_state["customer_name"] = cust
                                    st.session_state["account_name"] = account_name.strip() if account_name.strip() else cust
                                    st.session_state["budget"] = budget
                                    st.session_state.pop("pov_text", None)
                                    st.session_state.pop("imported_doc_text", None)
                                persist_session_state()
                                st.rerun()
                            except Exception as e:
                                st.error(f"生成失败：{e}")

                    # 若已生成，显示下载按钮
                    target_key = "solution_text" if current_doc_type == "AI" else "infra_text"
                    if target_key in st.session_state:
                        customer = st.session_state["customer_name"]
                        acct = st.session_state.get("account_name") or account_name.strip() or customer
                        if current_doc_type == "AI":
                            docx_bytes = create_solution_docx(
                                content=st.session_state["solution_text"], customer_name=customer
                            )
                            st.download_button(
                                label="下载 AI 解决方案架构文档 (.docx)",
                                data=docx_bytes,
                                file_name=f"{acct}-Solution Architecture.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="dl_sol_import",
                            )
                        else:
                            docx_bytes = create_infra_docx(
                                content=st.session_state["infra_text"], customer_name=customer
                            )
                            st.download_button(
                                label="下载 Infra 基础设施架构文档 (.docx)",
                                data=docx_bytes,
                                file_name=f"{acct}-Infra Solution Architecture.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="dl_infra_import",
                            )

            else:
                # ── AI 生成文档 ──
                if current_doc_type == "AI":
                    # AI 解决方案文档逻辑
                    has_solution = "solution_text" in st.session_state
                    sol_label = "重新生成" if has_solution else "生成 AI 解决方案架构文档"
                    if st.button(sol_label, type="primary", use_container_width=True, key="btn_sol"):
                        if not customer_name.strip():
                            st.warning("请输入客户名称。")
                            st.stop()
                        if not customer_bg.strip():
                            st.warning("请输入客户背景信息。")
                            st.stop()
                        try:
                            with st.spinner("正在生成 AI 解决方案架构文档..."):
                                user_ctx = (
                                    f"## 客户信息\n- **客户名称**：{customer_name}\n\n"
                                    f"## 客户背景\n{customer_bg}"
                                )
                                if solution_ref:
                                    user_ctx += (
                                        f"\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
                                        f"{solution_ref}"
                                    )
                                sol_text = call_azure_openai(SOLUTION_SYSTEM_PROMPT, user_ctx)
                                st.session_state["solution_text"] = sol_text
                                st.session_state["customer_name"] = customer_name
                                st.session_state["account_name"] = account_name.strip() if account_name.strip() else customer_name
                                st.session_state["budget"] = budget
                                st.session_state.pop("pov_text", None)
                                st.session_state.pop("svg_code", None)
                            persist_session_state()
                            st.rerun()
                        except Exception as e:
                            st.error(f"生成失败：{e}")

                    if "solution_text" in st.session_state:
                        customer = st.session_state["customer_name"]
                        acct = st.session_state.get("account_name") or account_name.strip() or customer
                        docx_sol = create_solution_docx(
                            content=st.session_state["solution_text"], customer_name=customer
                        )
                        st.download_button(
                            label="下载 AI 解决方案架构文档 (.docx)",
                            data=docx_sol,
                            file_name=f"{acct}-Solution Architecture.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                else:
                    # Infra 基础设施文档逻辑
                    has_infra = "infra_text" in st.session_state
                    infra_label = "重新生成" if has_infra else "生成 Infra 基础设施架构文档"
                    if st.button(infra_label, type="primary", use_container_width=True, key="btn_infra"):
                        if not customer_name.strip():
                            st.warning("请输入客户名称。")
                            st.stop()
                        if not customer_bg.strip():
                            st.warning("请输入客户背景信息。")
                            st.stop()
                        try:
                            with st.spinner("正在生成 Infra 基础设施架构文档..."):
                                user_ctx = (
                                    f"## 客户信息\n- **客户名称**：{customer_name}\n\n"
                                    f"## 客户背景\n{customer_bg}"
                                )
                                if infra_ref:
                                    user_ctx += (
                                        f"\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
                                        f"{infra_ref}"
                                    )
                                infra_text = call_azure_openai(INFRA_SYSTEM_PROMPT, user_ctx)
                                st.session_state["infra_text"] = infra_text
                                st.session_state["customer_name"] = customer_name
                                st.session_state["account_name"] = account_name.strip() if account_name.strip() else customer_name
                                st.session_state["budget"] = budget
                                st.session_state.pop("pov_text", None)
                                st.session_state.pop("svg_code", None)
                            persist_session_state()
                            st.rerun()
                        except Exception as e:
                            st.error(f"生成失败：{e}")

                    if "infra_text" in st.session_state:
                        customer = st.session_state["customer_name"]
                        acct = st.session_state.get("account_name") or account_name.strip() or customer
                        docx_infra = create_infra_docx(
                            content=st.session_state["infra_text"], customer_name=customer
                        )
                        st.download_button(
                            label="下载 Infra 基础设施架构文档 (.docx)",
                            data=docx_infra,
                            file_name=f"{acct}-Infra Solution Architecture.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )

        with right:
            if current_doc_type == "AI":
                if "solution_text" in st.session_state:
                    st.markdown("**AI 解决方案文档预览**")
                    st.markdown(st.session_state["solution_text"], unsafe_allow_html=True)
                else:
                    st.info("请先生成或导入 AI 解决方案文档")
            else:
                if "infra_text" in st.session_state:
                    st.markdown("**Infra 基础设施文档预览**")
                    st.markdown(st.session_state["infra_text"], unsafe_allow_html=True)
                else:
                    st.info("请先生成或导入 Infra 基础设施文档")

    # ─────────── Tab 2: POV 部署计划 ───────────
    with tab_pov:
        # 根据当前文档类型确定使用哪个解决方案文档
        current_doc_type = st.session_state.get("doc_type", "AI")
        has_base_doc = ("solution_text" in st.session_state) if current_doc_type == "AI" else ("infra_text" in st.session_state)
        
        if not has_base_doc:
            doc_type_name = "AI 解决方案" if current_doc_type == "AI" else "Infra 基础设施"
            st.info(f"请先在「解决方案文档」标签页中生成或导入 {doc_type_name} 文档")
        else:
            customer = st.session_state["customer_name"]
            solution = st.session_state["solution_text"] if current_doc_type == "AI" else st.session_state["infra_text"]
            left, right = st.columns([1, 1])
            with left:
                st.caption(f"📄 当前基于: **{current_doc_type}** 解决方案文档")

                pov_start = st.session_state.get("pov_start_date")
                pov_end = st.session_state.get("pov_end_date")
                vendor_team = st.session_state.get("pov_vendor_team", "")

                if pov_start and pov_end:
                    st.info(f"POV 周期：{pov_start} ~ {pov_end}")
                else:
                    st.warning("请在上方公共输入区域填写 POV 开始日期和结束日期。")

                has_pov = "pov_text" in st.session_state
                pov_label = "重新生成" if has_pov else "生成 POV 部署计划"
                if st.button(pov_label, type="primary", use_container_width=True, key="btn_pov"):
                    if not pov_start or not pov_end:
                        st.warning("请先在上方公共输入区域选择 POV 开始日期和结束日期。")
                        st.stop()
                    if pov_end < pov_start:
                        st.warning("POV 结束日期不能早于开始日期。")
                        st.stop()
                    if not has_meaningful_pov_team(vendor_team):
                        st.warning("请在上方公共输入区域填写乙方项目人员。")
                        st.stop()
                    try:
                        pov_prompt = build_pov_prompt(solution, customer, pov_start, pov_end, vendor_team, pov_ref)
                        with st.spinner("正在生成 POV 部署计划..."):
                            pov_text = call_azure_openai(POV_SYSTEM_PROMPT, pov_prompt)
                            st.session_state["pov_text"] = pov_text
                            st.session_state["pov_source_doc_type"] = current_doc_type
                        persist_session_state()
                        st.rerun()
                    except Exception as e:
                        st.error(f"生成失败：{e}")

                if "pov_text" in st.session_state:
                    acct = st.session_state.get("account_name") or account_name.strip() or customer
                    docx_pov = create_pov_docx(
                        content=st.session_state["pov_text"], customer_name=customer
                    )
                    st.download_button(
                        label="下载 POV 部署计划 (.docx)",
                        data=docx_pov,
                        file_name=f"{acct}-PostAssessment POVdeployment.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

            with right:
                if "pov_text" in st.session_state:
                    st.markdown("**文档预览**")
                    st.markdown(st.session_state["pov_text"], unsafe_allow_html=True)
                else:
                    st.info("请填写信息后点击生成")

    # ─────────── Tab 3: Azure Migrate CSV ───────────
    with tab_csv:
        current_doc_type = st.session_state.get("doc_type", "AI")
        has_base_doc = ("solution_text" in st.session_state) if current_doc_type == "AI" else ("infra_text" in st.session_state)

        if not has_base_doc:
            doc_type_name = "AI 解决方案" if current_doc_type == "AI" else "Infra 基础设施"
            st.info(f"请先在「解决方案文档」标签页中生成或导入 {doc_type_name} 文档")
        else:
            customer = st.session_state["customer_name"]
            bdgt = st.session_state.get("budget", budget)
            left, right = st.columns([1, 1])
            with left:
                st.caption(f"📄 当前基于: **{current_doc_type}** 解决方案文档")
                migrate_csv_header = ""
                if os.path.exists(MIGRATE_TEMPLATE_PATH):
                    with open(MIGRATE_TEMPLATE_PATH, "r", encoding="utf-8-sig") as f:
                        migrate_csv_header = f.readline().strip()

                uploaded_excel = st.file_uploader(
                    "上传价格估算表 (.xlsx)",
                    type=["xlsx"],
                    help="上传包含 Azure 资源估算金额的 Excel 文件",
                )

                has_csv = "csv_code" in st.session_state
                csv_label = "重新生成 CSV" if has_csv else "生成 Azure Migrate CSV"
                if st.button(csv_label, type="primary", use_container_width=True, key="btn_csv"):
                    if not uploaded_excel:
                        st.warning("请先上传价格估算表 Excel 文件。")
                        st.stop()
                    if not migrate_csv_header:
                        st.warning("Azure Migrate CSV 模板未找到。")
                        st.stop()
                    try:
                        import openpyxl
                        wb = openpyxl.load_workbook(uploaded_excel, data_only=True)
                        excel_text_parts = []
                        for sheet_name in wb.sheetnames:
                            ws = wb[sheet_name]
                            rows = list(ws.iter_rows(values_only=True))
                            if not rows:
                                continue
                            excel_text_parts.append(f"### Sheet: {sheet_name}")
                            headers = [str(c) if c is not None else "" for c in rows[0]]
                            excel_text_parts.append("| " + " | ".join(headers) + " |")
                            excel_text_parts.append("| " + " | ".join(["---"] * len(headers)) + " |")
                            for row in rows[1:]:
                                cells = [str(c) if c is not None else "" for c in row]
                                excel_text_parts.append("| " + " | ".join(cells) + " |")
                        excel_text = "\n".join(excel_text_parts)

                        csv_prompt = (
                            f"以下是客户的 Azure 价格估算表内容：\n\n{excel_text}\n\n"
                            f"客户预估年消耗：{bdgt}\n\n"
                            f"Azure Migrate CSV 模板表头：\n{migrate_csv_header}\n\n"
                            f"请根据价格估算表倒推本地 VM 配置，按模板格式生成 CSV。"
                        )

                        with st.spinner("正在生成 Azure Migrate CSV..."):
                            csv_raw = call_azure_openai(CSV_SYSTEM_PROMPT, csv_prompt)
                            csv_clean = csv_raw.strip()
                            if csv_clean.startswith("```"):
                                csv_clean = csv_clean.split("\n", 1)[1] if "\n" in csv_clean else csv_clean
                            if csv_clean.endswith("```"):
                                csv_clean = csv_clean[:-3].strip()
                            st.session_state["csv_code"] = csv_clean
                        persist_session_state()
                        st.rerun()
                    except Exception as e:
                        st.error(f"生成失败：{e}")

                if "csv_code" in st.session_state:
                    acct = st.session_state.get("account_name") or account_name.strip() or customer
                    csv_data = st.session_state["csv_code"]
                    st.download_button(
                        label="下载 Azure Migrate CSV",
                        data=csv_data.encode("utf-8-sig"),
                        file_name=f"{acct}-Azure migrate report.csv",
                        mime="text/csv",
                        use_container_width=True,
                    )

            with right:
                if "csv_code" in st.session_state:
                    csv_data = st.session_state["csv_code"]
                    st.markdown("**CSV 预览**")
                    try:
                        import csv as csv_mod
                        csv_lines = csv_data.strip().split("\n")
                        reader = csv_mod.reader(csv_lines)
                        all_rows = list(reader)
                        if len(all_rows) > 1:
                            header = all_rows[0]
                            num_cols = len(header)
                            # 对齐列数：补齐或截断
                            data_rows = []
                            for row in all_rows[1:]:
                                if len(row) < num_cols:
                                    row = row + [""] * (num_cols - len(row))
                                elif len(row) > num_cols:
                                    row = row[:num_cols]
                                data_rows.append(row)
                            import pandas as pd
                            df = pd.DataFrame(data_rows, columns=header)
                            st.dataframe(df, use_container_width=True)
                    except Exception as e:
                        st.warning(f"预览失败，请使用下载查看: {e}")
                        st.code(csv_data, language="csv")
                else:
                    st.info("上传 Excel 后点击生成")


    # ─────────── Tab 4: 年度价格表 ───────────
    with tab_yearly:
        st.markdown(
            "上传从 Azure 定价计算器导出的原始 Excel，自动新增 **Estimated yearly cost** 列（月费用 × 12）并在 Total 行汇总。"
        )

        st.divider()

        uploaded_price = st.file_uploader(
            "上传原始价格表 (.xlsx)",
            type=["xlsx"],
            key="upload_price_excel",
            help="支持标准 Azure 定价计算器导出格式",
        )

        if uploaded_price is not None:
            if st.button("生成年度价格表", type="primary", use_container_width=True, key="btn_gen_yearly"):
                import openpyxl
                from copy import copy as _copy
                from openpyxl.styles import Font as _Font

                def _col_letter(n):
                    result = ""
                    while n:
                        n, rem = divmod(n - 1, 26)
                        result = chr(65 + rem) + result
                    return result

                def _copy_cell_style(src, dst):
                    if src.has_style:
                        dst.font      = _copy(src.font)
                        dst.fill      = _copy(src.fill)
                        dst.border    = _copy(src.border)
                        dst.alignment = _copy(src.alignment)
                        dst.number_format = src.number_format

                def _find_header_row(ws):
                    for i, row in enumerate(ws.iter_rows(values_only=True), 1):
                        if row and "Estimated monthly cost" in row:
                            return i
                    return None

                def _find_total_row(ws, hrow):
                    for i, row in enumerate(ws.iter_rows(min_row=hrow + 1, values_only=True), hrow + 1):
                        if row and "Total" in row:
                            return i
                    return None

                def _get_account_name(ws):
                    """从 Sheet 第 2 行前 5 列取账号名（非空的第一个值）。"""
                    for col in range(1, 6):
                        v = ws.cell(2, col).value
                        if v and str(v).strip():
                            return str(v).strip().rstrip("\t").strip()
                    return None

                def _process_sheet(ws):
                    hrow = _find_header_row(ws)
                    if hrow is None:
                        return False, "未找到标题行（含 'Estimated monthly cost'）", None
                    trow = _find_total_row(ws, hrow)
                    if trow is None:
                        return False, "未找到 Total 行", None

                    header_vals = [ws.cell(hrow, c).value for c in range(1, ws.max_column + 1)]
                    try:
                        monthly_col = header_vals.index("Estimated monthly cost") + 1
                        upfront_col = header_vals.index("Estimated upfront cost") + 1
                    except ValueError:
                        return False, "未找到必要列名", None

                    yearly_col     = upfront_col + 1
                    ws.insert_cols(yearly_col)
                    monthly_letter = _col_letter(monthly_col)
                    yearly_letter  = _col_letter(yearly_col)

                    # 标题行：复制 upfront 列样式
                    hcell = ws.cell(hrow, yearly_col, "Estimated yearly cost")
                    _copy_cell_style(ws.cell(hrow, upfront_col), hcell)
                    src_hdr = ws.cell(hrow, upfront_col)
                    hcell.font = _Font(
                        name=src_hdr.font.name or "Calibri",
                        bold=True,
                        size=src_hdr.font.size or 11,
                    )

                    data_start = hrow + 1
                    data_end   = trow - 1

                    # 数据行：写公式，复制样式并特别保留 number_format（用于显示 $）
                    for r in range(data_start, data_end + 1):
                        mv = ws.cell(r, monthly_col).value
                        if mv is not None and (isinstance(mv, (int, float)) or (isinstance(mv, str) and mv.startswith("="))):
                            cell = ws.cell(r, yearly_col)
                            cell.value = f"={monthly_letter}{r}*12"
                            src_cell = ws.cell(r, monthly_col)
                            _copy_cell_style(src_cell, cell)
                            # 显式保留原始单元格的 number_format，以带上 $ 符号
                            if src_cell.number_format and src_cell.number_format != 'General':
                                cell.number_format = src_cell.number_format
                            else:
                                cell.number_format = '"$"#,##0.00'
                        else:
                            ws.cell(r, yearly_col).value = None

                    # Total 行
                    tcell = ws.cell(trow, yearly_col)
                    tcell.value = f"=SUM({yearly_letter}{data_start}:{yearly_letter}{data_end})"
                    src_total = ws.cell(trow, monthly_col)
                    _copy_cell_style(src_total, tcell)
                    if src_total.number_format and src_total.number_format != 'General':
                        tcell.number_format = src_total.number_format
                    else:
                        tcell.number_format = '"$"#,##0.00'
                    tcell.font = _Font(bold=True, name="Calibri", size=11)

                    ws.column_dimensions[yearly_letter].width = 22
                    
                    account = _get_account_name(ws)
                    return True, "处理成功", account

                try:
                    with st.spinner("正在处理 Excel..."):
                        wb = openpyxl.load_workbook(uploaded_price)
                        messages = []
                        account_name = None
                        for sname in wb.sheetnames:
                            ok, msg, acct = _process_sheet(wb[sname])
                            messages.append(f"**{sname}**: {msg}")
                            if acct and not account_name:
                                account_name = acct

                        # 优先使用用户输入的账户名，其次使用 Excel 中提取的名称
                        _budget = st.session_state.get("budget", budget) or "未填写"
                        _acct_from_input = st.session_state.get("account_name") or account_name.strip()
                        _acct_final = _acct_from_input or account_name or uploaded_price.name.replace(".xlsx", "")
                        new_dl_name = f"{_acct_final}-Azure calculator.xlsx"

                        out_buf = io.BytesIO()
                        wb.save(out_buf)
                        out_buf.seek(0)
                        st.session_state["yearly_excel_bytes"] = out_buf.getvalue()
                        st.session_state["yearly_excel_name"]  = new_dl_name
                        st.session_state["yearly_messages"]    = messages

                    st.rerun()
                except Exception as e:
                    st.error(f"处理失败：{e}")
        else:
            st.info("请先上传 Excel 文件")

        # 处理结果与下载
        if "yearly_excel_bytes" in st.session_state:
            st.divider()
            for msg in st.session_state.get("yearly_messages", []):
                st.markdown(msg)
            st.download_button(
                label="下载任务年度价格表 (.xlsx)",
                data=st.session_state["yearly_excel_bytes"],
                file_name=st.session_state["yearly_excel_name"],
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="dl_yearly",
            )


# ──────────────────────────────────────────────
# 入口
# ──────────────────────────────────────────────
if __name__ == "__main__":
    main()
