"""Shared Word document utilities."""

import io
import os
import re
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config import CN_FONT, CN_FONT_ALT

def svg_to_png_bytes(svg_code: str) -> Optional[bytes]:
    """将 SVG 字符串转换为 PNG 字节。优先 cairosvg → svglib → Edge headless。"""
    try:
        import cairosvg
        png_bytes = cairosvg.svg2png(bytestring=svg_code.encode("utf-8"), output_width=1200)
        return png_bytes
    except (ImportError, OSError):
        pass
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".svg", delete=False, mode="w", encoding="utf-8") as f:
            f.write(svg_code)
            tmp_path = f.name
        drawing = svg2rlg(tmp_path)
        os.unlink(tmp_path)
        if drawing:
            png_bytes = renderPM.drawToString(drawing, fmt="PNG")
            return png_bytes
    except (ImportError, OSError):
        pass
    # 最后尝试使用 Edge 浏览器 headless 模式渲染
    return _svg_to_png_via_edge(svg_code)


def _svg_to_png_via_edge(svg_code: str) -> Optional[bytes]:
    """使用 Edge 浏览器 headless 模式将 SVG 渲染为 PNG。"""
    import subprocess
    import tempfile

    edge_paths = [
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    ]
    edge_exe = None
    for p in edge_paths:
        if os.path.exists(p):
            edge_exe = p
            break
    if not edge_exe:
        return None

    # 解析 viewBox 确定合适的窗口尺寸
    vb_match = re.search(r'viewBox\s*=\s*"([^"]*)"', svg_code)
    if vb_match:
        parts = vb_match.group(1).split()
        if len(parts) == 4:
            vb_w, vb_h = int(float(parts[2])), int(float(parts[3]))
        else:
            vb_w, vb_h = 1200, 800
    else:
        vb_w, vb_h = 1200, 800

    # 确保 SVG 填满整个页面
    html = (
        '<!DOCTYPE html><html><head><meta charset="utf-8">'
        '<style>*{margin:0;padding:0}body{background:white}'
        'svg{display:block;width:100vw;height:100vh}</style></head>'
        f'<body>{svg_code}</body></html>'
    )

    tmp_html = tempfile.mktemp(suffix=".html")
    tmp_png = tempfile.mktemp(suffix=".png")
    tmp_user_data = tempfile.mkdtemp(prefix="edge_svg_")
    try:
        with open(tmp_html, "w", encoding="utf-8") as f:
            f.write(html)

        file_url = "file:///" + tmp_html.replace("\\", "/")
        result = subprocess.run(
            [
                edge_exe,
                "--headless",
                "--disable-gpu",
                "--no-sandbox",
                f"--screenshot={tmp_png}",
                f"--window-size={vb_w},{vb_h}",
                "--default-background-color=00000000",
                "--hide-scrollbars",
                f"--user-data-dir={tmp_user_data}",
                file_url,
            ],
            capture_output=True,
            timeout=30,
        )
        if os.path.exists(tmp_png) and os.path.getsize(tmp_png) > 0:
            with open(tmp_png, "rb") as f:
                return f.read()
    except (subprocess.TimeoutExpired, OSError):
        pass
    finally:
        for p in (tmp_html, tmp_png):
            try:
                os.unlink(p)
            except OSError:
                pass
        try:
            import shutil
            shutil.rmtree(tmp_user_data, ignore_errors=True)
        except Exception:
            pass
    return None


def add_svg_image_to_doc(doc, svg_code: str, width_cm: float = 16) -> bool:
    """
    将 SVG 直接插入到 Word 文档。
    优先转为 PNG 插入（兼容性最好），若无法转换则通过 docx XML 直接嵌入 SVG。
    返回 True 如果成功插入。
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    # 先尝试 PNG 转换
    png_bytes = svg_to_png_bytes(svg_code)
    if png_bytes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(png_bytes), width=Cm(width_cm))
        return True

    # PNG 转换失败，直接嵌入 SVG 为 docx image part（Word 365+ 支持）
    try:
        from docx.opc.part import Part as OpcPart
        from docx.opc.packuri import PackURI

        svg_bytes = svg_code.encode("utf-8")
        # 提取 viewBox 尺寸来计算比例
        vb_match = re.search(r'viewBox\s*=\s*"([^"]*)"', svg_code)
        if vb_match:
            parts = vb_match.group(1).split()
            if len(parts) == 4:
                vb_w, vb_h = float(parts[2]), float(parts[3])
            else:
                vb_w, vb_h = 1200, 800
        else:
            vb_w, vb_h = 1200, 800

        width_emu = int(width_cm * 360000)  # cm to EMU
        height_emu = int(width_emu * vb_h / max(vb_w, 1))

        # 添加 SVG 作为 image part
        part = doc.part
        svg_part = OpcPart(
            PackURI("/word/media/architecture.svg"),
            "image/svg+xml",
            svg_bytes,
            part.package,
        )
        r_id = part.relate_to(svg_part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")

        # 创建内联图片 XML（不要用 .set() 设置 xmlns 属性，OxmlElement 通过前缀自动处理命名空间）
        inline = OxmlElement("wp:inline")
        inline.set("distT", "0")
        inline.set("distB", "0")
        inline.set("distL", "0")
        inline.set("distR", "0")

        extent = OxmlElement("wp:extent")
        extent.set("cx", str(width_emu))
        extent.set("cy", str(height_emu))
        inline.append(extent)

        docPr = OxmlElement("wp:docPr")
        docPr.set("id", "1")
        docPr.set("name", "Architecture Diagram")
        inline.append(docPr)

        graphic = OxmlElement("a:graphic")

        graphicData = OxmlElement("a:graphicData")
        graphicData.set("uri", "http://schemas.openxmlformats.org/drawingml/2006/picture")

        pic = OxmlElement("pic:pic")

        nvPicPr = OxmlElement("pic:nvPicPr")
        cNvPr = OxmlElement("pic:cNvPr")
        cNvPr.set("id", "0")
        cNvPr.set("name", "architecture.svg")
        nvPicPr.append(cNvPr)
        nvPicPr.append(OxmlElement("pic:cNvPicPr"))
        pic.append(nvPicPr)

        blipFill = OxmlElement("pic:blipFill")
        blip = OxmlElement("a:blip")
        blip.set(_qn("r:embed"), r_id)
        blipFill.append(blip)
        stretch = OxmlElement("a:stretch")
        stretch.append(OxmlElement("a:fillRect"))
        blipFill.append(stretch)
        pic.append(blipFill)

        spPr = OxmlElement("pic:spPr")
        xfrm = OxmlElement("a:xfrm")
        off = OxmlElement("a:off")
        off.set("x", "0")
        off.set("y", "0")
        xfrm.append(off)
        ext = OxmlElement("a:ext")
        ext.set("cx", str(width_emu))
        ext.set("cy", str(height_emu))
        xfrm.append(ext)
        spPr.append(xfrm)
        prstGeom = OxmlElement("a:prstGeom")
        prstGeom.set("prst", "rect")
        spPr.append(prstGeom)
        pic.append(spPr)

        graphicData.append(pic)
        graphic.append(graphicData)
        inline.append(graphic)

        # 将 inline 放到段落 run 的 drawing 元素中
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        drawing = OxmlElement("w:drawing")
        drawing.append(inline)
        run._element.append(drawing)
        return True
    except Exception:
        return False


def set_run_font(run, font_name=CN_FONT, size_pt=None, bold=None, color_rgb=None):
    """为 run 设置字体（含中文 eastAsia 字体）。"""
    run.font.name = font_name
    # python-docx 需要同时设置 eastAsia 字体才能在 Word 中正确显示中文
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb


def add_styled_paragraph(doc, text, font_name=CN_FONT, size_pt=9, bold=False,
                          color_rgb=None, alignment=None, indent=True):
    """添加一个带完整样式的段落。indent=True 时添加首行缩进。"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    # 首行缩进（约 1 个 Tab = 0.74cm）
    if indent and alignment is None:
        p.paragraph_format.first_line_indent = Cm(0.74)
    # 处理 **加粗** 和普通文字的混合
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        is_bold = bold or (i % 2 == 1)
        set_run_font(run, font_name=font_name, size_pt=size_pt, bold=is_bold,
                       color_rgb=color_rgb)
    return p


def add_styled_heading(doc, text, level=1):
    """添加一个使用中文字体的标题。"""
    heading = doc.add_heading("", level=level)
    run = heading.add_run(text)
    size_map = {1: 18, 2: 14, 3: 12}
    set_run_font(run, font_name=CN_FONT, size_pt=size_map.get(level, 12), bold=True)
    return heading


def parse_markdown_table(lines: List[str]) -> Optional[List[List[str]]]:
    """
    尝试从 Markdown 行列表中解析表格。
    返回二维数组 (包含表头)，如果不是表格则返回 None。
    """
    if len(lines) < 2:
        return None
    # 检查是否是 Markdown 表格（至少有 | 分隔符和分隔行 ---）
    if "|" not in lines[0]:
        return None

    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # 跳过分隔行 |---|---|
        if re.match(r"^\|[\s\-:|]+\|$", stripped):
            continue
        # 解析单元格
        cells = [c.strip() for c in stripped.split("|")]
        # 去掉首尾空元素（因为 | 在开头和结尾会产生空字符串）
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows if len(rows) >= 2 else None


def add_word_table(doc, table_data: list[list[str]]):
    """将二维数组写入 Word 表格，应用专业样式。"""
    if not table_data:
        return

    num_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = "Table Grid"

    for ri, row_data in enumerate(table_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""  # 清空默认段落文本
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            is_header = (ri == 0)
            set_run_font(
                run,
                font_name=CN_FONT,
                size_pt=9,
                bold=is_header,
            )
            # 表头行背景色
            if is_header:
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(
                    qn("w:shd"),
                    {qn("w:fill"): "156082", qn("w:val"): "clear"},
                )
                shading.append(shading_elem)
                run.font.color.rgb = RGBColor(255, 255, 255)


def markdown_to_docx(doc, markdown_text: str, body_size=9):
    """
    将 AI 返回的 Markdown 文本解析并写入 Word 文档。
    支持: 标题 (#/##/###)、列表 (-/*)、Markdown 表格、加粗 (**)、普通段落。
    """
    # 预处理：将 <br> 变体转换为换行符
    markdown_text = re.sub(r"<br\s*/?>", "\n", markdown_text)
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # ── 跳过 --- 分隔线 ──
        if stripped == '---' or stripped == '***' or stripped == '___':
            i += 1
            continue

        # ── 标题 ──
        if stripped.startswith("#### "):
            add_styled_heading(doc, stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith("### "):
            add_styled_heading(doc, stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            add_styled_heading(doc, stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            add_styled_heading(doc, stripped[2:], level=1)
            i += 1
            continue

        # ── 独立的 **加粗行**（如阶段标题），转为三级标题 ──
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            title_text = stripped[2:-2]
            add_styled_heading(doc, title_text, level=3)
            i += 1
            continue

        # ── Markdown 表格 ──
        if "|" in stripped and not stripped.startswith("-"):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            table_data = parse_markdown_table(table_lines)
            if table_data:
                add_word_table(doc, table_data)
                doc.add_paragraph()  # 表格后空行
            else:
                # 不是表格，作为普通文本处理
                for tl in table_lines:
                    add_styled_paragraph(doc, tl.strip(), size_pt=body_size)
            continue

        # ── 无序列表 ──
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            add_styled_paragraph(doc, f"•  {text}", size_pt=body_size)
            i += 1
            continue

        # ── 有序列表 ──
        if stripped[0].isdigit() and ". " in stripped[:5]:
            add_styled_paragraph(doc, stripped, size_pt=body_size)
            i += 1
            continue

        # ── 普通段落 ──
        add_styled_paragraph(doc, stripped, size_pt=body_size)
        i += 1


# ──────────────────────────────────────────────
# Word 文档生成 —— 基于模板
# ──────────────────────────────────────────────
def load_template(template_path: str) -> Document:
    """
    加载 .docx 模板文件作为基础文档。
    如果模板不存在，则返回一个空白 Document。
    """
    if os.path.exists(template_path):
        doc = Document(template_path)
        # 清空模板中的所有正文段落（保留样式定义、页面设置、页眉页脚）
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        # 清空表格
        for t in doc.tables:
            t._element.getparent().remove(t._element)
        return doc
    else:
        return Document()


def extract_title(content: str, fallback: str = "") -> str:
    """从 AI 生成的 Markdown 内容中提取第一个 # 标题作为文档标题。"""
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            return stripped[2:].strip()
    return fallback


def strip_first_heading(content: str) -> str:
    """去掉 Markdown 内容中的第一个 # 标题行（因为封面已经显示了标题）。"""
    lines = content.split("\n")
    result = []
    found = False
    for line in lines:
        stripped = line.strip()
        if not found and stripped.startswith("# ") and not stripped.startswith("## "):
            found = True
            continue  # 跳过第一个 # 标题
        result.append(line)
    return "\n".join(result)


def add_page_break(doc):
    """在文档中添加分页符。"""
    from docx.oxml.ns import qn as _qn
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(_qn("w:br"), {_qn("w:type"): "page"})
    run._element.append(br)


def add_toc(doc):
    """插入 Word 目录域（用户打开文档后按 Ctrl+A → F9 即可更新）。"""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    # 目录标题
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目录")
    set_run_font(run, font_name=CN_FONT, size_pt=16, bold=True)

    doc.add_paragraph()  # 空行

    # 插入 TOC 域代码
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(_qn("w:fldCharType"), "begin")
    run._element.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(_qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(_qn("w:fldCharType"), "separate")
    run._element.append(fldChar_separate)

    # 占位文本（打开 Word 后会自动替换）
    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "（请右键点击此处 → 更新域，生成目录）"
    placeholder.append(placeholder_text)
    run._element.append(placeholder)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(_qn("w:fldCharType"), "end")
    run._element.append(fldChar_end)
