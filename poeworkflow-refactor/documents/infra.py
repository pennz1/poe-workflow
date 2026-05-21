"""Infrastructure document generation."""

import io

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from config import CN_FONT, CN_FONT_ALT, INFRA_TEMPLATE_PATH
from documents.docx_utils import (
    add_page_break,
    add_toc,
    extract_title,
    load_template,
    markdown_to_docx,
    set_run_font,
    strip_first_heading,
)

def create_infra_docx(content: str, customer_name: str) -> bytes:
    """
    基于 Infra 模板生成基础设施解决方案 Word 文档。
    布局: 封面标题（独占一页） → 目录（独占一页） → 正文
    """
    doc = load_template(INFRA_TEMPLATE_PATH)
    title = extract_title(content, f"{customer_name} - 基础设施解决方案架构文档")
    body_content = strip_first_heading(content)

    # ---- 第 1 页：封面标题 ----
    for _ in range(8):
        doc.add_paragraph()

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    # 与 AI 解决方案一致: 18pt #4874CB
    set_run_font(run, font_name=CN_FONT_ALT, size_pt=18,
                   bold=True, color_rgb=RGBColor(0x48, 0x74, 0xCB))

    # 封面分页
    add_page_break(doc)

    # ---- 第 2 页：目录 ----
    add_toc(doc)

    # 目录分页
    add_page_break(doc)

    # ---- 第 3 页起：正文内容（已去掉第一个 # 标题） ----
    markdown_to_docx(doc, body_content, body_size=9)

    # 导出
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
