"""POV deployment plan document generation."""

import datetime
import io
import re
from typing import Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from config import CN_FONT, CN_FONT_ALT, POV_TEMPLATE_PATH
from documents.docx_utils import (
    add_page_break,
    extract_title,
    load_template,
    markdown_to_docx,
    set_run_font,
    strip_first_heading,
)

def create_pov_docx(content: str, customer_name: str) -> bytes:
    """
    基于 POV 模板生成 POV 部署计划 Word 文档。
    布局: 封面标题（独占一页） → 正文
    """
    doc = load_template(POV_TEMPLATE_PATH)
    title = extract_title(content, f"{customer_name} - POV 部署计划")
    body_content = strip_first_heading(content)

    # ---- 第 1 页：封面标题 ----
    for _ in range(8):
        doc.add_paragraph()

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    # 与模板一致: 22pt #156082
    set_run_font(run, font_name=CN_FONT_ALT, size_pt=22,
                   bold=True, color_rgb=RGBColor(0x15, 0x60, 0x82))

    # 封面分页
    add_page_break(doc)

    # ---- 第 2 页起：正文内容（已去掉第一个 # 标题） ----
    markdown_to_docx(doc, body_content, body_size=9)

    # 导出
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _workday_info(start_date: datetime.date, end_date: datetime.date) -> tuple[list[str], list[str]]:
    workdays = []
    weekends = []
    current = start_date
    while current <= end_date:
        target = f"{current.month}月{current.day}日"
        if current.weekday() < 5:
            workdays.append(target)
        else:
            weekends.append(target)
        current += datetime.timedelta(days=1)
    return workdays, weekends


def has_meaningful_pov_team(vendor_team: Optional[str]) -> bool:
    text = str(vendor_team or "").strip()
    if not text:
        return False
    placeholders = {"技术负责人", "Azure架构师", "Azure 架构师", "项目经理", "负责人"}
    for raw_line in text.splitlines():
        line = raw_line.strip().strip("：:")
        if not line:
            continue
        if ":" in raw_line or "：" in raw_line:
            _, value = re.split(r"[:：]", raw_line, maxsplit=1)
            if value.strip():
                return True
            if line in placeholders:
                continue
        elif line not in placeholders:
            return True
    return False


def build_pov_prompt(
    solution_text: str,
    customer_name: str,
    pov_start: datetime.date,
    pov_end: datetime.date,
    vendor_team: str,
    pov_ref: str,
) -> str:
    workdays, weekends = _workday_info(pov_start, pov_end)
    pov_prompt = (
        f"以下是已生成的解决方案架构文档，请据此生成 POV 部署计划：\n\n"
        f"{solution_text}\n\n"
        f"## 补充信息\n- **客户名称**：{customer_name}\n"
        f"- **POV 周期**：{pov_start.strftime('%Y/%m/%d')} - {pov_end.strftime('%Y/%m/%d')}\n\n"
        f"## 可用工作日清单（共 {len(workdays)} 天，必须且只能使用这些日期）\n"
        f"{'、'.join(workdays)}\n\n"
        f"## 禁用日期（周末，严禁安排任何任务）\n"
        f"{'、'.join(weekends) if weekends else '无'}\n\n"
        f"## 乙方项目人员\n{vendor_team.strip()}\n\n"
        f"请根据客户背景信息自动生成合理的甲方人员（2-3人，包含项目负责人和技术对接人，一定要中文名！）。"
    )
    if pov_ref:
        pov_prompt += (
            "\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
            f"{pov_ref}"
        )
    return pov_prompt
