"""Solution architecture document generation."""

import io
from typing import Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from config import CN_FONT, CN_FONT_ALT, SOLUTION_TEMPLATE_PATH
from documents.docx_utils import (
    add_page_break,
    add_svg_image_to_doc,
    add_toc,
    extract_title,
    load_template,
    markdown_to_docx,
    set_run_font,
    strip_first_heading,
)
from llm.client import call_azure_openai
from llm.prompts import SVG_SYSTEM_PROMPT, _extract_svg_from_response

def generate_svg_architecture(solution_text: str, customer_name: str) -> Optional[str]:
    """
    根据解决方案文本生成 SVG 架构图。
    提取第 2、5、6、7、8 章节的内容作为输入。
    返回 SVG 字符串，失败返回 None。
    """
    # 提取相关章节
    lines = solution_text.split("\n")
    relevant_sections = []
    current_section = None
    capture = False
    target_prefixes = ("二、", "五、", "六、", "七、", "八、", "2.", "5.", "6.", "7.", "8.")

    for line in lines:
        stripped = line.strip()
        # 检测标题行（## 开头）
        if stripped.startswith("## ") or stripped.startswith("# "):
            heading = stripped.lstrip("#").strip()
            if any(heading.startswith(p) for p in target_prefixes):
                capture = True
                current_section = heading
                relevant_sections.append(f"\n## {heading}\n")
            else:
                capture = False
        elif capture:
            relevant_sections.append(line)

    if not relevant_sections:
        # 如果没找到编号章节，使用全文
        context_text = solution_text[:8000]
    else:
        context_text = "\n".join(relevant_sections)[:8000]

    user_prompt = (
        f"请为以下客户 **{customer_name}** 的 Azure 解决方案生成 SVG 架构图。\n"
        f"图表标题请包含客户名称：\"{customer_name} - Azure AI 解决方案架构\"。\n\n"
        f"以下是方案的关键章节内容：\n\n{context_text}"
    )

    try:
        svg_response = call_azure_openai(SVG_SYSTEM_PROMPT, user_prompt)
        svg_code = _extract_svg_from_response(svg_response)
        # 基本验证
        if "<svg" in svg_code and "</svg>" in svg_code:
            return svg_code
        return None
    except Exception:
        return None


def create_solution_docx(content: str, customer_name: str, svg_code: Optional[str] = None) -> bytes:
    """
    基于 solution 模板生成解决方案架构 Word 文档。
    布局: 封面标题（独占一页） → 目录（独占一页） → 正文
    如果提供 svg_code，则在第二章节结束后插入架构图。
    """
    doc = load_template(SOLUTION_TEMPLATE_PATH)
    title = extract_title(content, f"{customer_name} - AI 解决方案架构文档")
    body_content = strip_first_heading(content)

    # ---- 第 1 页：封面标题 ----
    for _ in range(8):
        doc.add_paragraph()

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    # 与模板一致: 18pt #4874CB
    set_run_font(run, font_name=CN_FONT_ALT, size_pt=18,
                   bold=True, color_rgb=RGBColor(0x48, 0x74, 0xCB))

    # 封面分页
    add_page_break(doc)

    # ---- 第 2 页：目录 ----
    add_toc(doc)

    # 目录分页
    add_page_break(doc)

    # ---- 第 3 页起：正文内容（已去掉第一个 # 标题） ----
    if svg_code:
        # 在第二章节后插入架构图
        # 查找第三个 ## 标题（即第三章开头），在其前面插入图片
        lines = body_content.split("\n")
        h2_count = 0
        split_idx = len(lines)
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith("## ") or (stripped.startswith("# ") and not stripped.startswith("## ")):
                h2_count += 1
                if h2_count == 3:  # 第三章开头
                    split_idx = idx
                    break

        # 渲染第二章及之前的内容
        part1 = "\n".join(lines[:split_idx])
        markdown_to_docx(doc, part1, body_size=9)

        # 插入架构图
        doc.add_paragraph()  # 空行
        add_svg_image_to_doc(doc, svg_code, width_cm=16)

        # 图注
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"图：{customer_name} Azure AI 解决方案架构图")
        set_run_font(cap_run, font_name=CN_FONT, size_pt=8, bold=False)

        doc.add_paragraph()  # 空行

        # 渲染剩余内容
        part2 = "\n".join(lines[split_idx:])
        if part2.strip():
            markdown_to_docx(doc, part2, body_size=9)
    else:
        markdown_to_docx(doc, body_content, body_size=9)

    # 导出
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
