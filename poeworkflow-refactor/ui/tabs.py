"""Streamlit tab renderers."""

import io
import os

import streamlit as st
from docx import Document

from config import MIGRATE_TEMPLATE_PATH
from documents.infra import create_infra_docx
from documents.pov import build_pov_prompt, create_pov_docx, has_meaningful_pov_team
from documents.solution import create_solution_docx
from llm.client import call_azure_openai
from llm.prompts import CSV_SYSTEM_PROMPT, INFRA_SYSTEM_PROMPT, POV_SYSTEM_PROMPT, SOLUTION_SYSTEM_PROMPT
from session import persist_session_state

def render_solution_tab(customer_name, account_name, budget, customer_bg, solution_ref, infra_ref):
    # 文档类型切换
    doc_type = st.radio(
        "选择文档类型",
        ["AI 解决方案", "Infra 基础设施"],
        horizontal=True,
        key="doc_type_radio",
        index=0 if st.session_state.get("doc_type", "AI") == "AI" else 1,
    )
    current_doc_type = "AI" if doc_type == "AI 解决方案" else "Infra"
    st.session_state["doc_type"] = current_doc_type

    # 文档来源切换
    doc_source = st.radio(
        "文档来源",
        ["AI 生成", "手动导入"],
        horizontal=True,
        key="doc_source_radio",
    )

    left, right = st.columns([1, 1])
    with left:
        if doc_source == "手动导入":
            # ── 手动导入：两步流程 ──
            # Step 1：上传 / 粘贴，确认后暂存原文
            # Step 2：AI 按模板格式重新生成
            if "imported_doc_text" not in st.session_state:
                # ── Step 1：上传或粘贴 ──
                uploaded_doc = st.file_uploader(
                    "上传已有的 .docx 文档",
                    type=["docx"],
                    key="upload_existing_doc",
                    help="上传后将自动提取文档文本内容",
                )
                manual_text = st.text_area(
                    "或直接粘贴文本内容",
                    height=200,
                    key="manual_doc_text",
                    placeholder="将已有的解决方案文档内容粘贴到此处...",
                )

                if st.button("确认导入", type="primary", use_container_width=True, key="btn_import"):
                    imported_text = ""
                    if uploaded_doc is not None:
                        doc = Document(uploaded_doc)
                        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                        for table in doc.tables:
                            for row in table.rows:
                                cells = [cell.text.strip() for cell in row.cells]
                                paragraphs.append(" | ".join(cells))
                        imported_text = "\n\n".join(paragraphs)
                    elif manual_text.strip():
                        imported_text = manual_text.strip()
                    else:
                        st.warning("请上传文档或粘贴文本。")
                        st.stop()

                    st.session_state["imported_doc_text"] = imported_text
                    # 同步写入 solution_text / infra_text，使 POV 等后续步骤可立即识别到文档
                    target_key = "solution_text" if current_doc_type == "AI" else "infra_text"
                    st.session_state[target_key] = imported_text
                    st.session_state["customer_name"] = customer_name.strip() if customer_name.strip() else "未命名客户"
                    st.session_state["account_name"] = account_name.strip() if account_name.strip() else (customer_name.strip() or "未命名客户")
                    st.session_state["budget"] = budget
                    st.session_state.pop("pov_text", None)
                    persist_session_state()
                    st.rerun()

            else:
                # ── Step 2：确认内容 + AI 重新生成 ──
                imported_text = st.session_state["imported_doc_text"]
                st.success(f"文档已导入（共 {len(imported_text)} 字符）")
                st.text_area(
                    "导入内容预览",
                    value=imported_text[:600] + "\n\n..." if len(imported_text) > 600 else imported_text,
                    height=160,
                    disabled=True,
                    key="preview_imported",
                )

                c_reimport, c_regen = st.columns(2)
                with c_reimport:
                    if st.button("重新上传", use_container_width=True, key="btn_reimport"):
                        st.session_state.pop("imported_doc_text", None)
                        st.rerun()
                with c_regen:
                    if st.button("AI 重新生成", type="primary", use_container_width=True, key="btn_regen_import"):
                        cust = customer_name.strip() or st.session_state.get("customer_name", "未命名客户")
                        system_prompt = SOLUTION_SYSTEM_PROMPT if current_doc_type == "AI" else INFRA_SYSTEM_PROMPT
                        ref_text = solution_ref if current_doc_type == "AI" else infra_ref
                        user_ctx = (
                            f"## 客户信息\n- **客户名称**：{cust}\n\n"
                        )
                        if customer_bg.strip():
                            user_ctx += (
                                f"## 客户背景信息\n{customer_bg.strip()}\n\n"
                            )
                        user_ctx += (
                            f"## 已有解决方案文档（请基于以上客户信息和以下已有文档，按照要求的章节格式重新整理生成，不要照抄原文）\n\n"
                            f"{imported_text}"
                        )
                        if ref_text:
                            user_ctx += (
                                f"\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
                                f"{ref_text}"
                            )
                        try:
                            with st.spinner("正在基于导入内容 AI 重新生成..."):
                                result_text = call_azure_openai(system_prompt, user_ctx)
                                target_key = "solution_text" if current_doc_type == "AI" else "infra_text"
                                st.session_state[target_key] = result_text
                                st.session_state["customer_name"] = cust
                                st.session_state["account_name"] = account_name.strip() if account_name.strip() else cust
                                st.session_state["budget"] = budget
                                st.session_state.pop("pov_text", None)
                                st.session_state.pop("imported_doc_text", None)
                            persist_session_state()
                            st.rerun()
                        except Exception as e:
                            st.error(f"生成失败：{e}")

                # 若已生成，显示下载按钮
                target_key = "solution_text" if current_doc_type == "AI" else "infra_text"
                if target_key in st.session_state:
                    customer = st.session_state["customer_name"]
                    acct = st.session_state.get("account_name") or account_name.strip() or customer
                    if current_doc_type == "AI":
                        docx_bytes = create_solution_docx(
                            content=st.session_state["solution_text"], customer_name=customer
                        )
                        st.download_button(
                            label="下载 AI 解决方案架构文档 (.docx)",
                            data=docx_bytes,
                            file_name=f"{acct}-Solution Architecture.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="dl_sol_import",
                        )
                    else:
                        docx_bytes = create_infra_docx(
                            content=st.session_state["infra_text"], customer_name=customer
                        )
                        st.download_button(
                            label="下载 Infra 基础设施架构文档 (.docx)",
                            data=docx_bytes,
                            file_name=f"{acct}-Infra Solution Architecture.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="dl_infra_import",
                        )

        else:
            # ── AI 生成文档 ──
            if current_doc_type == "AI":
                # AI 解决方案文档逻辑
                has_solution = "solution_text" in st.session_state
                sol_label = "重新生成" if has_solution else "生成 AI 解决方案架构文档"
                if st.button(sol_label, type="primary", use_container_width=True, key="btn_sol"):
                    if not customer_name.strip():
                        st.warning("请输入客户名称。")
                        st.stop()
                    if not customer_bg.strip():
                        st.warning("请输入客户背景信息。")
                        st.stop()
                    try:
                        with st.spinner("正在生成 AI 解决方案架构文档..."):
                            user_ctx = (
                                f"## 客户信息\n- **客户名称**：{customer_name}\n\n"
                                f"## 客户背景\n{customer_bg}"
                            )
                            if solution_ref:
                                user_ctx += (
                                    f"\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
                                    f"{solution_ref}"
                                )
                            sol_text = call_azure_openai(SOLUTION_SYSTEM_PROMPT, user_ctx)
                            st.session_state["solution_text"] = sol_text
                            st.session_state["customer_name"] = customer_name
                            st.session_state["account_name"] = account_name.strip() if account_name.strip() else customer_name
                            st.session_state["budget"] = budget
                            st.session_state.pop("pov_text", None)
                            st.session_state.pop("svg_code", None)
                        persist_session_state()
                        st.rerun()
                    except Exception as e:
                        st.error(f"生成失败：{e}")

                if "solution_text" in st.session_state:
                    customer = st.session_state["customer_name"]
                    acct = st.session_state.get("account_name") or account_name.strip() or customer
                    docx_sol = create_solution_docx(
                        content=st.session_state["solution_text"], customer_name=customer
                    )
                    st.download_button(
                        label="下载 AI 解决方案架构文档 (.docx)",
                        data=docx_sol,
                        file_name=f"{acct}-Solution Architecture.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
            else:
                # Infra 基础设施文档逻辑
                has_infra = "infra_text" in st.session_state
                infra_label = "重新生成" if has_infra else "生成 Infra 基础设施架构文档"
                if st.button(infra_label, type="primary", use_container_width=True, key="btn_infra"):
                    if not customer_name.strip():
                        st.warning("请输入客户名称。")
                        st.stop()
                    if not customer_bg.strip():
                        st.warning("请输入客户背景信息。")
                        st.stop()
                    try:
                        with st.spinner("正在生成 Infra 基础设施架构文档..."):
                            user_ctx = (
                                f"## 客户信息\n- **客户名称**：{customer_name}\n\n"
                                f"## 客户背景\n{customer_bg}"
                            )
                            if infra_ref:
                                user_ctx += (
                                    f"\n\n---\n\n## 【参考模板文档 —— 请学习其风格和结构，不要照抄具体数据】\n\n"
                                    f"{infra_ref}"
                                )
                            infra_text = call_azure_openai(INFRA_SYSTEM_PROMPT, user_ctx)
                            st.session_state["infra_text"] = infra_text
                            st.session_state["customer_name"] = customer_name
                            st.session_state["account_name"] = account_name.strip() if account_name.strip() else customer_name
                            st.session_state["budget"] = budget
                            st.session_state.pop("pov_text", None)
                            st.session_state.pop("svg_code", None)
                        persist_session_state()
                        st.rerun()
                    except Exception as e:
                        st.error(f"生成失败：{e}")

                if "infra_text" in st.session_state:
                    customer = st.session_state["customer_name"]
                    acct = st.session_state.get("account_name") or account_name.strip() or customer
                    docx_infra = create_infra_docx(
                        content=st.session_state["infra_text"], customer_name=customer
                    )
                    st.download_button(
                        label="下载 Infra 基础设施架构文档 (.docx)",
                        data=docx_infra,
                        file_name=f"{acct}-Infra Solution Architecture.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

    with right:
        if current_doc_type == "AI":
            if "solution_text" in st.session_state:
                st.markdown("**AI 解决方案文档预览**")
                st.markdown(st.session_state["solution_text"], unsafe_allow_html=True)
            else:
                st.info("请先生成或导入 AI 解决方案文档")
        else:
            if "infra_text" in st.session_state:
                st.markdown("**Infra 基础设施文档预览**")
                st.markdown(st.session_state["infra_text"], unsafe_allow_html=True)
            else:
                st.info("请先生成或导入 Infra 基础设施文档")


def render_pov_tab(pov_ref, account_name=""):
    # 根据当前文档类型确定使用哪个解决方案文档
    current_doc_type = st.session_state.get("doc_type", "AI")
    has_base_doc = ("solution_text" in st.session_state) if current_doc_type == "AI" else ("infra_text" in st.session_state)

    if not has_base_doc:
        doc_type_name = "AI 解决方案" if current_doc_type == "AI" else "Infra 基础设施"
        st.info(f"请先在「解决方案文档」标签页中生成或导入 {doc_type_name} 文档")
    else:
        customer = st.session_state["customer_name"]
        solution = st.session_state["solution_text"] if current_doc_type == "AI" else st.session_state["infra_text"]
        left, right = st.columns([1, 1])
        with left:
            st.caption(f"📄 当前基于: **{current_doc_type}** 解决方案文档")

            pov_start = st.session_state.get("pov_start_date")
            pov_end = st.session_state.get("pov_end_date")
            vendor_team = st.session_state.get("pov_vendor_team", "")

            if pov_start and pov_end:
                st.info(f"POV 周期：{pov_start} ~ {pov_end}")
            else:
                st.warning("请在上方公共输入区域填写 POV 开始日期和结束日期。")

            has_pov = "pov_text" in st.session_state
            pov_label = "重新生成" if has_pov else "生成 POV 部署计划"
            if st.button(pov_label, type="primary", use_container_width=True, key="btn_pov"):
                if not pov_start or not pov_end:
                    st.warning("请先在上方公共输入区域选择 POV 开始日期和结束日期。")
                    st.stop()
                if pov_end < pov_start:
                    st.warning("POV 结束日期不能早于开始日期。")
                    st.stop()
                if not has_meaningful_pov_team(vendor_team):
                    st.warning("请在上方公共输入区域填写乙方项目人员。")
                    st.stop()
                try:
                    pov_prompt = build_pov_prompt(solution, customer, pov_start, pov_end, vendor_team, pov_ref)
                    with st.spinner("正在生成 POV 部署计划..."):
                        pov_text = call_azure_openai(POV_SYSTEM_PROMPT, pov_prompt)
                        st.session_state["pov_text"] = pov_text
                        st.session_state["pov_source_doc_type"] = current_doc_type
                    persist_session_state()
                    st.rerun()
                except Exception as e:
                    st.error(f"生成失败：{e}")

            if "pov_text" in st.session_state:
                acct = st.session_state.get("account_name") or account_name.strip() or customer
                docx_pov = create_pov_docx(
                    content=st.session_state["pov_text"], customer_name=customer
                )
                st.download_button(
                    label="下载 POV 部署计划 (.docx)",
                    data=docx_pov,
                    file_name=f"{acct}-PostAssessment POVdeployment.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

        with right:
            if "pov_text" in st.session_state:
                st.markdown("**文档预览**")
                st.markdown(st.session_state["pov_text"], unsafe_allow_html=True)
            else:
                st.info("请填写信息后点击生成")


def render_csv_tab(budget, account_name=""):
    current_doc_type = st.session_state.get("doc_type", "AI")
    has_base_doc = ("solution_text" in st.session_state) if current_doc_type == "AI" else ("infra_text" in st.session_state)

    if not has_base_doc:
        doc_type_name = "AI 解决方案" if current_doc_type == "AI" else "Infra 基础设施"
        st.info(f"请先在「解决方案文档」标签页中生成或导入 {doc_type_name} 文档")
    else:
        customer = st.session_state["customer_name"]
        bdgt = st.session_state.get("budget", budget)
        left, right = st.columns([1, 1])
        with left:
            st.caption(f"📄 当前基于: **{current_doc_type}** 解决方案文档")
            migrate_csv_header = ""
            if os.path.exists(MIGRATE_TEMPLATE_PATH):
                with open(MIGRATE_TEMPLATE_PATH, "r", encoding="utf-8-sig") as f:
                    migrate_csv_header = f.readline().strip()

            uploaded_excel = st.file_uploader(
                "上传价格估算表 (.xlsx)",
                type=["xlsx"],
                help="上传包含 Azure 资源估算金额的 Excel 文件",
            )

            has_csv = "csv_code" in st.session_state
            csv_label = "重新生成 CSV" if has_csv else "生成 Azure Migrate CSV"
            if st.button(csv_label, type="primary", use_container_width=True, key="btn_csv"):
                if not uploaded_excel:
                    st.warning("请先上传价格估算表 Excel 文件。")
                    st.stop()
                if not migrate_csv_header:
                    st.warning("Azure Migrate CSV 模板未找到。")
                    st.stop()
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(uploaded_excel, data_only=True)
                    excel_text_parts = []
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        rows = list(ws.iter_rows(values_only=True))
                        if not rows:
                            continue
                        excel_text_parts.append(f"### Sheet: {sheet_name}")
                        headers = [str(c) if c is not None else "" for c in rows[0]]
                        excel_text_parts.append("| " + " | ".join(headers) + " |")
                        excel_text_parts.append("| " + " | ".join(["---"] * len(headers)) + " |")
                        for row in rows[1:]:
                            cells = [str(c) if c is not None else "" for c in row]
                            excel_text_parts.append("| " + " | ".join(cells) + " |")
                    excel_text = "\n".join(excel_text_parts)

                    csv_prompt = (
                        f"以下是客户的 Azure 价格估算表内容：\n\n{excel_text}\n\n"
                        f"客户预估年消耗：{bdgt}\n\n"
                        f"Azure Migrate CSV 模板表头：\n{migrate_csv_header}\n\n"
                        f"请根据价格估算表倒推本地 VM 配置，按模板格式生成 CSV。"
                    )

                    with st.spinner("正在生成 Azure Migrate CSV..."):
                        csv_raw = call_azure_openai(CSV_SYSTEM_PROMPT, csv_prompt)
                        csv_clean = csv_raw.strip()
                        if csv_clean.startswith("```"):
                            csv_clean = csv_clean.split("\n", 1)[1] if "\n" in csv_clean else csv_clean
                        if csv_clean.endswith("```"):
                            csv_clean = csv_clean[:-3].strip()
                        st.session_state["csv_code"] = csv_clean
                    persist_session_state()
                    st.rerun()
                except Exception as e:
                    st.error(f"生成失败：{e}")

            if "csv_code" in st.session_state:
                acct = st.session_state.get("account_name") or account_name.strip() or customer
                csv_data = st.session_state["csv_code"]
                st.download_button(
                    label="下载 Azure Migrate CSV",
                    data=csv_data.encode("utf-8-sig"),
                    file_name=f"{acct}-Azure migrate report.csv",
                    mime="text/csv",
                    use_container_width=True,
                )

        with right:
            if "csv_code" in st.session_state:
                csv_data = st.session_state["csv_code"]
                st.markdown("**CSV 预览**")
                try:
                    import csv as csv_mod
                    csv_lines = csv_data.strip().split("\n")
                    reader = csv_mod.reader(csv_lines)
                    all_rows = list(reader)
                    if len(all_rows) > 1:
                        header = all_rows[0]
                        num_cols = len(header)
                        # 对齐列数：补齐或截断
                        data_rows = []
                        for row in all_rows[1:]:
                            if len(row) < num_cols:
                                row = row + [""] * (num_cols - len(row))
                            elif len(row) > num_cols:
                                row = row[:num_cols]
                            data_rows.append(row)
                        import pandas as pd
                        df = pd.DataFrame(data_rows, columns=header)
                        st.dataframe(df, use_container_width=True)
                except Exception as e:
                    st.warning(f"预览失败，请使用下载查看: {e}")
                    st.code(csv_data, language="csv")
            else:
                st.info("上传 Excel 后点击生成")



def render_yearly_tab(budget, account_name=""):
    st.markdown(
        "上传从 Azure 定价计算器导出的原始 Excel，自动新增 **Estimated yearly cost** 列（月费用 × 12）并在 Total 行汇总。"
    )

    st.divider()

    uploaded_price = st.file_uploader(
        "上传原始价格表 (.xlsx)",
        type=["xlsx"],
        key="upload_price_excel",
        help="支持标准 Azure 定价计算器导出格式",
    )

    if uploaded_price is not None:
        if st.button("生成年度价格表", type="primary", use_container_width=True, key="btn_gen_yearly"):
            import openpyxl
            from copy import copy as _copy
            from openpyxl.styles import Font as _Font

            def _col_letter(n):
                result = ""
                while n:
                    n, rem = divmod(n - 1, 26)
                    result = chr(65 + rem) + result
                return result

            def _copy_cell_style(src, dst):
                if src.has_style:
                    dst.font      = _copy(src.font)
                    dst.fill      = _copy(src.fill)
                    dst.border    = _copy(src.border)
                    dst.alignment = _copy(src.alignment)
                    dst.number_format = src.number_format

            def _find_header_row(ws):
                for i, row in enumerate(ws.iter_rows(values_only=True), 1):
                    if row and "Estimated monthly cost" in row:
                        return i
                return None

            def _find_total_row(ws, hrow):
                for i, row in enumerate(ws.iter_rows(min_row=hrow + 1, values_only=True), hrow + 1):
                    if row and "Total" in row:
                        return i
                return None

            def _get_account_name(ws):
                """从 Sheet 第 2 行前 5 列取账号名（非空的第一个值）。"""
                for col in range(1, 6):
                    v = ws.cell(2, col).value
                    if v and str(v).strip():
                        return str(v).strip().rstrip("\t").strip()
                return None

            def _process_sheet(ws):
                hrow = _find_header_row(ws)
                if hrow is None:
                    return False, "未找到标题行（含 'Estimated monthly cost'）", None
                trow = _find_total_row(ws, hrow)
                if trow is None:
                    return False, "未找到 Total 行", None

                header_vals = [ws.cell(hrow, c).value for c in range(1, ws.max_column + 1)]
                try:
                    monthly_col = header_vals.index("Estimated monthly cost") + 1
                    upfront_col = header_vals.index("Estimated upfront cost") + 1
                except ValueError:
                    return False, "未找到必要列名", None

                yearly_col     = upfront_col + 1
                ws.insert_cols(yearly_col)
                monthly_letter = _col_letter(monthly_col)
                yearly_letter  = _col_letter(yearly_col)

                # 标题行：复制 upfront 列样式
                hcell = ws.cell(hrow, yearly_col, "Estimated yearly cost")
                _copy_cell_style(ws.cell(hrow, upfront_col), hcell)
                src_hdr = ws.cell(hrow, upfront_col)
                hcell.font = _Font(
                    name=src_hdr.font.name or "Calibri",
                    bold=True,
                    size=src_hdr.font.size or 11,
                )

                data_start = hrow + 1
                data_end   = trow - 1

                # 数据行：写公式，复制样式并特别保留 number_format（用于显示 $）
                for r in range(data_start, data_end + 1):
                    mv = ws.cell(r, monthly_col).value
                    if mv is not None and (isinstance(mv, (int, float)) or (isinstance(mv, str) and mv.startswith("="))):
                        cell = ws.cell(r, yearly_col)
                        cell.value = f"={monthly_letter}{r}*12"
                        src_cell = ws.cell(r, monthly_col)
                        _copy_cell_style(src_cell, cell)
                        # 显式保留原始单元格的 number_format，以带上 $ 符号
                        if src_cell.number_format and src_cell.number_format != 'General':
                            cell.number_format = src_cell.number_format
                        else:
                            cell.number_format = '"$"#,##0.00'
                    else:
                        ws.cell(r, yearly_col).value = None

                # Total 行
                tcell = ws.cell(trow, yearly_col)
                tcell.value = f"=SUM({yearly_letter}{data_start}:{yearly_letter}{data_end})"
                src_total = ws.cell(trow, monthly_col)
                _copy_cell_style(src_total, tcell)
                if src_total.number_format and src_total.number_format != 'General':
                    tcell.number_format = src_total.number_format
                else:
                    tcell.number_format = '"$"#,##0.00'
                tcell.font = _Font(bold=True, name="Calibri", size=11)

                ws.column_dimensions[yearly_letter].width = 22
                
                account = _get_account_name(ws)
                return True, "处理成功", account

            try:
                with st.spinner("正在处理 Excel..."):
                    wb = openpyxl.load_workbook(uploaded_price)
                    messages = []
                    account_name = None
                    for sname in wb.sheetnames:
                        ok, msg, acct = _process_sheet(wb[sname])
                        messages.append(f"**{sname}**: {msg}")
                        if acct and not account_name:
                            account_name = acct

                    # 优先使用用户输入的账户名，其次使用 Excel 中提取的名称
                    _budget = st.session_state.get("budget", budget) or "未填写"
                    _acct_from_input = st.session_state.get("account_name") or account_name.strip()
                    _acct_final = _acct_from_input or account_name or uploaded_price.name.replace(".xlsx", "")
                    new_dl_name = f"{_acct_final}-Azure calculator.xlsx"

                    out_buf = io.BytesIO()
                    wb.save(out_buf)
                    out_buf.seek(0)
                    st.session_state["yearly_excel_bytes"] = out_buf.getvalue()
                    st.session_state["yearly_excel_name"]  = new_dl_name
                    st.session_state["yearly_messages"]    = messages

                st.rerun()
            except Exception as e:
                st.error(f"处理失败：{e}")
    else:
        st.info("请先上传 Excel 文件")

    # 处理结果与下载
    if "yearly_excel_bytes" in st.session_state:
        st.divider()
        for msg in st.session_state.get("yearly_messages", []):
            st.markdown(msg)
        st.download_button(
            label="下载任务年度价格表 (.xlsx)",
            data=st.session_state["yearly_excel_bytes"],
            file_name=st.session_state["yearly_excel_name"],
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="dl_yearly",
        )
